"""Generate case27: image interpolation test with different resolution photos.

Tests three scaling scenarios:
- Low-res (80x60 px) upscaled ~3x
- Medium-res (400x300 px) at ~1:1 (72 dpi)
- High-res (1800x1200 px) downscaled ~10x
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document()

# Fix page to standard Letter
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.orientation = WD_ORIENT.PORTRAIT
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(6)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Image Interpolation Test")
run.bold = True
run.font.size = Pt(16)
run.font.name = "Arial"

doc.add_paragraph()

# --- 1. Low-res upscaled ~3x ---
p = doc.add_paragraph()
run = p.add_run("Low resolution (80\u00d760 px) at 3.3\u00d72.5 in (upscaled ~3x):")
run.bold = True
run.font.name = "Arial"

doc.add_picture("/tmp/test_images/small.jpg", width=Inches(3.33), height=Inches(2.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# --- 2. High-res downscaled ~10x ---
p = doc.add_paragraph()
run = p.add_run("High resolution (1800\u00d71200 px) at 2.5\u00d71.67 in (downscaled ~10x):")
run.bold = True
run.font.name = "Arial"

doc.add_picture("/tmp/test_images/large.jpg", width=Inches(2.5), height=Inches(1.67))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# --- 3. Medium-res at ~1:1 at 72dpi ---
p = doc.add_paragraph()
run = p.add_run("Medium resolution (400\u00d7300 px) at 5.5\u00d74.1 in (~1:1 at 72 dpi):")
run.bold = True
run.font.name = "Arial"

doc.add_picture("/tmp/test_images/medium.jpg", width=Inches(5.5), height=Inches(4.12))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

out = "tests/fixtures/cases/case27/input.docx"
doc.save(out)
print(f"Created {out}")
