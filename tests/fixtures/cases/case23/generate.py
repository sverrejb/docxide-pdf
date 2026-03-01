"""Generate case23: Character Spacing and Double Strikethrough

Tests:
- w:spacing w:val (character spacing / letter-spacing) at various values
- w:dstrike (double strikethrough)
- Combination of char spacing with other formatting
"""

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_spacing_run(paragraph, text, spacing_twips, bold=False, dstrike=False):
    """Add a run with explicit character spacing."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    rpr = run._element.get_or_add_rPr()
    if spacing_twips != 0:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing_twips))
        rpr.append(sp)
    if dstrike:
        ds = OxmlElement("w:dstrike")
        rpr.append(ds)
    return run


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Title
doc.add_heading("Character Spacing & Double Strikethrough", level=1)

# Section 1: Character spacing examples
doc.add_heading("Character Spacing", level=2)

p = doc.add_paragraph()
add_spacing_run(p, "Normal spacing (0pt). ", 0)
add_spacing_run(p, "Expanded +2pt. ", 40)
add_spacing_run(p, "Expanded +5pt. ", 100)

p = doc.add_paragraph()
add_spacing_run(p, "Condensed -1pt. ", -20)
add_spacing_run(p, "Normal again. ", 0)
add_spacing_run(p, "Condensed -0.5pt. ", -10)

p = doc.add_paragraph()
add_spacing_run(p, "This entire paragraph has expanded spacing of 3 points applied to every "
    "character. This tests how character spacing affects line breaking and text width "
    "calculations across a longer block of text.", 60)

p = doc.add_paragraph()
add_spacing_run(p, "Tight condensed text at -1.5pt spacing to test negative values.", -30)

# Section 2: Double strikethrough
doc.add_heading("Double Strikethrough", level=2)

p = doc.add_paragraph()
p.add_run("Normal text, then ")
add_spacing_run(p, "double strikethrough text", 0, dstrike=True)
p.add_run(", then normal again.")

p = doc.add_paragraph()
add_spacing_run(p, "Bold double strikethrough combined.", 0, bold=True, dstrike=True)

# Section 3: Combinations
doc.add_heading("Combined Formatting", level=2)

p = doc.add_paragraph()
add_spacing_run(p, "Expanded +3pt with double strike", 60, dstrike=True)
p.add_run(" mixed with ")
add_spacing_run(p, "condensed -1pt no strike", -20)
p.add_run(".")

p = doc.add_paragraph()
add_spacing_run(p, "A longer paragraph mixing expanded and condensed spacing. ", 40)
add_spacing_run(p, "This part is condensed. ", -20)
add_spacing_run(p, "Back to expanded. ", 40)
add_spacing_run(p, "And normal to finish.", 0)

# Save
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(
    out_buf, "w", zipfile.ZIP_DEFLATED
) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/settings.xml":
            tree = etree.fromstring(data)
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(
                tree, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
