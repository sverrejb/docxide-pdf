"""Generate a multi-section DOCX testing various page numbering styles for case28.

Section 1: Centered footer with "- PAGE -" format
Section 2: Right-aligned footer with "Page X of Y"
Section 3: Header page number (left) + footer page number (right)
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_field(paragraph, field_code):
    """Insert a field code (PAGE, NUMPAGES, etc.) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_end)


def set_section_margins(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


doc = Document()

# ── Section 1: Centered footer with "- PAGE -" ──

section1 = doc.sections[0]
set_section_margins(section1)

footer1 = section1.footer
footer1.is_linked_to_previous = False
p = footer1.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("- ")
run.font.size = Pt(10)
add_field(p, "PAGE")
run = p.add_run(" -")
run.font.size = Pt(10)

doc.add_heading("Section One: Centered Page Numbers", level=1)
doc.add_paragraph(
    "This section demonstrates centered page numbers in the footer using the "
    "classic dash-surrounded format. The page number appears centered at the "
    "bottom of each page, formatted as '- 1 -', '- 2 -', etc."
)
doc.add_paragraph(
    "Page numbering is one of the most common features in word processing "
    "documents. It helps readers navigate through multi-page documents and "
    "provides a reference point for discussions and reviews."
)
doc.add_paragraph(
    "The field code mechanism in OOXML uses w:fldChar elements to mark the "
    "beginning and end of a field, with w:instrText containing the field "
    "instruction. The PAGE instruction returns the current page number, while "
    "NUMPAGES returns the total number of pages in the document."
)

# ── Section 2: Right-aligned "Page X of Y" ──

doc.add_section(start_type=2)  # 2 = new page
section2 = doc.sections[1]
set_section_margins(section2)

footer2 = section2.footer
footer2.is_linked_to_previous = False
p = footer2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Page ")
run.font.size = Pt(10)
add_field(p, "PAGE")
run = p.add_run(" of ")
run.font.size = Pt(10)
add_field(p, "NUMPAGES")

doc.add_heading("Section Two: Page X of Y", level=1)
doc.add_paragraph(
    "This section uses a right-aligned footer showing 'Page X of Y' format. "
    "This is a common page numbering style in formal documents, reports, and "
    "legal filings where readers need to know both their current position and "
    "the total length of the document."
)
doc.add_paragraph(
    "The NUMPAGES field code returns the total page count for the entire "
    "document, not just the current section. This means in a 3-page document, "
    "all pages will show 'of 3' regardless of which section they belong to."
)
doc.add_paragraph(
    "Right-aligned page numbers are particularly common in headers of academic "
    "papers and business reports. They stay out of the way of the main content "
    "while remaining easily visible when thumbing through printed pages."
)

# ── Section 3: Header (left) + Footer (right) page numbers ──

doc.add_section(start_type=2)
section3 = doc.sections[2]
set_section_margins(section3)

header3 = section3.header
header3.is_linked_to_previous = False
p = header3.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Page ")
run.font.size = Pt(9)
run.italic = True
add_field(p, "PAGE")

footer3 = section3.footer
footer3.is_linked_to_previous = False
p = footer3.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Document total: ")
run.font.size = Pt(9)
add_field(p, "NUMPAGES")
run = p.add_run(" pages")
run.font.size = Pt(9)

doc.add_heading("Section Three: Header and Footer Numbers", level=1)
doc.add_paragraph(
    "This final section demonstrates page numbers in both the header and the "
    "footer simultaneously. The header shows the current page number in italic "
    "on the left side, while the footer shows the total page count on the right."
)
doc.add_paragraph(
    "Having page information in multiple locations is common in reference "
    "materials, manuals, and lengthy reports. Headers often carry the page "
    "number along with chapter or section titles, while footers may contain "
    "supplementary information like document version or total pages."
)
doc.add_paragraph(
    "This pattern tests that our renderer correctly handles field codes in both "
    "headers and footers within the same section, and that page numbering "
    "continues correctly across section boundaries."
)

doc.save("tests/fixtures/cases/case28/input.docx")
print("Generated tests/fixtures/cases/case28/input.docx")
