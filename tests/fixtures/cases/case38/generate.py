#!/usr/bin/env python3
"""case38: Gradient fill gallery — 6 shapes with diverse linear gradients.

Tests gradient color interpolation (linear-light blending via CalRGB).
Each shape is a floating wps:wsp anchor with a:gradFill.

Shapes:
1. Blue→Orange (90°) — the vaccine case gradient, should show pink midtone
2. Red→Cyan (0°) — maximum hue shift, most sensitive to interpolation space
3. 3-stop Red→Green→Blue (90°) — tests stitching function path
4. Black→White (90°) — control case, should look identical in both spaces
5. Dark Blue→Bright Yellow (45°) — large luminance range
6. Purple→Orange (180°) — another cross-hue blend
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

OUT = Path("tests/fixtures/cases/case38/input.docx")

# XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
GRAPHIC_DATA_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

# Gradient definitions: (label, angle_deg, stops)
# stops: list of (position_percent, srgb_hex)
GRADIENTS = [
    (
        "Blue to Orange\n90deg",
        90,
        [(0, "4472C4"), (100, "ED7D31")],
    ),
    (
        "Red to Cyan\n0deg",
        0,
        [(0, "FF0000"), (100, "00FFFF")],
    ),
    (
        "Red Green Blue\n3-stop 90deg",
        90,
        [(0, "FF0000"), (50, "00FF00"), (100, "0000FF")],
    ),
    (
        "Black to White\n90deg",
        90,
        [(0, "000000"), (100, "FFFFFF")],
    ),
    (
        "Dark Blue to Yellow\n45deg",
        45,
        [(0, "002060"), (100, "FFFF00")],
    ),
    (
        "Purple to Orange\n180deg",
        180,
        [(0, "7030A0"), (100, "ED7D31")],
    ),
]

# Layout constants (EMUs: 1 inch = 914400 EMU)
EMU_PER_INCH = 914400

COLS = 3
ROWS = 2
SHAPE_W_EMU = int(1.8 * EMU_PER_INCH)
SHAPE_H_EMU = int(1.4 * EMU_PER_INCH)

# Content area: 6.5" wide, 9" tall (letter page with 1" margins)
CONTENT_W = 6.5 * EMU_PER_INCH
CONTENT_H = 9.0 * EMU_PER_INCH
COL_PITCH = CONTENT_W / COLS
ROW_PITCH = CONTENT_H / ROWS


def gradient_fill_xml(stops, angle_deg):
    """Build a:gradFill XML from stops and angle."""
    gs_items = ""
    for pos_pct, color in stops:
        pos_ooxml = pos_pct * 1000  # OOXML uses 1/1000th of a percent
        gs_items += (
            f'<a:gs pos="{pos_ooxml}">'
            f'<a:srgbClr val="{color}"/>'
            f'</a:gs>'
        )
    # angle in 60,000ths of a degree
    ang = int(angle_deg * 60000)
    return (
        f'<a:gradFill>'
        f'<a:gsLst>{gs_items}</a:gsLst>'
        f'<a:lin ang="{ang}" scaled="0"/>'
        f'</a:gradFill>'
    )


def shape_anchor_xml(idx, label, angle_deg, stops):
    """Build wp:anchor XML for a single floating shape with gradient fill."""
    row = idx // COLS
    col = idx % COLS
    x_emu = int(col * COL_PITCH + (COL_PITCH - SHAPE_W_EMU) / 2)
    y_emu = int(row * ROW_PITCH + (ROW_PITCH - SHAPE_H_EMU) / 2)

    fill_xml = gradient_fill_xml(stops, angle_deg)

    # Build text body with label (split by newline for multi-line labels)
    label_lines = label.split("\n")
    text_paras = ""
    for line in label_lines:
        text_paras += (
            f'<w:p>'
            f'<w:pPr><w:jc w:val="center"/></w:pPr>'
            f'<w:r>'
            f'<w:rPr>'
            f'<w:sz w:val="18"/>'
            f'<w:b/>'
            f'<w:color w:val="000000"/>'
            f'</w:rPr>'
            f'<w:t>{line}</w:t>'
            f'</w:r>'
            f'</w:p>'
        )

    return (
        f'<w:r>'
        f'<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f' simplePos="0" relativeHeight="{idx}" behindDoc="0"'
        f' locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="margin">'
        f'<wp:posOffset>{x_emu}</wp:posOffset>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="margin">'
        f'<wp:posOffset>{y_emu}</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{idx + 1}" name="Gradient {idx + 1}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="{A_NS}">'
        f'<a:graphicData uri="{GRAPHIC_DATA_URI}">'
        f'<wps:wsp xmlns:wps="{WPS_NS}">'
        f'<wps:cNvSpPr/>'
        f'<wps:spPr>'
        f'<a:xfrm>'
        f'<a:off x="0" y="0"/>'
        f'<a:ext cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'</a:xfrm>'
        f'<a:prstGeom prst="rect">'
        f'<a:avLst/>'
        f'</a:prstGeom>'
        f'{fill_xml}'
        f'<a:ln w="12700">'
        f'<a:solidFill>'
        f'<a:srgbClr val="000000"/>'
        f'</a:solidFill>'
        f'</a:ln>'
        f'</wps:spPr>'
        f'<wps:txbx>'
        f'<w:txbxContent xmlns:w="{W_NS}">'
        f'{text_paras}'
        f'</w:txbxContent>'
        f'</wps:txbx>'
        f'<wps:bodyPr anchor="t" anchorCtr="0" lIns="91440" tIns="45720" rIns="91440" bIns="45720"/>'
        f'</wps:wsp>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
        f'</w:r>'
    )


# Step 1: Create base document with python-docx
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Single paragraph to hold all floating shape anchors
p = doc.add_paragraph()
run = p.add_run("GRADIENT_PLACEHOLDER")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Step 2: Post-process ZIP to inject gradient shapes
all_shapes_xml = "".join(
    shape_anchor_xml(i, label, angle, stops)
    for i, (label, angle, stops) in enumerate(GRADIENTS)
)

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        # Add namespace declarations to root element
        ns_decls = (
            f' xmlns:wp="{WP_NS}"'
            f' xmlns:a="{A_NS}"'
            f' xmlns:wps="{WPS_NS}"'
            f' xmlns:r="{R_NS}"'
            f' xmlns:mc="{MC_NS}"'
        )
        if 'xmlns:wp=' not in doc_xml:
            doc_xml = doc_xml.replace(
                '<w:document ',
                f'<w:document {ns_decls} ',
                1,
            )

        # Replace the placeholder run with shape anchors
        placeholder_pattern = r'<w:r>.*?<w:t>GRADIENT_PLACEHOLDER</w:t>\s*</w:r>'
        doc_xml = re.sub(
            placeholder_pattern,
            all_shapes_xml,
            doc_xml,
            count=1,
            flags=re.DOTALL,
        )

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)
print(f"Generated {OUT}")
