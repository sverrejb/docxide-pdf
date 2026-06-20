#!/usr/bin/env python3
"""case69: w:vAlign (Vertical Text Alignment on Page, OOXML 17.6.23) — center.

Isolates the section-level page vertical-justification feature: a sectPr child
`<w:vAlign w:val="center"/>` that vertically centers the section's body text
between the top and bottom margins (classic title-page layout). Distinct from
the cell-level w:vAlign already handled in tables. Content is deliberately tiny
so center-vs-top placement is unmistakable in the reference.
"""

import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case69/input.docx")

# 17.6.23: ST_VerticalJc = top (default) | center | both | bottom.
# In CT_SectPr child order, w:vAlign follows w:cols and precedes w:docGrid.
V_ALIGN = '<w:vAlign w:val="center"/>'


def main():
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = 1  # center horizontally so it reads as a title page
    run = title.add_run("Vertically Centered Title")
    run.bold = True
    run.font.size = Pt(24)
    doc.add_paragraph(
        "This short body should sit in the vertical middle of the page, "
        "not at the top, because the section sets w:vAlign=center."
    ).alignment = 1

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode("utf-8")
        other = {n: zin.read(n) for n in zin.namelist() if n != "word/document.xml"}

    # vAlign precedes w:docGrid in the CT_SectPr sequence; fall back to end.
    if "<w:docGrid" in doc_xml:
        doc_xml = doc_xml.replace("<w:docGrid", V_ALIGN + "<w:docGrid", 1)
    else:
        doc_xml = doc_xml.replace("</w:sectPr>", V_ALIGN + "</w:sectPr>", 1)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
