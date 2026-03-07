#!/usr/bin/env python3
"""case33: Clickable hyperlinks — external URLs, styled hyperlink text, mixed inline content."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Helper to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text, font_size=None, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Hyperlink style
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink

# 1. Simple hyperlink
p = doc.add_paragraph()
run = p.add_run("Visit ")
add_hyperlink(p, "https://www.example.com", "Example.com")
run2 = p.add_run(" for more information.")

# 2. Heading with description
doc.add_heading("Hyperlink Test Cases", level=1)

# 3. Multiple hyperlinks in one paragraph
p = doc.add_paragraph()
p.add_run("Search engines: ")
add_hyperlink(p, "https://www.google.com", "Google")
p.add_run(", ")
add_hyperlink(p, "https://www.bing.com", "Bing")
p.add_run(", and ")
add_hyperlink(p, "https://duckduckgo.com", "DuckDuckGo")
p.add_run(".")

# 4. Hyperlink with long URL text
p = doc.add_paragraph()
p.add_run("Full URL as text: ")
add_hyperlink(p, "https://en.wikipedia.org/wiki/Office_Open_XML", "https://en.wikipedia.org/wiki/Office_Open_XML")

# 5. Bold hyperlink
p = doc.add_paragraph()
p.add_run("Bold link: ")
add_hyperlink(p, "https://www.rust-lang.org", "The Rust Programming Language", bold=True)

# 6. Hyperlink in a bullet list
for label, url in [
    ("Rust documentation", "https://doc.rust-lang.org"),
    ("crates.io", "https://crates.io"),
    ("Rust playground", "https://play.rust-lang.org"),
]:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, label)

# 7. Paragraph with text before, link, and text after
p = doc.add_paragraph()
p.add_run("The PDF specification is maintained by ")
add_hyperlink(p, "https://www.iso.org", "ISO")
p.add_run(" and the latest version is ")
add_hyperlink(p, "https://pdfa.org/resource/iso-32000-2/", "ISO 32000-2:2020")
p.add_run(", which defines PDF 2.0.")

# 8. Larger font hyperlink
p = doc.add_paragraph()
run = p.add_run("Large link: ")
run.font.size = Pt(16)
add_hyperlink(p, "https://www.anthropic.com", "Anthropic", font_size=Pt(16))

doc.save("tests/fixtures/cases/case33/input.docx")
print("Generated case33/input.docx")
