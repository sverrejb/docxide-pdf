"""Generate case22: Three Columns with Separator Lines

Tests:
- Three equal-width columns (w:cols w:num="3" w:sep="1")
- Column separator lines between columns
- Natural text overflow from column to column
- More text to test multi-column flow across the page
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Title
p = doc.add_paragraph("Three-Column Newsletter")
p.style = doc.styles["Heading 1"]

# Enough text to flow across all three columns
paragraphs = [
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor "
    "incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud "
    "exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.",

    "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu "
    "fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
    "culpa qui officia deserunt mollit anim id est laborum.",

    "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium "
    "doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore "
    "veritatis et quasi architecto beatae vitae dicta sunt explicabo.",

    "Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, "
    "sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt.",

    "Neque porro quisquam est, qui dolorem ipsum quia dolor sit amet, consectetur, "
    "adipisci velit, sed quia non numquam eius modi tempora incidunt ut labore et "
    "dolore magnam aliquam quaerat voluptatem.",

    "Ut enim ad minima veniam, quis nostrum exercitationem ullam corporis suscipit "
    "laboriosam, nisi ut aliquid ex ea commodi consequatur. Quis autem vel eum iure "
    "reprehenderit qui in ea voluptate velit esse quam nihil molestiae consequatur.",

    "At vero eos et accusamus et iusto odio dignissimos ducimus qui blanditiis "
    "praesentium voluptatum deleniti atque corrupti quos dolores et quas molestias "
    "excepturi sint occaecati cupiditate non provident.",

    "Similique sunt in culpa qui officia deserunt mollitia animi, id est laborum et "
    "dolorum fuga. Et harum quidem rerum facilis est et expedita distinctio.",

    "Nam libero tempore, cum soluta nobis est eligendi optio cumque nihil impedit quo "
    "minus id quod maxime placeat facere possimus, omnis voluptas assumenda est, omnis "
    "dolor repellendus. Temporibus autem quibusdam et aut officiis debitis aut rerum "
    "necessitatibus saepe eveniet ut et voluptates repudiandae sint et molestiae non "
    "recusandae. Itaque earum rerum hic tenetur a sapiente delectus.",

    "Ut aut reiciendis voluptatibus maiores alias consequatur aut perferendis doloribus "
    "asperiores repellat. Hanc ego cum teneam sententiam, quid est cur verear ne ad eam "
    "non possim accommodare Torquatos nostros? Quos tu paulo ante cum memoriter, tum "
    "etiam erga nos amice et benivole collegisti.",

    "Quis est tam dissimile homini. Quodsi ipsam honestatem undique pertectam atque "
    "absolutam. Tecum optime, deinde etiam cum mediocri amico. De ingenio eius in his "
    "disputationibus, non de moribus quaeritur. Duo Reges: constructio interrete.",

    "Quamquam te quidem video minime esse deterritum. Primum in nostrane potestate est "
    "quid meminerimus? Quae cum dixisset, finem ille. Sed residamus, inquit, si placet. "
    "Eam tum adesse cum dolor omnis absit. Tum Torquatus: Prorsus, inquit, assentior.",

    "Idemne potest esse dies saepius qui semel fuit? Nihil est enim rerum omnium quod "
    "non natura elaborata concipere soleat. Haec para ad eos qui non student. Quae "
    "diligentissime contra Aristonem dicuntur a Chryippo. Quod ea non occurrentia "
    "fingunt, vincunt Stoicos.",

    "Conferam tecum, quam cuique verso rem subicias. Tum ille: Tu autem cum ipse tantum "
    "librorum habeas, quos hic tandem requiris? An hoc usque quaque, aliter in vita? "
    "Quae cum dixisset paulumque institisset. Quae similitudo in genere etiam humano "
    "apparet. Quid, de quo nulla dissensio est?",

    "Itaque hic ipse iam pridem est reiectus. Non quaero quid dicat, sed quid convenit. "
    "Potius inflammat, ut coercendi magis quam dedocendi esse videantur. Hoc loco "
    "tenere se Triarius non potuit. Quid iudicant sensus? Primum cur ista res digna "
    "odio est, nisi quod est turpis?",

    "Quid ergo aliud intellegetur nisi uti ne quae pars naturae neglegatur? Si longus, "
    "levis. Varietates autem iniurasque fortunae facile veterum philosophorum "
    "praeceptis instituta vita superabat. Ut alios omittam, hunc appello, quem ille "
    "unum secutus est.",
]

for text in paragraphs:
    doc.add_paragraph(text)

# Save, then post-process to set 3 columns with separator
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/document.xml":
            tree = etree.fromstring(data)
            nsmap = {"w": WML}

            # Replace w:cols in sectPr with 3-column separator config
            body = tree.find("w:body", nsmap)
            sect_pr = body.find("w:sectPr", nsmap)
            for old_cols in sect_pr.findall("w:cols", nsmap):
                sect_pr.remove(old_cols)
            cols = etree.SubElement(sect_pr, qn("w:cols"))
            cols.set(qn("w:num"), "3")
            cols.set(qn("w:space"), "480")  # ~1/3 inch gap
            cols.set(qn("w:sep"), "1")      # draw separator lines

            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        elif item.filename == "word/settings.xml":
            tree = etree.fromstring(data)
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
