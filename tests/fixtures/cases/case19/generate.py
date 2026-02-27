"""Generate case19: Nested Lists

Tests nested numbered and bullet lists with multiple indent levels,
counter restart on returning to parent level, and custom start values.
"""

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NUM_BASE = 100
ABS_BASE = 100


def make_abstract_num_xml(abstract_num_id, levels):
    lvls = []
    for ilvl, start, num_fmt, lvl_text, left, hanging in levels:
        lvls.append(
            f'<w:lvl w:ilvl="{ilvl}">'
            f'<w:start w:val="{start}"/>'
            f'<w:numFmt w:val="{num_fmt}"/>'
            f'<w:lvlText w:val="{lvl_text}"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="{left}" w:hanging="{hanging}"/></w:pPr>'
            f'</w:lvl>'
        )
    return (
        f'<w:abstractNum w:abstractNumId="{abstract_num_id}" xmlns:w="{WML}">'
        f'<w:multiLevelType w:val="hybridMultilevel"/>'
        + "".join(lvls)
        + '</w:abstractNum>'
    )


def make_num_xml(num_id, abstract_num_id):
    return (
        f'<w:num w:numId="{num_id}" xmlns:w="{WML}">'
        f'<w:abstractNumId w:val="{abstract_num_id}"/>'
        f'</w:num>'
    )


def add_list_para(doc, text, num_id, ilvl):
    p = doc.add_paragraph(text)
    ppr = p._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_pr.append(ilvl_el)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    ppr.append(num_pr)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# 3-level numbered list with counter restart
doc.add_heading("Nested Numbered List", level=2)
add_list_para(doc, "First item", NUM_BASE, 0)
add_list_para(doc, "Sub-item a", NUM_BASE, 1)
add_list_para(doc, "Sub-item b", NUM_BASE, 1)
add_list_para(doc, "Detail i", NUM_BASE, 2)
add_list_para(doc, "Detail ii", NUM_BASE, 2)
add_list_para(doc, "Sub-item c", NUM_BASE, 1)
add_list_para(doc, "Second item", NUM_BASE, 0)
add_list_para(doc, "Sub-item a again", NUM_BASE, 1)
add_list_para(doc, "Third item", NUM_BASE, 0)

# Custom start value
doc.add_heading("Custom Start", level=2)
add_list_para(doc, "Starts at five", NUM_BASE + 1, 0)
add_list_para(doc, "Then six", NUM_BASE + 1, 0)
add_list_para(doc, "Then seven", NUM_BASE + 1, 0)

# Save, then post-process to inject custom numbering
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/numbering.xml":
            tree = etree.fromstring(data)
            nsmap = {"w": WML}
            first_num = tree.find("w:num", nsmap)

            defs = [
                (ABS_BASE, [
                    (0, 1, "decimal",     "%1.", 720, 360),
                    (1, 1, "lowerLetter", "%2)", 1440, 360),
                    (2, 1, "lowerRoman",  "%3)", 2160, 360),
                ]),
                (ABS_BASE + 1, [
                    (0, 5, "decimal", "%1.", 720, 360),
                ]),
            ]
            for abs_id, levels in defs:
                el = etree.fromstring(make_abstract_num_xml(abs_id, levels))
                if first_num is not None:
                    tree.insert(list(tree).index(first_num), el)
                else:
                    tree.append(el)

            for i in range(2):
                tree.append(etree.fromstring(make_num_xml(NUM_BASE + i, ABS_BASE + i)))

            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        elif item.filename == "word/settings.xml":
            # Upgrade compatibility mode from 14 (Word 2010) to 15 (Word 2013+)
            tree = etree.fromstring(data)
            nsmap = {"w": WML}
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
