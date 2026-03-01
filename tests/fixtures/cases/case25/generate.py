"""Generate case25: Landscape orientation and non-Letter page sizes

Tests:
- Landscape US Letter (11 x 8.5")
- A4 portrait (210 x 297mm)
- A4 landscape (297 x 210mm)
- Legal portrait (8.5 x 14")
Each as a separate section with sample text to verify layout.
"""

from docx import Document
from docx.shared import Inches, Pt, Mm, Twips, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_section(doc, width, height, orient, margin=Inches(1)):
    """Add a new section with given page dimensions."""
    section = doc.add_section()
    section.page_width = width
    section.page_height = height
    section.orientation = orient
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    return section


doc = Document()

# Section 1: Landscape US Letter
section = doc.sections[0]
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.orientation = WD_ORIENT.LANDSCAPE
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

doc.add_heading("Landscape US Letter (11 x 8.5 inches)", level=1)
doc.add_paragraph(
    "This page uses US Letter paper in landscape orientation. The text area is "
    "9 inches wide by 6.5 inches tall. This tests that page dimensions are "
    "correctly swapped for landscape mode and that text reflows to fill the "
    "wider page. Words should wrap at approximately 9 inches."
)
doc.add_paragraph(
    "A second paragraph to add more content. The quick brown fox jumps over "
    "the lazy dog. Pack my box with five dozen liquor jugs. How vexingly "
    "quick daft zebras jump."
)

# Section 2: A4 portrait
section = add_section(doc, Mm(210), Mm(297), WD_ORIENT.PORTRAIT)

doc.add_heading("A4 Portrait (210 x 297 mm)", level=1)
doc.add_paragraph(
    "This page uses A4 paper in portrait orientation. A4 is 210mm wide by "
    "297mm tall, which is slightly narrower and taller than US Letter. "
    "European documents commonly use this size. Text should reflow to fit "
    "the narrower width compared to Letter."
)
doc.add_paragraph(
    "Additional content to verify vertical spacing and page height. "
    "The quick brown fox jumps over the lazy dog. Pack my box with five "
    "dozen liquor jugs."
)

# Section 3: A4 landscape
section = add_section(doc, Mm(297), Mm(210), WD_ORIENT.LANDSCAPE)

doc.add_heading("A4 Landscape (297 x 210 mm)", level=1)
doc.add_paragraph(
    "This page uses A4 paper rotated to landscape. The text area should be "
    "significantly wider than A4 portrait. This paragraph tests that the "
    "wider layout works correctly with A4 dimensions."
)

# Section 4: US Legal portrait
section = add_section(doc, Inches(8.5), Inches(14), WD_ORIENT.PORTRAIT)

doc.add_heading("US Legal Portrait (8.5 x 14 inches)", level=1)
doc.add_paragraph(
    "This page uses US Legal paper, which is 8.5 inches wide (same as Letter) "
    "but 14 inches tall (3 inches taller than Letter). Legal paper is commonly "
    "used for contracts and legal documents in the United States."
)
doc.add_paragraph(
    "More text to take up some vertical space on this taller page. The quick "
    "brown fox jumps over the lazy dog. Pack my box with five dozen liquor "
    "jugs. How vexingly quick daft zebras jump. The five boxing wizards jump "
    "quickly. Sphinx of black quartz, judge my vow."
)

# Save with compat mode 15
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(
    out_buf, "w", zipfile.ZIP_DEFLATED
) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/settings.xml":
            tree = etree.fromstring(data)
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(
                tree, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
