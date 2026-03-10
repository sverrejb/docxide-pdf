#!/usr/bin/env python3
"""case36: Custom geometry shapes — shapes defined with a:custGeom paths.

Tests custom geometry parsing and rendering: moveTo, lnTo, cubicBezTo, arcTo, close,
plus guide formulas (gdLst) and adjustment defaults (avLst).
Each shape is a floating wps:wsp anchor with a:custGeom instead of a:prstGeom.
"""

import math
import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

OUT = Path("tests/fixtures/cases/case36/input.docx")

# XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
GRAPHIC_DATA_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

EMU_PER_INCH = 914400

# Grid layout: 3 columns, 2 rows
COLS = 3
ROWS = 2
SHAPE_W_EMU = int(1.8 * EMU_PER_INCH)
SHAPE_H_EMU = int(1.8 * EMU_PER_INCH)
CONTENT_W = 6.5 * EMU_PER_INCH
CONTENT_H = 9.0 * EMU_PER_INCH
COL_PITCH = CONTENT_W / COLS
ROW_PITCH = CONTENT_H / ROWS

# Path coordinate space (same for all shapes for simplicity)
CS = 10000


def make_star_path():
    """6-pointed star using moveTo + lnTo + close. No guides needed."""
    cx, cy = CS // 2, CS // 2
    outer_r = CS // 2
    inner_r = CS // 4
    pts = []
    for i in range(6):
        # Outer point
        angle = math.radians(-90 + i * 60)
        pts.append((int(cx + outer_r * math.cos(angle)), int(cy + outer_r * math.sin(angle))))
        # Inner point
        angle2 = math.radians(-90 + i * 60 + 30)
        pts.append((int(cx + inner_r * math.cos(angle2)), int(cy + inner_r * math.sin(angle2))))

    cmds = f'<a:moveTo><a:pt x="{pts[0][0]}" y="{pts[0][1]}"/></a:moveTo>'
    for x, y in pts[1:]:
        cmds += f'<a:lnTo><a:pt x="{x}" y="{y}"/></a:lnTo>'
    cmds += '<a:close/>'
    return "", cmds


def make_heart_path():
    """Heart shape using moveTo + cubicBezTo + close."""
    # Heart: start at bottom point, two cubic bezier curves forming left and right lobes
    cmds = (
        '<a:moveTo><a:pt x="5000" y="9500"/></a:moveTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="5000" y="7500"/>'
        '  <a:pt x="0" y="4500"/>'
        '  <a:pt x="0" y="3000"/>'
        '</a:cubicBezTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="0" y="1000"/>'
        '  <a:pt x="2500" y="0"/>'
        '  <a:pt x="5000" y="2500"/>'
        '</a:cubicBezTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="7500" y="0"/>'
        '  <a:pt x="10000" y="1000"/>'
        '  <a:pt x="10000" y="3000"/>'
        '</a:cubicBezTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="10000" y="4500"/>'
        '  <a:pt x="5000" y="7500"/>'
        '  <a:pt x="5000" y="9500"/>'
        '</a:cubicBezTo>'
        '<a:close/>'
    )
    return "", cmds


def make_arrow_path():
    """Block arrow using moveTo + lnTo + close. Uses guide formulas."""
    # Guides compute shaft width and head start from adjustments
    guides = (
        '<a:gdLst>'
        '  <a:gd name="shaft" fmla="*/ h adj1 100000"/>'
        '  <a:gd name="headStart" fmla="*/ w adj2 100000"/>'
        '  <a:gd name="shaftTop" fmla="+- vc 0 shaft"/>'
        '  <a:gd name="shaftBot" fmla="+- vc shaft 0"/>'
        '</a:gdLst>'
    )
    avlst = (
        '<a:avLst>'
        '  <a:gd name="adj1" fmla="val 25000"/>'
        '  <a:gd name="adj2" fmla="val 60000"/>'
        '</a:avLst>'
    )
    cmds = (
        '<a:moveTo><a:pt x="0" y="shaftTop"/></a:moveTo>'
        '<a:lnTo><a:pt x="headStart" y="shaftTop"/></a:lnTo>'
        '<a:lnTo><a:pt x="headStart" y="0"/></a:lnTo>'
        '<a:lnTo><a:pt x="w" y="vc"/></a:lnTo>'
        '<a:lnTo><a:pt x="headStart" y="h"/></a:lnTo>'
        '<a:lnTo><a:pt x="headStart" y="shaftBot"/></a:lnTo>'
        '<a:lnTo><a:pt x="0" y="shaftBot"/></a:lnTo>'
        '<a:close/>'
    )
    return avlst + guides, cmds


def make_wave_path():
    """Wave/ribbon using moveTo + cubicBezTo + lnTo + close."""
    cmds = (
        '<a:moveTo><a:pt x="0" y="5000"/></a:moveTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="2500" y="0"/>'
        '  <a:pt x="5000" y="0"/>'
        '  <a:pt x="5000" y="5000"/>'
        '</a:cubicBezTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="5000" y="10000"/>'
        '  <a:pt x="7500" y="10000"/>'
        '  <a:pt x="10000" y="5000"/>'
        '</a:cubicBezTo>'
        '<a:lnTo><a:pt x="10000" y="7000"/></a:lnTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="7500" y="10000"/>'
        '  <a:pt x="5000" y="10000"/>'
        '  <a:pt x="5000" y="7000"/>'
        '</a:cubicBezTo>'
        '<a:cubicBezTo>'
        '  <a:pt x="5000" y="2000"/>'
        '  <a:pt x="2500" y="2000"/>'
        '  <a:pt x="0" y="7000"/>'
        '</a:cubicBezTo>'
        '<a:close/>'
    )
    return "", cmds


def make_cross_path():
    """Plus/cross shape using guide formulas for arm thickness."""
    guides = (
        '<a:gdLst>'
        '  <a:gd name="arm" fmla="*/ w adj1 100000"/>'
        '  <a:gd name="x1" fmla="+- hc 0 arm"/>'
        '  <a:gd name="x2" fmla="+- hc arm 0"/>'
        '  <a:gd name="y1" fmla="+- vc 0 arm"/>'
        '  <a:gd name="y2" fmla="+- vc arm 0"/>'
        '</a:gdLst>'
    )
    avlst = (
        '<a:avLst>'
        '  <a:gd name="adj1" fmla="val 20000"/>'
        '</a:avLst>'
    )
    cmds = (
        '<a:moveTo><a:pt x="x1" y="0"/></a:moveTo>'
        '<a:lnTo><a:pt x="x2" y="0"/></a:lnTo>'
        '<a:lnTo><a:pt x="x2" y="y1"/></a:lnTo>'
        '<a:lnTo><a:pt x="w" y="y1"/></a:lnTo>'
        '<a:lnTo><a:pt x="w" y="y2"/></a:lnTo>'
        '<a:lnTo><a:pt x="x2" y="y2"/></a:lnTo>'
        '<a:lnTo><a:pt x="x2" y="h"/></a:lnTo>'
        '<a:lnTo><a:pt x="x1" y="h"/></a:lnTo>'
        '<a:lnTo><a:pt x="x1" y="y2"/></a:lnTo>'
        '<a:lnTo><a:pt x="0" y="y2"/></a:lnTo>'
        '<a:lnTo><a:pt x="0" y="y1"/></a:lnTo>'
        '<a:lnTo><a:pt x="x1" y="y1"/></a:lnTo>'
        '<a:close/>'
    )
    return avlst + guides, cmds


def make_rounded_rect_path():
    """Rounded rectangle using arcTo commands + lnTo."""
    # Uses arcTo for corners: 90-degree arcs at each corner
    # Coordinate space: 10000x10000, corner radius = 1500
    guides = (
        '<a:gdLst>'
        '  <a:gd name="r" fmla="val 1500"/>'
        '  <a:gd name="x1" fmla="+- w 0 r"/>'
        '  <a:gd name="y1" fmla="+- h 0 r"/>'
        '</a:gdLst>'
    )
    # Start at top-left after first arc, go clockwise
    # stAng/swAng in 60000ths of a degree
    cmds = (
        # Start at (r, 0) — top edge, right of top-left corner
        '<a:moveTo><a:pt x="r" y="0"/></a:moveTo>'
        # Top edge to top-right corner
        '<a:lnTo><a:pt x="x1" y="0"/></a:lnTo>'
        # Top-right corner arc (90° clockwise from 270°)
        '<a:arcTo wR="r" hR="r" stAng="16200000" swAng="5400000"/>'
        # Right edge to bottom-right corner
        '<a:lnTo><a:pt x="w" y="y1"/></a:lnTo>'
        # Bottom-right corner arc
        '<a:arcTo wR="r" hR="r" stAng="0" swAng="5400000"/>'
        # Bottom edge to bottom-left corner
        '<a:lnTo><a:pt x="r" y="h"/></a:lnTo>'
        # Bottom-left corner arc
        '<a:arcTo wR="r" hR="r" stAng="5400000" swAng="5400000"/>'
        # Left edge to top-left corner
        '<a:lnTo><a:pt x="0" y="r"/></a:lnTo>'
        # Top-left corner arc
        '<a:arcTo wR="r" hR="r" stAng="10800000" swAng="5400000"/>'
        '<a:close/>'
    )
    return guides, cmds


# Shape definitions: (label, color, extra_xml_before_path, path_commands, path_w, path_h)
SHAPES = [
    ("Star (lnTo)", "4472C4", *make_star_path(), CS, CS),
    ("Heart (cubicBez)", "C00000", *make_heart_path(), CS, CS),
    ("Arrow (guides)", "264478", *make_arrow_path(), None, None),
    ("Wave (mixed)", "70AD47", *make_wave_path(), CS, CS),
    ("Cross (guides)", "7030A0", *make_cross_path(), None, None),
    ("RndRect (arcTo)", "ED7D31", *make_rounded_rect_path(), None, None),
]


def shape_anchor_xml(idx, label, color, extra_xml, path_cmds, path_w, path_h):
    """Build wp:anchor XML for a floating shape with a:custGeom."""
    row = idx // COLS
    col = idx % COLS
    x_emu = int(col * COL_PITCH + (COL_PITCH - SHAPE_W_EMU) / 2)
    y_emu = int(row * ROW_PITCH + (ROW_PITCH - SHAPE_H_EMU) / 2)

    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    luma = 0.299 * r + 0.587 * g + 0.114 * b
    text_color = "FFFFFF" if luma < 140 else "000000"

    path_attrs = ""
    if path_w is not None:
        path_attrs += f' w="{path_w}"'
    if path_h is not None:
        path_attrs += f' h="{path_h}"'

    return (
        f'<w:r>'
        f'<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f' simplePos="0" relativeHeight="{idx}" behindDoc="0"'
        f' locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="margin">'
        f'<wp:posOffset>{x_emu}</wp:posOffset>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="margin">'
        f'<wp:posOffset>{y_emu}</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{idx + 1}" name="Shape {idx + 1}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="{A_NS}">'
        f'<a:graphicData uri="{GRAPHIC_DATA_URI}">'
        f'<wps:wsp xmlns:wps="{WPS_NS}">'
        f'<wps:cNvSpPr/>'
        f'<wps:spPr>'
        f'<a:xfrm>'
        f'<a:off x="0" y="0"/>'
        f'<a:ext cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'</a:xfrm>'
        f'<a:custGeom>'
        f'{extra_xml}'
        f'<a:pathLst>'
        f'<a:path{path_attrs}>'
        f'{path_cmds}'
        f'</a:path>'
        f'</a:pathLst>'
        f'</a:custGeom>'
        f'<a:solidFill>'
        f'<a:srgbClr val="{color}"/>'
        f'</a:solidFill>'
        f'<a:ln w="12700">'
        f'<a:solidFill>'
        f'<a:srgbClr val="000000"/>'
        f'</a:solidFill>'
        f'</a:ln>'
        f'</wps:spPr>'
        f'<wps:txbx>'
        f'<w:txbxContent xmlns:w="{W_NS}">'
        f'<w:p>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:sz w:val="14"/>'
        f'<w:color w:val="{text_color}"/>'
        f'</w:rPr>'
        f'<w:t>{label}</w:t>'
        f'</w:r>'
        f'</w:p>'
        f'</w:txbxContent>'
        f'</wps:txbx>'
        f'<wps:bodyPr anchor="ctr" anchorCtr="0" lIns="0" tIns="0" rIns="0" bIns="0"/>'
        f'</wps:wsp>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
        f'</w:r>'
    )


# Step 1: Create base document with python-docx
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p = doc.add_paragraph()
run = p.add_run("SHAPE_PLACEHOLDER")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Step 2: Post-process ZIP to inject custom geometry shapes
all_shapes_xml = "".join(
    shape_anchor_xml(i, label, color, extra, cmds, pw, ph)
    for i, (label, color, extra, cmds, pw, ph) in enumerate(SHAPES)
)

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        ns_decls = (
            f' xmlns:wp="{WP_NS}"'
            f' xmlns:a="{A_NS}"'
            f' xmlns:wps="{WPS_NS}"'
            f' xmlns:r="{R_NS}"'
            f' xmlns:mc="{MC_NS}"'
        )
        if 'xmlns:wp=' not in doc_xml:
            doc_xml = doc_xml.replace(
                '<w:document ',
                f'<w:document {ns_decls} ',
                1,
            )

        placeholder_pattern = r'<w:r>.*?<w:t>SHAPE_PLACEHOLDER</w:t>\s*</w:r>'
        doc_xml = re.sub(
            placeholder_pattern,
            all_shapes_xml,
            doc_xml,
            count=1,
            flags=re.DOTALL,
        )

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)
print(f"Generated {OUT}")
