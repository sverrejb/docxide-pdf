"""
Case 61: gridSpan inference edge cases

Tests tables where gridSpan is absent from cells, requiring inference from
tcW vs tblGrid column widths. Covers:
  1. Explicit gridSpan (control) — equal columns
  2. Narrow column trap — cell width matches single col, but narrow neighbor
     makes span=2 cumulative almost match too
  3. Multiple narrow columns — inference across several thin columns
  4. Asymmetric last row — like japanese_interlibrary_loan: 7-col grid,
     most rows use [1][1][span=5], last row uses [span=3][1][span=2][1]
"""

from docx import Document
from docx.shared import Pt, Twips
from docx.enum.table import WD_TABLE_ALIGNMENT
import zipfile, os
from lxml import etree

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns = {'w': NS}

def set_grid(tbl_elem, widths_twips):
    """Replace tblGrid with custom column widths."""
    old = tbl_elem.find('w:tblGrid', ns)
    if old is not None:
        tbl_elem.remove(old)
    grid = etree.SubElement(tbl_elem, f'{{{NS}}}tblGrid')
    # Insert grid right after tblPr
    tbl_pr = tbl_elem.find('w:tblPr', ns)
    if tbl_pr is not None:
        tbl_pr.addnext(grid)
    for w in widths_twips:
        col = etree.SubElement(grid, f'{{{NS}}}gridCol')
        col.set(f'{{{NS}}}w', str(w))

def set_cell_widths(row_elem, widths_twips):
    """Set tcW on each cell in a row."""
    cells = row_elem.findall('w:tc', ns)
    for cell, w in zip(cells, widths_twips):
        tc_pr = cell.find('w:tcPr', ns)
        if tc_pr is None:
            tc_pr = etree.SubElement(cell, f'{{{NS}}}tcPr')
            cell.insert(0, tc_pr)
        tcw = tc_pr.find('w:tcW', ns)
        if tcw is None:
            tcw = etree.SubElement(tc_pr, f'{{{NS}}}tcW')
        tcw.set(f'{{{NS}}}w', str(w))
        tcw.set(f'{{{NS}}}type', 'dxa')

def strip_gridspan(row_elem):
    """Remove all gridSpan elements from a row."""
    for gs in row_elem.findall('.//w:gridSpan', ns):
        gs.getparent().remove(gs)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

doc.add_paragraph('Case 61: gridSpan inference', style='Heading 1')

# ── Table 1: Control — all gridSpans explicit ──
doc.add_paragraph('Table 1: Explicit gridSpan (control)', style='Heading 2')
t1 = doc.add_table(rows=3, cols=5)
t1.style = 'Table Grid'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, text in enumerate(['A', 'B', 'C', 'D', 'E']):
    t1.cell(0, i).text = text
t1.cell(1, 0).merge(t1.cell(1, 1)).text = 'AB merged'
t1.cell(1, 2).text = 'C alone'
t1.cell(1, 3).merge(t1.cell(1, 4)).text = 'DE merged'
t1.cell(2, 0).merge(t1.cell(2, 2)).text = 'ABC merged'
t1.cell(2, 3).merge(t1.cell(2, 4)).text = 'DE merged'

# ── Table 2: Narrow column trap ──
# Grid: [2000, 2000, 200, 2000, 2000] twips
# Row 1: [2000][2000][200+2000=2200][2000] — 4 cells, cell 2 spans cols 2+3
# The trap: cell 1 width=2000, col1=2000, cumulative(col1+col2)=2200.
# With 10% tolerance, 2200 is within 10% of 2000 (diff=200 < 220) — WRONG.
# With closest-match, col1 alone (diff=0) beats col1+col2 (diff=200) — CORRECT.
doc.add_paragraph('Table 2: Narrow column trap', style='Heading 2')
t2 = doc.add_table(rows=3, cols=5)
t2.style = 'Table Grid'
for i, text in enumerate(['Col0', 'Col1', 'Narrow', 'Col3', 'Col4']):
    t2.cell(0, i).text = text
t2.cell(1, 0).text = 'A'
t2.cell(1, 1).text = 'B'
t2.cell(1, 2).merge(t2.cell(1, 3)).text = 'Merged C+D'
t2.cell(1, 4).text = 'E'
t2.cell(2, 0).merge(t2.cell(2, 1)).text = 'AB merged'
t2.cell(2, 2).merge(t2.cell(2, 3)).text = 'CD merged'
t2.cell(2, 4).text = 'E alone'

# ── Table 3: Multiple narrow columns ──
# Grid: [3000, 150, 150, 150, 3000] twips
doc.add_paragraph('Table 3: Multiple narrow columns', style='Heading 2')
t3 = doc.add_table(rows=2, cols=5)
t3.style = 'Table Grid'
for i, text in enumerate(['Wide1', 'N1', 'N2', 'N3', 'Wide2']):
    t3.cell(0, i).text = text
t3.cell(1, 0).text = 'Left'
t3.cell(1, 1).merge(t3.cell(1, 4)).text = 'Rest merged (4 cols)'

# ── Table 4: Asymmetric last row (like japanese_interlibrary_loan) ──
# Grid: [600, 1400, 550, 2550, 210, 2340, 3400] twips
doc.add_paragraph('Table 4: Asymmetric rows (7-col grid)', style='Heading 2')
t4 = doc.add_table(rows=4, cols=7)
t4.style = 'Table Grid'
for i, text in enumerate(['C0', 'C1', 'C2', 'C3', 'C4', 'C5', 'C6']):
    t4.cell(0, i).text = text
for r in [1, 2]:
    t4.cell(r, 0).text = f'R{r}C0'
    t4.cell(r, 1).text = f'R{r}C1'
    t4.cell(r, 2).merge(t4.cell(r, 6)).text = f'R{r} cols 2-6 merged'
t4.cell(3, 0).merge(t4.cell(3, 2)).text = 'Span 0-2'
t4.cell(3, 3).text = 'Col3 alone'
t4.cell(3, 4).merge(t4.cell(3, 5)).text = 'Span 4-5'
t4.cell(3, 6).text = 'Col6 alone'

# Save initial
script_dir = os.path.dirname(os.path.abspath(__file__))
tmp_path = os.path.join(script_dir, '_tmp.docx')
final_path = os.path.join(script_dir, 'input.docx')
doc.save(tmp_path)

# Post-process: set custom grid widths, cell widths, and strip gridSpan
with zipfile.ZipFile(tmp_path, 'r') as zin:
    with zipfile.ZipFile(final_path, 'w') as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                root = etree.fromstring(data)
                tables = root.findall('.//w:tbl', ns)

                # Table 1: equal columns, keep all gridSpans (control)
                if len(tables) > 0:
                    set_grid(tables[0], [1728, 1728, 1728, 1728, 1728])

                # Table 2: narrow column trap
                if len(tables) > 1:
                    tbl = tables[1]
                    set_grid(tbl, [2000, 2000, 200, 2000, 2000])
                    rows = tbl.findall('w:tr', ns)
                    # Row 0: 5 cells with individual widths
                    set_cell_widths(rows[0], [2000, 2000, 200, 2000, 2000])
                    # Row 1: [2000][2000][2200][2000] — strip gridSpan
                    set_cell_widths(rows[1], [2000, 2000, 2200, 2000])
                    strip_gridspan(rows[1])
                    # Row 2: [4000][2200][2000] — strip gridSpan
                    set_cell_widths(rows[2], [4000, 2200, 2000])
                    strip_gridspan(rows[2])

                # Table 3: multiple narrow columns
                if len(tables) > 2:
                    tbl = tables[2]
                    set_grid(tbl, [3000, 150, 150, 150, 3000])
                    rows = tbl.findall('w:tr', ns)
                    set_cell_widths(rows[0], [3000, 150, 150, 150, 3000])
                    # Row 1: [3000][3450] — strip gridSpan
                    set_cell_widths(rows[1], [3000, 3450])
                    strip_gridspan(rows[1])

                # Table 4: asymmetric (like japanese form)
                if len(tables) > 3:
                    tbl = tables[3]
                    grid = [600, 1400, 550, 2550, 210, 2340, 3400]
                    set_grid(tbl, grid)
                    rows = tbl.findall('w:tr', ns)
                    # Row 0: 7 individual cells
                    set_cell_widths(rows[0], grid)
                    # Row 1: [600][1400][550+2550+210+2340+3400=9050] — strip gridSpan
                    set_cell_widths(rows[1], [600, 1400, 9050])
                    strip_gridspan(rows[1])
                    # Row 2: same as row 1
                    set_cell_widths(rows[2], [600, 1400, 9050])
                    strip_gridspan(rows[2])
                    # Row 3: [2550][2550][2550][3400] — strip gridSpan
                    set_cell_widths(rows[3], [2550, 2550, 2550, 3400])
                    strip_gridspan(rows[3])

                data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
            zout.writestr(item, data)

os.remove(tmp_path)
print(f'Created {final_path}')
