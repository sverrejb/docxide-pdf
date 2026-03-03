"""Generate a DOCX with line, pie, and area chart variations for case30.

Chart variations:
- Chart 1: Line chart (3 series, 6 data points, legend right)
- Chart 2: Pie chart (single series, 5 slices, legend right)
- Chart 3: Area chart (2 series, 6 data points, legend bottom)
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT = Path("tests/fixtures/cases/case30/input.docx")

CHART_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
DML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CHART_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
CT_CHART = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"


def build_series_xml(idx, label, color_hex, categories, values):
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{c}</c:v></c:pt>' for i, c in enumerate(categories)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill>
    <a:ln><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill></a:ln>
  </c:spPr>
  <c:cat><c:strRef><c:f>cats</c:f>
    <c:strCache><c:ptCount val="{len(categories)}"/>{cat_pts}</c:strCache>
  </c:strRef></c:cat>
  <c:val><c:numRef><c:f>{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
  </c:numRef></c:val>
</c:ser>"""


def build_axis_xml(ax_id, cross_id, position, axis_type="cat", gridlines=False, line_color="b3b3b3"):
    grid_xml = ""
    if gridlines:
        grid_xml = f"""<c:majorGridlines><c:spPr><a:ln>
      <a:solidFill><a:srgbClr val="{line_color}"/></a:solidFill>
    </a:ln></c:spPr></c:majorGridlines>"""
    tag = "c:valAx" if axis_type == "val" else "c:catAx"
    return f"""<{tag}>
  <c:axId val="{ax_id}"/>
  <c:scaling><c:orientation val="minMax"/></c:scaling>
  <c:delete val="0"/>
  <c:axPos val="{position}"/>
  {grid_xml}
  <c:majorTickMark val="out"/><c:minorTickMark val="none"/>
  <c:tickLblPos val="nextTo"/>
  <c:spPr><a:ln><a:solidFill><a:srgbClr val="{line_color}"/></a:solidFill></a:ln></c:spPr>
  <c:crossAx val="{cross_id}"/>
  <c:crossesAt val="0"/>
</{tag}>"""


def build_line_chart_xml(series_list, categories, legend_pos="r", colors=None):
    default_colors = ["4472C4", "ED7D31", "A5A5A5", "FFC000", "5B9BD5", "70AD47"]
    if colors is None:
        colors = default_colors

    series_xml = ""
    for i, (label, values) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_series_xml(i, label, color, categories, values)

    cat_ax = build_axis_xml(100, 200, "b", axis_type="cat")
    val_ax = build_axis_xml(200, 100, "l", axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:lineChart>
        <c:grouping val="standard"/>
        {series_xml}
        <c:axId val="100"/><c:axId val="200"/>
      </c:lineChart>
      {cat_ax}
      {val_ax}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_pie_chart_xml(labels, values, legend_pos="r"):
    # Pie chart has a single series with per-slice categories
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{l}</c:v></c:pt>' for i, l in enumerate(labels)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:pieChart>
        <c:varyColors val="1"/>
        <c:ser>
          <c:idx val="0"/><c:order val="0"/>
          <c:tx><c:strRef><c:f>label</c:f>
            <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>Market Share</c:v></c:pt></c:strCache>
          </c:strRef></c:tx>
          <c:cat><c:strRef><c:f>cats</c:f>
            <c:strCache><c:ptCount val="{len(labels)}"/>{cat_pts}</c:strCache>
          </c:strRef></c:cat>
          <c:val><c:numRef><c:f>0</c:f>
            <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
          </c:numRef></c:val>
        </c:ser>
      </c:pieChart>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_area_chart_xml(series_list, categories, legend_pos="b", colors=None):
    default_colors = ["4472C4", "ED7D31", "A5A5A5", "FFC000"]
    if colors is None:
        colors = default_colors

    series_xml = ""
    for i, (label, values) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_series_xml(i, label, color, categories, values)

    cat_ax = build_axis_xml(100, 200, "b", axis_type="cat")
    val_ax = build_axis_xml(200, 100, "l", axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:areaChart>
        <c:grouping val="standard"/>
        {series_xml}
        <c:axId val="100"/><c:axId val="200"/>
      </c:areaChart>
      {cat_ax}
      {val_ax}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_drawing_xml(rel_id, cx_emu, cy_emu):
    return (
        f'<w:drawing xmlns:w="{W_NS}" xmlns:wp="{WP_NS}" '
        f'xmlns:a="{DML_NS}" xmlns:c="{CHART_NS}" xmlns:r="{REL_NS}">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx_emu}" cy="{cy_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{rel_id.replace("rId","")}" name="Chart {rel_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic><a:graphicData uri="{CHART_NS}">'
        f'<c:chart r:id="{rel_id}"/>'
        f'</a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing>'
    )


# ── Step 1: Build document ──

doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

doc.add_heading("Line, Pie, and Area Charts", level=1)

doc.add_paragraph("Chart 1: Line chart with 3 series tracking monthly trends")
p1 = doc.add_paragraph()
p1.add_run("CHART_PLACEHOLDER_1")

doc.add_paragraph("")

doc.add_paragraph("Chart 2: Pie chart showing market share distribution")
p2 = doc.add_paragraph()
p2.add_run("CHART_PLACEHOLDER_2")

doc.add_paragraph("")

doc.add_paragraph("Chart 3: Area chart comparing two metrics over time")
p3 = doc.add_paragraph()
p3.add_run("CHART_PLACEHOLDER_3")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# ── Step 2: Inject charts ──

charts = {
    1: build_line_chart_xml(
        categories=["Jan", "Feb", "Mar", "Apr", "May", "Jun"],
        series_list=[
            ("Website", [1200, 1350, 1100, 1500, 1800, 2100]),
            ("Mobile App", [800, 950, 1050, 1200, 1400, 1650]),
            ("Desktop", [600, 580, 620, 550, 500, 480]),
        ],
        legend_pos="r",
        colors=["4472C4", "ED7D31", "70AD47"],
    ),
    2: build_pie_chart_xml(
        labels=["Chrome", "Safari", "Firefox", "Edge", "Other"],
        values=[63, 20, 8, 5, 4],
        legend_pos="r",
    ),
    3: build_area_chart_xml(
        categories=["Q1 '23", "Q2 '23", "Q3 '23", "Q4 '23", "Q1 '24", "Q2 '24"],
        series_list=[
            ("Revenue", [320, 380, 410, 450, 520, 590]),
            ("Costs", [280, 310, 340, 360, 390, 420]),
        ],
        legend_pos="b",
        colors=["4472C4", "ED7D31"],
    ),
}

chart_sizes = {
    1: (4572000, 2743200),  # ~5.0 x 3.0 inches
    2: (4572000, 2743200),  # ~5.0 x 3.0 inches
    3: (5486400, 2743200),  # ~6.0 x 3.0 inches
}

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        rels_xml = zin.read("word/_rels/document.xml.rels").decode()
        ct_xml = zin.read("[Content_Types].xml").decode()

        existing_rids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_xml)]
        next_rid = max(existing_rids, default=0) + 1

        chart_rids = {}
        for chart_num in sorted(charts.keys()):
            rid = f"rId{next_rid}"
            chart_rids[chart_num] = rid
            next_rid += 1

        new_rels = ""
        for chart_num, rid in chart_rids.items():
            new_rels += (
                f'<Relationship Id="{rid}" '
                f'Type="{CHART_REL_TYPE}" '
                f'Target="charts/chart{chart_num}.xml"/>'
            )
        rels_xml = rels_xml.replace("</Relationships>", new_rels + "</Relationships>")

        new_ct = ""
        for chart_num in charts:
            new_ct += (
                f'<Override PartName="/word/charts/chart{chart_num}.xml" '
                f'ContentType="{CT_CHART}"/>'
            )
        ct_xml = ct_xml.replace("</Types>", new_ct + "</Types>")

        doc_xml = zin.read("word/document.xml").decode()
        for chart_num, rid in chart_rids.items():
            placeholder = f"CHART_PLACEHOLDER_{chart_num}"
            cx, cy = chart_sizes[chart_num]
            drawing = build_drawing_xml(rid, cx, cy)
            run_pattern = f'<w:r><w:rPr></w:rPr><w:t>{placeholder}</w:t></w:r>'
            run_replacement = f'<w:r>{drawing}</w:r>'
            if run_pattern in doc_xml:
                doc_xml = doc_xml.replace(run_pattern, run_replacement)
            else:
                run_pattern2 = f'<w:r><w:t>{placeholder}</w:t></w:r>'
                if run_pattern2 in doc_xml:
                    doc_xml = doc_xml.replace(run_pattern2, run_replacement)
                else:
                    doc_xml = re.sub(
                        rf'<w:r[^>]*>.*?<w:t[^>]*>{placeholder}</w:t>.*?</w:r>',
                        run_replacement,
                        doc_xml,
                        flags=re.DOTALL,
                    )

        for item in zin.infolist():
            if item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, ct_xml)
            elif item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

        for chart_num, chart_xml in charts.items():
            zout.writestr(f"word/charts/chart{chart_num}.xml", chart_xml)

os.unlink(tmp)
print(f"Generated {OUT}")
