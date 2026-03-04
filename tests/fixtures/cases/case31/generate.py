"""Generate a DOCX with scatter, doughnut, radar, and bubble charts for case31.

Chart variations:
- Chart 1: Scatter chart (2 series, XY data points, legend right)
- Chart 2: Doughnut chart (single series, 5 slices, legend right)
- Chart 3: Radar chart (3 series, 5 categories, legend bottom)
- Chart 4: Bubble chart (2 series, XY + size, legend right)
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT = Path("tests/fixtures/cases/case31/input.docx")

CHART_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
DML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CHART_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
CT_CHART = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"


def build_axis_xml(ax_id, cross_id, position, axis_type="cat", gridlines=False, line_color="b3b3b3"):
    grid_xml = ""
    if gridlines:
        grid_xml = f"""<c:majorGridlines><c:spPr><a:ln>
      <a:solidFill><a:srgbClr val="{line_color}"/></a:solidFill>
    </a:ln></c:spPr></c:majorGridlines>"""
    tag = "c:valAx" if axis_type == "val" else "c:catAx"
    return f"""<{tag}>
  <c:axId val="{ax_id}"/>
  <c:scaling><c:orientation val="minMax"/></c:scaling>
  <c:delete val="0"/>
  <c:axPos val="{position}"/>
  {grid_xml}
  <c:majorTickMark val="out"/><c:minorTickMark val="none"/>
  <c:tickLblPos val="nextTo"/>
  <c:spPr><a:ln><a:solidFill><a:srgbClr val="{line_color}"/></a:solidFill></a:ln></c:spPr>
  <c:crossAx val="{cross_id}"/>
  <c:crossesAt val="0"/>
</{tag}>"""


def build_scatter_series_xml(idx, label, color_hex, x_values, y_values):
    x_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(x_values)
    )
    y_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(y_values)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr><a:noFill/>
    <a:ln><a:noFill/></a:ln>
  </c:spPr>
  <c:marker>
    <c:symbol val="circle"/>
    <c:size val="5"/>
    <c:spPr>
      <a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill>
      <a:ln><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill></a:ln>
    </c:spPr>
  </c:marker>
  <c:xVal><c:numRef><c:f>x{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(x_values)}"/>{x_pts}</c:numCache>
  </c:numRef></c:xVal>
  <c:yVal><c:numRef><c:f>y{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(y_values)}"/>{y_pts}</c:numCache>
  </c:numRef></c:yVal>
</c:ser>"""


def build_scatter_chart_xml(series_list, legend_pos="r"):
    series_xml = ""
    colors = ["4472C4", "ED7D31", "A5A5A5", "FFC000"]
    for i, (label, x_vals, y_vals) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_scatter_series_xml(i, label, color, x_vals, y_vals)

    val_ax_x = build_axis_xml(100, 200, "b", axis_type="val", gridlines=False)
    val_ax_y = build_axis_xml(200, 100, "l", axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:scatterChart>
        <c:scatterStyle val="lineMarker"/>
        {series_xml}
        <c:axId val="100"/><c:axId val="200"/>
      </c:scatterChart>
      {val_ax_x}
      {val_ax_y}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_doughnut_chart_xml(labels, values, hole_size=50, legend_pos="r"):
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{l}</c:v></c:pt>' for i, l in enumerate(labels)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:doughnutChart>
        <c:varyColors val="1"/>
        <c:ser>
          <c:idx val="0"/><c:order val="0"/>
          <c:tx><c:strRef><c:f>label</c:f>
            <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>Budget</c:v></c:pt></c:strCache>
          </c:strRef></c:tx>
          <c:cat><c:strRef><c:f>cats</c:f>
            <c:strCache><c:ptCount val="{len(labels)}"/>{cat_pts}</c:strCache>
          </c:strRef></c:cat>
          <c:val><c:numRef><c:f>0</c:f>
            <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
          </c:numRef></c:val>
        </c:ser>
        <c:holeSize val="{hole_size}"/>
      </c:doughnutChart>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_radar_series_xml(idx, label, color_hex, categories, values):
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{c}</c:v></c:pt>' for i, c in enumerate(categories)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr><a:noFill/>
    <a:ln w="25400"><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill></a:ln>
  </c:spPr>
  <c:cat><c:strRef><c:f>cats</c:f>
    <c:strCache><c:ptCount val="{len(categories)}"/>{cat_pts}</c:strCache>
  </c:strRef></c:cat>
  <c:val><c:numRef><c:f>{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
  </c:numRef></c:val>
</c:ser>"""


def build_radar_chart_xml(series_list, categories, legend_pos="b", radar_style="marker"):
    colors = ["4F81BD", "C0504D", "9BBB59"]
    series_xml = ""
    for i, (label, values) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_radar_series_xml(i, label, color, categories, values)

    cat_ax = build_axis_xml(100, 200, "b", axis_type="cat")
    val_ax = build_axis_xml(200, 100, "l", axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:radarChart>
        <c:radarStyle val="{radar_style}"/>
        {series_xml}
        <c:axId val="100"/><c:axId val="200"/>
      </c:radarChart>
      {cat_ax}
      {val_ax}
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_bubble_series_xml(idx, label, color_hex, x_values, y_values, sizes):
    x_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(x_values)
    )
    y_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(y_values)
    )
    s_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(sizes)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr>
    <a:solidFill><a:srgbClr val="{color_hex}"><a:alpha val="70000"/></a:srgbClr></a:solidFill>
    <a:ln><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill></a:ln>
  </c:spPr>
  <c:xVal><c:numRef><c:f>x{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(x_values)}"/>{x_pts}</c:numCache>
  </c:numRef></c:xVal>
  <c:yVal><c:numRef><c:f>y{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(y_values)}"/>{y_pts}</c:numCache>
  </c:numRef></c:yVal>
  <c:bubbleSize><c:numRef><c:f>s{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(sizes)}"/>{s_pts}</c:numCache>
  </c:numRef></c:bubbleSize>
</c:ser>"""


def build_bubble_chart_xml(series_list, legend_pos="r"):
    colors = ["4472C4", "ED7D31", "A5A5A5", "FFC000"]
    series_xml = ""
    for i, (label, x_vals, y_vals, sizes) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_bubble_series_xml(i, label, color, x_vals, y_vals, sizes)

    val_ax_x = build_axis_xml(100, 200, "b", axis_type="val", gridlines=False)
    val_ax_y = build_axis_xml(200, 100, "l", axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:bubbleChart>
        <c:varyColors val="0"/>
        {series_xml}
        <c:axId val="100"/><c:axId val="200"/>
      </c:bubbleChart>
      {val_ax_x}
      {val_ax_y}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_drawing_xml(rel_id, cx_emu, cy_emu):
    return (
        f'<w:drawing xmlns:w="{W_NS}" xmlns:wp="{WP_NS}" '
        f'xmlns:a="{DML_NS}" xmlns:c="{CHART_NS}" xmlns:r="{REL_NS}">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx_emu}" cy="{cy_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{rel_id.replace("rId","")}" name="Chart {rel_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic><a:graphicData uri="{CHART_NS}">'
        f'<c:chart r:id="{rel_id}"/>'
        f'</a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing>'
    )


# ── Step 1: Build document ──

doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

doc.add_heading("Scatter, Doughnut, Radar, and Bubble Charts", level=1)

doc.add_paragraph("Chart 1: Scatter chart showing test scores vs study hours")
p1 = doc.add_paragraph()
p1.add_run("CHART_PLACEHOLDER_1")

doc.add_paragraph("")

doc.add_paragraph("Chart 2: Doughnut chart showing budget allocation")
p2 = doc.add_paragraph()
p2.add_run("CHART_PLACEHOLDER_2")

doc.add_paragraph("")

doc.add_paragraph("Chart 3: Radar chart comparing product features")
p3 = doc.add_paragraph()
p3.add_run("CHART_PLACEHOLDER_3")

doc.add_paragraph("")

doc.add_paragraph("Chart 4: Bubble chart showing sales by region")
p4 = doc.add_paragraph()
p4.add_run("CHART_PLACEHOLDER_4")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# ── Step 2: Inject charts ──

charts = {
    1: build_scatter_chart_xml(
        series_list=[
            ("Class A", [2, 4, 5, 7, 8, 9], [55, 65, 72, 80, 88, 92]),
            ("Class B", [1, 3, 5, 6, 8, 10], [40, 58, 68, 75, 82, 95]),
        ],
        legend_pos="r",
    ),
    2: build_doughnut_chart_xml(
        labels=["Marketing", "Engineering", "Sales", "Operations", "HR"],
        values=[30, 35, 20, 10, 5],
        hole_size=50,
        legend_pos="r",
    ),
    3: build_radar_chart_xml(
        categories=["Speed", "Reliability", "Comfort", "Safety", "Price"],
        series_list=[
            ("Product A", [8, 7, 9, 6, 5]),
            ("Product B", [6, 9, 5, 8, 7]),
            ("Product C", [7, 6, 7, 9, 8]),
        ],
        legend_pos="b",
        radar_style="marker",
    ),
    4: build_bubble_chart_xml(
        series_list=[
            ("North", [10, 20, 30, 40], [50, 80, 65, 90], [15, 25, 20, 30]),
            ("South", [15, 25, 35, 45], [45, 70, 85, 60], [20, 10, 30, 25]),
        ],
        legend_pos="r",
    ),
}

chart_sizes = {
    1: (4572000, 2743200),  # ~5.0 x 3.0 inches
    2: (4572000, 2743200),  # ~5.0 x 3.0 inches
    3: (4572000, 2743200),  # ~5.0 x 3.0 inches
    4: (5486400, 2743200),  # ~6.0 x 3.0 inches
}

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        rels_xml = zin.read("word/_rels/document.xml.rels").decode()
        ct_xml = zin.read("[Content_Types].xml").decode()

        existing_rids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_xml)]
        next_rid = max(existing_rids, default=0) + 1

        chart_rids = {}
        for chart_num in sorted(charts.keys()):
            rid = f"rId{next_rid}"
            chart_rids[chart_num] = rid
            next_rid += 1

        new_rels = ""
        for chart_num, rid in chart_rids.items():
            new_rels += (
                f'<Relationship Id="{rid}" '
                f'Type="{CHART_REL_TYPE}" '
                f'Target="charts/chart{chart_num}.xml"/>'
            )
        rels_xml = rels_xml.replace("</Relationships>", new_rels + "</Relationships>")

        new_ct = ""
        for chart_num in charts:
            new_ct += (
                f'<Override PartName="/word/charts/chart{chart_num}.xml" '
                f'ContentType="{CT_CHART}"/>'
            )
        ct_xml = ct_xml.replace("</Types>", new_ct + "</Types>")

        doc_xml = zin.read("word/document.xml").decode()
        for chart_num, rid in chart_rids.items():
            placeholder = f"CHART_PLACEHOLDER_{chart_num}"
            cx, cy = chart_sizes[chart_num]
            drawing = build_drawing_xml(rid, cx, cy)
            run_pattern = f'<w:r><w:rPr></w:rPr><w:t>{placeholder}</w:t></w:r>'
            run_replacement = f'<w:r>{drawing}</w:r>'
            if run_pattern in doc_xml:
                doc_xml = doc_xml.replace(run_pattern, run_replacement)
            else:
                run_pattern2 = f'<w:r><w:t>{placeholder}</w:t></w:r>'
                if run_pattern2 in doc_xml:
                    doc_xml = doc_xml.replace(run_pattern2, run_replacement)
                else:
                    doc_xml = re.sub(
                        rf'<w:r[^>]*>.*?<w:t[^>]*>{placeholder}</w:t>.*?</w:r>',
                        run_replacement,
                        doc_xml,
                        flags=re.DOTALL,
                    )

        for item in zin.infolist():
            if item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, ct_xml)
            elif item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

        for chart_num, chart_xml in charts.items():
            zout.writestr(f"word/charts/chart{chart_num}.xml", chart_xml)

os.unlink(tmp)
print(f"Generated {OUT}")
