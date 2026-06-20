#!/usr/bin/env python3
"""case71: w:isLgl (§17.9.4) — "legal numbering" forces referenced levels to decimal.

isLgl on a w:lvl makes every %N reference in that level's lvlText render as a
decimal numeral regardless of the referenced level's own numFmt. The corpus only
ever carries isLgl on all-decimal multilevel templates, where it is a no-op, so
the feature is genuinely unexercised. This fixture isolates the visible case:

  ilvl 0: upperRoman, lvlText "%1."        -> "I.", "II."
  ilvl 1: decimal, <w:isLgl/>, "%1.%2."    -> "1.1.", "1.2.", "2.1."

With isLgl the %1 (an upperRoman level) is forced to decimal, so the sub-items
read 1.1 / 1.2 / 2.1. A renderer ignoring isLgl would print I.1 / I.2 / II.1.
The contrast between the level-0 "I."/"II." and the legal "1.x" prefixes proves
the override applied.
"""

import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case71/input.docx")

NUM_ID = 100
ABSTRACT_ID = 100

# isLgl level (ilvl 1) is decimal but references upperRoman level 0 via %1 ->
# the only configuration where isLgl changes rendered glyphs.
ABSTRACT_XML = f"""<w:abstractNum w:abstractNumId="{ABSTRACT_ID}">
  <w:multiLevelType w:val="multilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="upperRoman"/>
    <w:lvlText w:val="%1."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="432"/></w:pPr>
  </w:lvl>
  <w:lvl w:ilvl="1">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:isLgl/>
    <w:lvlText w:val="%1.%2."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="1440" w:hanging="432"/></w:pPr>
  </w:lvl>
</w:abstractNum>"""

NUM_XML = f'<w:num w:numId="{NUM_ID}"><w:abstractNumId w:val="{ABSTRACT_ID}"/></w:num>'


def set_num(paragraph, ilvl):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(NUM_ID))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def main():
    doc = Document()

    title = doc.add_paragraph()
    r = title.add_run("isLgl legal numbering")
    r.bold = True
    r.font.size = Pt(14)

    # (ilvl, text) — expect: I. / 1.1 / 1.2 / II. / 2.1
    items = [
        (0, "First section"),
        (1, "Sub one"),
        (1, "Sub two"),
        (0, "Second section"),
        (1, "Sub one again"),
    ]
    for ilvl, text in items:
        p = doc.add_paragraph(text)
        set_num(p, ilvl)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    num_xml = files["word/numbering.xml"].decode("utf-8")
    # CT_Numbering child order is abstractNum* then num*: drop the abstractNum in
    # with the other abstractNums, the num just before the closing tag.
    assert "<w:num " in num_xml, "default template should ship at least one w:num"
    num_xml = num_xml.replace("<w:num ", ABSTRACT_XML + "<w:num ", 1)
    num_xml = num_xml.replace("</w:numbering>", NUM_XML + "</w:numbering>", 1)
    files["word/numbering.xml"] = num_xml.encode("utf-8")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
