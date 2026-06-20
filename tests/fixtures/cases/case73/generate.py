#!/usr/bin/env python3
"""case73: w:pStyle in a numbering level (§17.9.23) — style->level reverse lookup.

A <w:lvl> may carry <w:pStyle w:val="StyleId"/>. When a paragraph style links to
a numbering definition via its own numPr WITHOUT an explicit ilvl, the level is
chosen by matching the paragraph's style against the abstractNum level whose
pStyle names it (the standard multilevel heading-numbering mechanism). A renderer
that ignores pStyle defaults the absent ilvl to 0, collapsing every such heading
to level 0's format.

The corpus only ever carries this in the latent form: styles mapped to level 0
(where default-0 is already correct) or deeper styles that ALSO write an explicit
ilvl (making pStyle redundant). No fixture maps a USED style to a level >0 with no
explicit ilvl, so the feature is genuinely unexercised. This fixture isolates it:

  lvl0  upperRoman  "%1."   pStyle=PHeadingA   -> "I.",  "II."
  lvl1  lowerLetter "%2."   pStyle=PHeadingB   -> "a.",  "b."
  lvl2  lowerRoman  "%3)"   pStyle=PHeadingC   -> "i)"

The styles reference numId 100 with NO ilvl, so only pStyle distinguishes them.
Correct (Word) output reads  I. / a. / b. / II. / a. / i)  — a mix of Roman,
letters and paren-roman. A renderer ignoring pStyle defaults all three to level 0
and prints one upperRoman sequence  I. II. III. IV. V. VI.  — unmistakably wrong.
"""

import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case73/input.docx")

NUM_ID = 100
ABSTRACT_ID = 100

# Each level: distinct numFmt + its own counter only, plus a pStyle naming the
# heading style that maps to it. Distinct formats make a wrong level obvious.
ABSTRACT_XML = f"""<w:abstractNum w:abstractNumId="{ABSTRACT_ID}">
  <w:multiLevelType w:val="multilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="upperRoman"/>
    <w:pStyle w:val="PHeadingA"/>
    <w:lvlText w:val="%1."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="432"/></w:pPr>
  </w:lvl>
  <w:lvl w:ilvl="1">
    <w:start w:val="1"/>
    <w:numFmt w:val="lowerLetter"/>
    <w:pStyle w:val="PHeadingB"/>
    <w:lvlText w:val="%2."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="1440" w:hanging="432"/></w:pPr>
  </w:lvl>
  <w:lvl w:ilvl="2">
    <w:start w:val="1"/>
    <w:numFmt w:val="lowerRoman"/>
    <w:pStyle w:val="PHeadingC"/>
    <w:lvlText w:val="%3)"/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="2160" w:hanging="432"/></w:pPr>
  </w:lvl>
</w:abstractNum>"""

NUM_XML = f'<w:num w:numId="{NUM_ID}"><w:abstractNumId w:val="{ABSTRACT_ID}"/></w:num>'


def heading_style(style_id, name, size):
    # numPr carries numId but NO ilvl: the level is resolvable ONLY via the
    # abstractNum level's pStyle that names this styleId.
    return (
        f'<w:style w:type="paragraph" w:styleId="{style_id}">'
        f'<w:name w:val="{name}"/>'
        f'<w:basedOn w:val="Normal"/>'
        f'<w:pPr><w:numPr><w:numId w:val="{NUM_ID}"/></w:numPr></w:pPr>'
        f'<w:rPr><w:b/><w:sz w:val="{size}"/></w:rPr>'
        f"</w:style>"
    )


STYLES_XML = (
    heading_style("PHeadingA", "P Heading A", 30)
    + heading_style("PHeadingB", "P Heading B", 26)
    + heading_style("PHeadingC", "P Heading C", 22)
)


def set_pstyle(paragraph, style_id):
    # Set pStyle directly: the custom styles are injected post-save, so python-docx
    # doesn't know them and paragraph.style would raise.
    pPr = paragraph._p.get_or_add_pPr()
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.insert(0, pStyle)


def main():
    doc = Document()

    title = doc.add_paragraph()
    r = title.add_run("pStyle numbering levels")
    r.bold = True
    r.font.size = Pt(14)

    # (styleId, text) — expect: I. / a. / b. / II. / a. / i)
    items = [
        ("PHeadingA", "Alpha section"),
        ("PHeadingB", "First subpoint"),
        ("PHeadingB", "Second subpoint"),
        ("PHeadingA", "Beta section"),
        ("PHeadingB", "Another subpoint"),
        ("PHeadingC", "Deep detail"),
    ]
    for style_id, text in items:
        p = doc.add_paragraph(text)
        set_pstyle(p, style_id)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    num_xml = files["word/numbering.xml"].decode("utf-8")
    # CT_Numbering child order is abstractNum* then num*.
    assert "<w:num " in num_xml, "default template should ship at least one w:num"
    num_xml = num_xml.replace("<w:num ", ABSTRACT_XML + "<w:num ", 1)
    num_xml = num_xml.replace("</w:numbering>", NUM_XML + "</w:numbering>", 1)
    files["word/numbering.xml"] = num_xml.encode("utf-8")

    styles_xml = files["word/styles.xml"].decode("utf-8")
    # style* is the trailing sequence in CT_Styles: append before the close tag.
    assert "</w:styles>" in styles_xml
    styles_xml = styles_xml.replace("</w:styles>", STYLES_XML + "</w:styles>", 1)
    files["word/styles.xml"] = styles_xml.encode("utf-8")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
