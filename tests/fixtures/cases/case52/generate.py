#!/usr/bin/env python3
"""case52: Stacked bar chart rendering.

Tests stacked and percent-stacked bar charts, which are parsed but currently
render as clustered. This fixture provides visual baselines for when stacked
rendering is implemented.

Charts:
1. Vertical stacked bar (3 series, 4 categories)
2. Horizontal stacked bar (3 series, 5 categories)
3. Vertical percent-stacked bar (3 series, 4 categories)
4. Horizontal percent-stacked bar (2 series, 6 categories)

Usage:
    uv run tests/fixtures/cases/case52/generate.py
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

OUT = Path("tests/fixtures/cases/case52/input.docx")

CHART_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
DML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CHART_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
CT_CHART = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"


def build_series_xml(idx, label, color_hex, categories, values):
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{c}</c:v></c:pt>' for i, c in enumerate(categories)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill><a:ln><a:noFill/></a:ln></c:spPr>
  <c:cat><c:strRef><c:f>cats</c:f>
    <c:strCache><c:ptCount val="{len(categories)}"/>{cat_pts}</c:strCache>
  </c:strRef></c:cat>
  <c:val><c:numRef><c:f>{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
  </c:numRef></c:val>
</c:ser>"""


def build_axis_xml(ax_id, cross_id, position, axis_type="cat", gridlines=False):
    grid_xml = ""
    if gridlines:
        grid_xml = """<c:majorGridlines><c:spPr><a:ln>
      <a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill>
    </a:ln></c:spPr></c:majorGridlines>"""
    tag = "c:valAx" if axis_type == "val" else "c:catAx"
    return f"""<{tag}>
  <c:axId val="{ax_id}"/>
  <c:scaling><c:orientation val="minMax"/></c:scaling>
  <c:delete val="0"/>
  <c:axPos val="{position}"/>
  {grid_xml}
  <c:majorTickMark val="out"/><c:minorTickMark val="none"/>
  <c:tickLblPos val="nextTo"/>
  <c:spPr><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
  <c:crossAx val="{cross_id}"/>
  <c:crossesAt val="0"/>
</{tag}>"""


def build_chart_xml(bar_dir, grouping, series_list, categories, gap_width,
                    legend_pos=None, colors=None):
    default_colors = ["4472C4", "ED7D31", "A5A5A5", "FFC000", "5B9BD5", "70AD47"]
    if colors is None:
        colors = default_colors

    series_xml = ""
    for i, (label, values) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_series_xml(i, label, color, categories, values)

    cat_ax_id, val_ax_id = 100, 200
    if bar_dir == "bar":
        cat_pos, val_pos = "l", "b"
    else:
        cat_pos, val_pos = "b", "l"

    cat_ax = build_axis_xml(cat_ax_id, val_ax_id, cat_pos, axis_type="cat")
    val_ax = build_axis_xml(val_ax_id, cat_ax_id, val_pos, axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    # Use overlap=100 for stacked charts (bars stack, no gap between series)
    overlap_xml = '<c:overlap val="100"/>' if "stacked" in grouping.lower() else ""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:barChart>
        <c:barDir val="{bar_dir}"/>
        <c:grouping val="{grouping}"/>
        {series_xml}
        <c:gapWidth val="{gap_width}"/>
        {overlap_xml}
        <c:axId val="{cat_ax_id}"/><c:axId val="{val_ax_id}"/>
      </c:barChart>
      {cat_ax}
      {val_ax}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_drawing_xml(rel_id, cx_emu, cy_emu):
    return (
        f'<w:drawing xmlns:w="{W_NS}" xmlns:wp="{WP_NS}" '
        f'xmlns:a="{DML_NS}" xmlns:c="{CHART_NS}" xmlns:r="{REL_NS}">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx_emu}" cy="{cy_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{rel_id.replace("rId","")}" name="Chart {rel_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic><a:graphicData uri="{CHART_NS}">'
        f'<c:chart r:id="{rel_id}"/>'
        f'</a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing>'
    )


def generate():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.add_heading("Stacked Bar Chart Tests", level=1)

    doc.add_paragraph("Chart 1: Vertical stacked bar (3 series, 4 categories)")
    p1 = doc.add_paragraph()
    p1.add_run("CHART_PLACEHOLDER_1")

    doc.add_paragraph("")

    doc.add_paragraph("Chart 2: Horizontal stacked bar (3 series, 5 categories)")
    p2 = doc.add_paragraph()
    p2.add_run("CHART_PLACEHOLDER_2")

    doc.add_paragraph("")

    doc.add_paragraph("Chart 3: Vertical percent-stacked bar (3 series, 4 categories)")
    p3 = doc.add_paragraph()
    p3.add_run("CHART_PLACEHOLDER_3")

    doc.add_paragraph("")

    doc.add_paragraph("Chart 4: Horizontal percent-stacked bar (2 series, 6 categories)")
    p4 = doc.add_paragraph()
    p4.add_run("CHART_PLACEHOLDER_4")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    # Chart data
    charts = {
        1: build_chart_xml(
            bar_dir="col", grouping="stacked", gap_width=100,
            legend_pos="r",
            categories=["Q1", "Q2", "Q3", "Q4"],
            series_list=[
                ("Hardware", [120, 145, 132, 168]),
                ("Software", [85, 92, 110, 95]),
                ("Services", [40, 55, 48, 72]),
            ],
        ),
        2: build_chart_xml(
            bar_dir="bar", grouping="stacked", gap_width=80,
            legend_pos="b",
            categories=["Engineering", "Marketing", "Sales", "Support", "Admin"],
            series_list=[
                ("Salary", [450, 220, 380, 150, 120]),
                ("Benefits", [135, 66, 114, 45, 36]),
                ("Training", [45, 22, 38, 15, 12]),
            ],
            colors=["2E75B6", "BF504D", "548235"],
        ),
        3: build_chart_xml(
            bar_dir="col", grouping="percentStacked", gap_width=100,
            legend_pos="r",
            categories=["North", "South", "East", "West"],
            series_list=[
                ("Product A", [30, 45, 25, 50]),
                ("Product B", [50, 30, 55, 25]),
                ("Product C", [20, 25, 20, 25]),
            ],
            colors=["4472C4", "ED7D31", "70AD47"],
        ),
        4: build_chart_xml(
            bar_dir="bar", grouping="percentStacked", gap_width=80,
            legend_pos="r",
            categories=["Jan", "Feb", "Mar", "Apr", "May", "Jun"],
            series_list=[
                ("Online", [60, 65, 58, 72, 68, 75]),
                ("In-Store", [40, 35, 42, 28, 32, 25]),
            ],
            colors=["5B9BD5", "FFC000"],
        ),
    }

    chart_sizes = {
        1: (4572000, 2743200),
        2: (5486400, 2743200),
        3: (4572000, 2743200),
        4: (5486400, 2743200),
    }

    with zipfile.ZipFile(tmp, "r") as zin:
        with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
            rels_xml = zin.read("word/_rels/document.xml.rels").decode()
            ct_xml = zin.read("[Content_Types].xml").decode()

            existing_rids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_xml)]
            next_rid = max(existing_rids, default=0) + 1

            chart_rids = {}
            for chart_num in sorted(charts.keys()):
                chart_rids[chart_num] = f"rId{next_rid}"
                next_rid += 1

            # Patch rels
            new_rels = ""
            for chart_num, rid in chart_rids.items():
                new_rels += (
                    f'<Relationship Id="{rid}" '
                    f'Type="{CHART_REL_TYPE}" '
                    f'Target="charts/chart{chart_num}.xml"/>'
                )
            rels_xml = rels_xml.replace("</Relationships>", new_rels + "</Relationships>")

            # Patch content types
            new_ct = ""
            for chart_num in charts:
                new_ct += (
                    f'<Override PartName="/word/charts/chart{chart_num}.xml" '
                    f'ContentType="{CT_CHART}"/>'
                )
            ct_xml = ct_xml.replace("</Types>", new_ct + "</Types>")

            # Replace placeholders with chart drawings
            doc_xml = zin.read("word/document.xml").decode()
            for chart_num, rid in chart_rids.items():
                placeholder = f"CHART_PLACEHOLDER_{chart_num}"
                cx, cy = chart_sizes[chart_num]
                drawing = build_drawing_xml(rid, cx, cy)
                run_replacement = f'<w:r>{drawing}</w:r>'
                # Try multiple patterns
                for pattern in [
                    f'<w:r><w:rPr></w:rPr><w:t>{placeholder}</w:t></w:r>',
                    f'<w:r><w:t>{placeholder}</w:t></w:r>',
                ]:
                    if pattern in doc_xml:
                        doc_xml = doc_xml.replace(pattern, run_replacement)
                        break
                else:
                    doc_xml = re.sub(
                        rf'<w:r[^>]*>.*?<w:t[^>]*>{placeholder}</w:t>.*?</w:r>',
                        run_replacement, doc_xml, flags=re.DOTALL,
                    )

            for item in zin.infolist():
                if item.filename == "word/_rels/document.xml.rels":
                    zout.writestr(item, rels_xml)
                elif item.filename == "[Content_Types].xml":
                    zout.writestr(item, ct_xml)
                elif item.filename == "word/document.xml":
                    zout.writestr(item, doc_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))

            for chart_num, chart_xml in charts.items():
                zout.writestr(f"word/charts/chart{chart_num}.xml", chart_xml)

    os.unlink(tmp)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    generate()
