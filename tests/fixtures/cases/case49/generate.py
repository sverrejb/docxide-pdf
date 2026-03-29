#!/usr/bin/env python3
"""case49: Table of Contents with right-aligned tabs, dot leaders, and PAGEREF fields.

Tests realistic TOC rendering:
- TOC Heading ("Table of Contents")
- TOC 1/2/3 styles with right-aligned tab stops and dot leaders
- Hyperlink anchors for each TOC entry
- PAGEREF field codes for page numbers
- Tab character between heading text and page number
- Bookmark targets on actual headings
- Enough body text to push content onto multiple pages

Usage:
    uv run tests/fixtures/cases/case49/generate.py
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("tests/fixtures/cases/case49/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

BODY_TEXT = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod "
    "tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, "
    "quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. "
    "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu "
    "fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
    "culpa qui officia deserunt mollit anim id est laborum."
)

BODY_TEXT_2 = (
    "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium "
    "doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore "
    "veritatis et quasi architecto beatae vitae dicta sunt explicabo. Nemo enim "
    "ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, sed quia "
    "consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt."
)

# Document structure: TOC + 3 chapters with subsections
CHAPTERS = [
    {
        "title": "Introduction",
        "bookmark": "_Toc100",
        "level": 1,
        "subsections": [
            {"title": "Background", "bookmark": "_Toc101", "level": 2},
            {"title": "Scope and Objectives", "bookmark": "_Toc102", "level": 2},
        ],
    },
    {
        "title": "Literature Review",
        "bookmark": "_Toc200",
        "level": 1,
        "subsections": [
            {"title": "Historical Context", "bookmark": "_Toc201", "level": 2},
            {"title": "Current Research", "bookmark": "_Toc202", "level": 2},
            {"title": "Methodology Overview", "bookmark": "_Toc203", "level": 2,
             "subsubsections": [
                 {"title": "Qualitative Methods", "bookmark": "_Toc204", "level": 3},
                 {"title": "Quantitative Methods", "bookmark": "_Toc205", "level": 3},
             ]},
        ],
    },
    {
        "title": "Conclusion",
        "bookmark": "_Toc300",
        "level": 1,
        "subsections": [
            {"title": "Summary of Findings", "bookmark": "_Toc301", "level": 2},
            {"title": "Future Work", "bookmark": "_Toc302", "level": 2},
        ],
    },
]


def collect_all_entries(chapters):
    """Flatten chapters into a list of (title, bookmark, level) tuples."""
    entries = []
    for ch in chapters:
        entries.append((ch["title"], ch["bookmark"], ch["level"]))
        for sub in ch.get("subsections", []):
            entries.append((sub["title"], sub["bookmark"], sub["level"]))
            for subsub in sub.get("subsubsections", []):
                entries.append((subsub["title"], subsub["bookmark"], subsub["level"]))
    return entries


def toc_entry_xml(display_text, anchor_name, level, bookmark_id, page_num="1"):
    """TOC paragraph with hyperlink, tab, dot leader, and PAGEREF field."""
    style = f"TOC{level}"
    # Indent: TOC1=0, TOC2=240twips (12pt), TOC3=480twips (24pt)
    indent = (level - 1) * 240

    indent_xml = f'<w:ind w:left="{indent}"/>' if indent > 0 else ""

    # Right-aligned tab at 9350 twips (6.5 inches from left = right margin with 1" margins)
    return (
        f'<w:p>'
        f'<w:pPr>'
        f'<w:pStyle w:val="{style}"/>'
        f'<w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9350"/></w:tabs>'
        f'{indent_xml}'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'</w:pPr>'
        f'<w:hyperlink w:anchor="{anchor_name}">'
        # Heading text run
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{display_text}</w:t>'
        f'</w:r>'
        # Tab run
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:webHidden/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:tab/>'
        f'</w:r>'
        # PAGEREF field begin
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:webHidden/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:webHidden/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:instrText xml:space="preserve"> PAGEREF {anchor_name} \\h </w:instrText>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:webHidden/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
        # Page number text (placeholder — Word replaces this)
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:webHidden/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{page_num}</w:t>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:noProof/>'
        f'<w:webHidden/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
        f'</w:hyperlink>'
        f'</w:p>'
    )


def heading_xml(text, level, bookmark_name, bookmark_id):
    """Heading paragraph with bookmark."""
    style = f"Heading{level}"
    return (
        f'<w:p>'
        f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>'
        f'<w:bookmarkStart w:id="{bookmark_id}" w:name="{bookmark_name}"/>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'</w:rPr>'
        f'<w:t>{text}</w:t>'
        f'</w:r>'
        f'<w:bookmarkEnd w:id="{bookmark_id}"/>'
        f'</w:p>'
    )


def body_para_xml(text):
    """Normal body paragraph."""
    return (
        f'<w:p>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )


def generate():
    # Step 1: Create base document with placeholder markers
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Placeholder paragraphs — will be replaced with XML
    doc.add_paragraph("TOC_HEADING_PLACEHOLDER")
    doc.add_paragraph("TOC_ENTRIES_PLACEHOLDER")
    doc.add_paragraph("BODY_CONTENT_PLACEHOLDER")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    # Step 2: Build the XML for TOC entries and body content
    all_entries = collect_all_entries(CHAPTERS)

    # TOC heading
    toc_heading = (
        '<w:p>'
        '<w:pPr><w:pStyle w:val="TOCHeading"/></w:pPr>'
        '<w:r>'
        '<w:rPr>'
        '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        '</w:rPr>'
        '<w:t>Table of Contents</w:t>'
        '</w:r>'
        '</w:p>'
    )

    # TOC entries
    toc_xml_parts = []
    for i, (title, bookmark, level) in enumerate(all_entries):
        toc_xml_parts.append(toc_entry_xml(title, bookmark, level, i + 10))

    # Body content with headings and text
    body_parts = []
    bookmark_id = 100
    for ch in CHAPTERS:
        body_parts.append(heading_xml(ch["title"], 1, ch["bookmark"], bookmark_id))
        bookmark_id += 1
        body_parts.append(body_para_xml(BODY_TEXT))
        body_parts.append(body_para_xml(BODY_TEXT_2))

        for sub in ch.get("subsections", []):
            body_parts.append(heading_xml(sub["title"], 2, sub["bookmark"], bookmark_id))
            bookmark_id += 1
            body_parts.append(body_para_xml(BODY_TEXT))

            for subsub in sub.get("subsubsections", []):
                body_parts.append(
                    heading_xml(subsub["title"], 3, subsub["bookmark"], bookmark_id)
                )
                bookmark_id += 1
                body_parts.append(body_para_xml(BODY_TEXT_2))

    # Step 3: Post-process ZIP
    with zipfile.ZipFile(tmp, "r") as zin:
        with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
            doc_xml = zin.read("word/document.xml").decode()
            styles_xml = zin.read("word/styles.xml").decode()

            # Replace placeholders using lambda to avoid regex escape issues in replacement
            for placeholder, replacement in [
                ("TOC_HEADING_PLACEHOLDER", toc_heading),
                ("TOC_ENTRIES_PLACEHOLDER", "\n".join(toc_xml_parts)),
                ("BODY_CONTENT_PLACEHOLDER", "\n".join(body_parts)),
            ]:
                doc_xml = re.sub(
                    r'<w:p[^>]*><w:r><w:t>' + placeholder + r'</w:t></w:r></w:p>',
                    lambda m, r=replacement: r,
                    doc_xml,
                    count=1,
                )

            # Add TOC styles to styles.xml if not present
            toc_styles = """
  <w:style w:type="paragraph" w:styleId="TOCHeading">
    <w:name w:val="TOC Heading"/>
    <w:basedOn w:val="Heading1"/>
    <w:next w:val="Normal"/>
    <w:uiPriority w:val="39"/>
    <w:qFormat/>
    <w:pPr>
      <w:spacing w:before="480" w:after="0"/>
      <w:outlineLvl w:val="9"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:b/>
      <w:sz w:val="28"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="TOC1">
    <w:name w:val="toc 1"/>
    <w:basedOn w:val="Normal"/>
    <w:next w:val="Normal"/>
    <w:uiPriority w:val="39"/>
    <w:pPr>
      <w:tabs>
        <w:tab w:val="right" w:leader="dot" w:pos="9350"/>
      </w:tabs>
      <w:spacing w:after="100"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:b/>
      <w:sz w:val="24"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="TOC2">
    <w:name w:val="toc 2"/>
    <w:basedOn w:val="Normal"/>
    <w:next w:val="Normal"/>
    <w:uiPriority w:val="39"/>
    <w:pPr>
      <w:tabs>
        <w:tab w:val="right" w:leader="dot" w:pos="9350"/>
      </w:tabs>
      <w:spacing w:after="100"/>
      <w:ind w:left="240"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:sz w:val="24"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="TOC3">
    <w:name w:val="toc 3"/>
    <w:basedOn w:val="Normal"/>
    <w:next w:val="Normal"/>
    <w:uiPriority w:val="39"/>
    <w:pPr>
      <w:tabs>
        <w:tab w:val="right" w:leader="dot" w:pos="9350"/>
      </w:tabs>
      <w:spacing w:after="100"/>
      <w:ind w:left="480"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:sz w:val="24"/>
    </w:rPr>
  </w:style>"""

            # Insert TOC styles before closing </w:styles>
            styles_xml = styles_xml.replace("</w:styles>", toc_styles + "\n</w:styles>")

            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, doc_xml)
                elif item.filename == "word/styles.xml":
                    zout.writestr(item, styles_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))

    os.unlink(tmp)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    generate()
