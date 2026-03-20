"""Generate a DOCX with floating images and text wrapping for case41.

Tests:
- Left-aligned floating image with text wrapping to the right (wrapSquare)
- Right-aligned floating image with text wrapping to the left
- Center-aligned floating image (should push text above/below)
- Text flowing back to full width below each image
"""

import io
import os
import re
import struct
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Emu

OUT = Path("tests/fixtures/cases/case41/input.docx")


def make_bmp(width, height, r, g, b):
    """Create a minimal BMP image (no PIL dependency)."""
    row_size = (width * 3 + 3) & ~3  # rows padded to 4-byte boundary
    pixel_data_size = row_size * height
    file_size = 54 + pixel_data_size
    buf = io.BytesIO()
    # BMP header
    buf.write(b"BM")
    buf.write(struct.pack("<I", file_size))
    buf.write(struct.pack("<HH", 0, 0))
    buf.write(struct.pack("<I", 54))
    # DIB header (BITMAPINFOHEADER)
    buf.write(struct.pack("<I", 40))
    buf.write(struct.pack("<i", width))
    buf.write(struct.pack("<i", height))
    buf.write(struct.pack("<HH", 1, 24))
    buf.write(struct.pack("<I", 0))  # no compression
    buf.write(struct.pack("<I", pixel_data_size))
    buf.write(struct.pack("<ii", 2835, 2835))  # 72 DPI
    buf.write(struct.pack("<II", 0, 0))
    # Pixel data (BMP stores rows bottom-to-top, BGR order)
    row = bytes([b, g, r] * width)
    padding = b"\x00" * (row_size - width * 3)
    for _ in range(height):
        buf.write(row)
        buf.write(padding)
    return buf.getvalue()


# Create simple colored test images
img_blue = make_bmp(200, 200, 70, 130, 200)  # ~2.78" at 72dpi
img_green = make_bmp(200, 150, 80, 180, 100)
img_red = make_bmp(350, 180, 200, 80, 80)  # wider, for center test

# Save temp images
tmp_dir = tempfile.mkdtemp()
blue_path = os.path.join(tmp_dir, "blue.bmp")
green_path = os.path.join(tmp_dir, "green.bmp")
red_path = os.path.join(tmp_dir, "red.bmp")
with open(blue_path, "wb") as f:
    f.write(img_blue)
with open(green_path, "wb") as f:
    f.write(img_green)
with open(red_path, "wb") as f:
    f.write(img_red)

doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

LOREM = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
    "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
    "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris "
    "nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in "
    "reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla "
    "pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
    "culpa qui officia deserunt mollit anim id est laborum. "
)

# --- Section 1: Left-aligned floating image ---

p1 = doc.add_paragraph(
    "This paragraph precedes a left-aligned floating image. "
    "The body text below should wrap to the right of the image."
)

# Add inline image (will be converted to floating in post-processing)
p_img1 = doc.add_paragraph()
p_img1.add_run().add_picture(blue_path, width=Inches(2.5), height=Inches(2.5))

p2 = doc.add_paragraph(
    "This paragraph should appear to the right of the blue image. "
    "Word wraps body text around floating images when wrapSquare is set "
    "and there is sufficient horizontal space beside the image. " + LOREM +
    "The text continues to fill the space beside the image and eventually "
    "flows below it at full page width."
)

p3 = doc.add_paragraph(
    "This paragraph is fully below the left-aligned floating image. "
    "Normal full-width layout resumes here."
)

# --- Section 2: Right-aligned floating image ---

doc.add_page_break()

p4 = doc.add_paragraph(
    "This section has a right-aligned floating image. "
    "The body text should wrap to the left of the image."
)

p_img2 = doc.add_paragraph()
p_img2.add_run().add_picture(green_path, width=Inches(2.5), height=Inches(2))

p5 = doc.add_paragraph(
    "This text wraps to the left of the right-aligned green image. "
    "The image is anchored to the right side of the column. "
    "Body text fills the available space on the left. " + LOREM +
    "When the text extends past the image, it returns to full width."
)

p6 = doc.add_paragraph(
    "After the right-aligned image, text returns to full page width."
)

# --- Section 3: Center-aligned floating image ---

doc.add_page_break()

p7 = doc.add_paragraph(
    "This section has a center-aligned floating image that is wide enough "
    "to leave insufficient room for text on either side. Text should flow "
    "above and below the image, not beside it."
)

p_img3 = doc.add_paragraph()
p_img3.add_run().add_picture(red_path, width=Inches(4.5), height=Inches(2.5))

p8 = doc.add_paragraph(
    "This paragraph should appear below the centered red image since there "
    "is not enough horizontal space on either side for text wrapping. " + LOREM
)

# --- Section 4: Left-aligned SMALL image (diagnostic: does size mask the shift?) ---

doc.add_page_break()

p9 = doc.add_paragraph(
    "Page 4: Left-aligned SMALL image (1.5x1 inch). Same structure as page 1 "
    "but smaller image to reveal any vertical shift more clearly."
)

# Reuse green image for small case
p_img4 = doc.add_paragraph()
p_img4.add_run().add_picture(green_path, width=Inches(1.5), height=Inches(1))

p10 = doc.add_paragraph(
    "This text wraps beside the small left-aligned image. " + LOREM +
    "The text should flow back to full width below the image."
)

p11 = doc.add_paragraph("Full-width text below the small image.")

# --- Section 5: Right-aligned TALL image (diagnostic: same height as page 1) ---

doc.add_page_break()

p12 = doc.add_paragraph(
    "Page 5: Right-aligned image same size as page 1 (2.5x2.5 inch). "
    "Compare with page 2 which uses a shorter image."
)

p_img5 = doc.add_paragraph()
p_img5.add_run().add_picture(blue_path, width=Inches(2.5), height=Inches(2.5))

p13 = doc.add_paragraph(
    "This text wraps to the left of the tall right-aligned image. " + LOREM +
    "Compare the image vertical position with page 2."
)

p14 = doc.add_paragraph("Full-width text below the tall right-aligned image.")

# --- Section 6: Image IN the text paragraph (not standalone) ---

doc.add_page_break()

p15 = doc.add_paragraph(
    "Page 6: Image is in the SAME paragraph as text, not in its own paragraph. "
    "This is the control case."
)

# Image placed in a paragraph that also has text
p_img6 = doc.add_paragraph("Text before image in same paragraph. ")
p_img6.add_run().add_picture(green_path, width=Inches(2.5), height=Inches(2))

p16 = doc.add_paragraph(
    "This text wraps beside the image that shares a paragraph with text. " + LOREM
)

p17 = doc.add_paragraph("Full-width text after the shared-paragraph image.")

# --- Section 7: Image-only paragraph at TOP of page (no preceding text) ---

doc.add_page_break()

p_img7 = doc.add_paragraph()
p_img7.add_run().add_picture(blue_path, width=Inches(2.5), height=Inches(2.5))

p18 = doc.add_paragraph(
    "Page 7: The left-aligned image is at the top of this page with no "
    "preceding text paragraph. " + LOREM +
    "This tests whether the image position depends on a preceding paragraph."
)

p19 = doc.add_paragraph("Full-width text below the top-of-page image.")

# Save initial DOCX
tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Post-process: convert inline images to floating anchored images
# python-docx creates wp:inline; we need wp:anchor with positioning

WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def make_anchor(inline_xml, h_align, dist_l_emu, dist_r_emu, dist_t_emu, dist_b_emu, cx, cy):
    """Replace wp:inline with wp:anchor for floating positioning."""
    # Extract the graphic content from inline
    # Find <a:graphic ... </a:graphic>
    graphic_match = re.search(r'(<a:graphic\b.*?</a:graphic>)', inline_xml, re.DOTALL)
    if not graphic_match:
        return inline_xml
    graphic = graphic_match.group(1)

    # Extract docPr
    docpr_match = re.search(r'(<wp:docPr[^/]*/\s*>)', inline_xml)
    docpr = docpr_match.group(1) if docpr_match else '<wp:docPr id="1" name="Picture"/>'

    # Extract effectExtent or use default
    effect_match = re.search(r'(<wp:effectExtent[^/]*/\s*>)', inline_xml)
    effect = effect_match.group(1) if effect_match else '<wp:effectExtent l="0" t="0" r="0" b="0"/>'

    anchor = (
        f'<wp:anchor distT="{dist_t_emu}" distB="{dist_b_emu}" '
        f'distL="{dist_l_emu}" distR="{dist_r_emu}" '
        f'simplePos="0" relativeHeight="0" behindDoc="0" '
        f'locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column">'
        f'<wp:align>{h_align}</wp:align>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph">'
        f'<wp:posOffset>0</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'{effect}'
        f'<wp:wrapSquare wrapText="bothSides"/>'
        f'{docpr}'
        f'<wp:cNvGraphicFramePr/>'
        f'{graphic}'
        f'</wp:anchor>'
    )
    return anchor


with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        # Image 1: left-aligned, 2.5" x 2.5"
        cx1 = int(2.5 * 914400)  # EMU
        cy1 = int(2.5 * 914400)
        gap = int(0.1 * 914400)  # 0.1 inch gap

        # Find first wp:inline and replace
        inline_re = re.compile(r'<wp:inline\b[^>]*>.*?</wp:inline>', re.DOTALL)
        inlines = list(inline_re.finditer(doc_xml))

        # Define all image conversions: (index, h_align, dist_l, dist_r, dist_t, dist_b, cx, cy)
        conversions = [
            # Image 1: left-aligned, 2.5" x 2.5"
            (0, "left", 0, gap, gap, gap, int(2.5 * 914400), int(2.5 * 914400)),
            # Image 2: right-aligned, 2.5" x 2"
            (1, "right", gap, 0, gap, gap, int(2.5 * 914400), int(2.0 * 914400)),
            # Image 3: center-aligned, 4.5" x 2.5"
            (2, "center", gap, gap, gap, gap, int(4.5 * 914400), int(2.5 * 914400)),
            # Image 4: left-aligned SMALL, 1.5" x 1"
            (3, "left", 0, gap, gap, gap, int(1.5 * 914400), int(1.0 * 914400)),
            # Image 5: right-aligned TALL, 2.5" x 2.5"
            (4, "right", gap, 0, gap, gap, int(2.5 * 914400), int(2.5 * 914400)),
            # Image 6: left-aligned (in paragraph WITH text), 2.5" x 2"
            (5, "left", 0, gap, gap, gap, int(2.5 * 914400), int(2.0 * 914400)),
            # Image 7: left-aligned at top of page, 2.5" x 2.5"
            (6, "left", 0, gap, gap, gap, int(2.5 * 914400), int(2.5 * 914400)),
        ]

        # Process in reverse order to preserve string positions
        for idx, h_align, dl, dr, dt, db, cx, cy in reversed(conversions):
            if idx < len(inlines):
                m = inlines[idx]
                anchor = make_anchor(m.group(), h_align, dl, dr, dt, db, cx, cy)
                doc_xml = doc_xml[:m.start()] + anchor + doc_xml[m.end():]

        # Ensure DrawingML namespace prefixes are declared at root level
        # (python-docx declares them locally on wp:inline which we replaced)
        if 'xmlns:a=' not in doc_xml:
            doc_xml = doc_xml.replace(
                '<w:document ',
                '<w:document '
                'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
                'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" ',
                1,
            )

        # Move image drawings into the following paragraph's first run
        # Currently: <w:p><w:r><w:drawing>...</w:drawing></w:r></w:p><w:p>text...</w:p>
        # We need: <w:p><w:r><w:drawing>...</w:drawing></w:r><w:r>text...</w:r></w:p>
        # Actually, floating images can stay in their own paragraph — Word handles
        # both patterns. The anchor positioning is relative to the paragraph.

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)

# Clean up temp images
for p in [blue_path, green_path, red_path]:
    os.unlink(p)
os.rmdir(tmp_dir)

print(f"Generated {OUT}")
