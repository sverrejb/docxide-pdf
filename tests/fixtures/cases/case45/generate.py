#!/usr/bin/env python3
"""case45: Floating table text wrapping.

Tests body text wrapping around positioned (floating) tables:
1. Left-aligned floating table with text wrapping to the right
2. Text continuing below the table at full width
3. Right-aligned floating table with text wrapping to the left
4. Floating table with explicit wrap distances (margins between table and text)
5. Tall floating table spanning multiple body paragraphs
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Twips

OUT = Path("tests/fixtures/cases/case45/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

LOREM = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
    "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
    "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris "
    "nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in "
    "reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla "
    "pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
    "culpa qui officia deserunt mollit anim id est laborum."
)


def main():
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Scenario 1: Left-aligned floating table ---
    p = doc.add_paragraph()
    run = p.add_run("Scenario 1: Left-Aligned Floating Table")
    run.bold = True
    run.font.size = Pt(14)

    # Table 1: small 2x2, will be positioned left with text wrapping right
    t1 = doc.add_table(rows=3, cols=2)
    for ri, row in enumerate([("Name", "Value"), ("Alpha", "100"), ("Beta", "200")]):
        for ci, text in enumerate(row):
            t1.rows[ri].cells[ci].text = text

    # Body text that should wrap to the right of the table
    doc.add_paragraph(LOREM)
    doc.add_paragraph(LOREM)

    # --- Scenario 2: Right-aligned floating table ---
    p = doc.add_paragraph()
    run = p.add_run("Scenario 2: Right-Aligned Floating Table")
    run.bold = True
    run.font.size = Pt(14)

    # Table 2: small 2x2, positioned right
    t2 = doc.add_table(rows=3, cols=2)
    for ri, row in enumerate([("Item", "Cost"), ("Widget", "$50"), ("Gadget", "$75")]):
        for ci, text in enumerate(row):
            t2.rows[ri].cells[ci].text = text

    # Body text that should wrap to the left of the table
    doc.add_paragraph(LOREM)

    # --- Scenario 3: Center floating table with wrap distances ---
    p = doc.add_paragraph()
    run = p.add_run("Scenario 3: Centered Table with Wrap Gaps")
    run.bold = True
    run.font.size = Pt(14)

    t3 = doc.add_table(rows=2, cols=3)
    for ri, row in enumerate([("X", "Y", "Z"), ("10", "20", "30")]):
        for ci, text in enumerate(row):
            t3.rows[ri].cells[ci].text = text

    doc.add_paragraph(
        "This text appears after the centered floating table. "
        "It should not be pushed down by the table height. "
        "Instead it should appear directly below the heading, "
        "with the table floating above it. " + LOREM
    )

    # --- Scenario 4: Tall table spanning multiple paragraphs ---
    p = doc.add_paragraph()
    run = p.add_run("Scenario 4: Tall Table Spanning Multiple Paragraphs")
    run.bold = True
    run.font.size = Pt(14)

    # Tall table with many rows — should span across several body paragraphs
    t4 = doc.add_table(rows=10, cols=2)
    row_data = [
        ("Month", "Sales"),
        ("January", "$1,200"),
        ("February", "$1,450"),
        ("March", "$1,100"),
        ("April", "$1,800"),
        ("May", "$2,100"),
        ("June", "$1,950"),
        ("July", "$2,300"),
        ("August", "$2,050"),
        ("September", "$1,700"),
    ]
    for ri, (c1, c2) in enumerate(row_data):
        t4.rows[ri].cells[0].text = c1
        t4.rows[ri].cells[1].text = c2

    # Multiple paragraphs that should wrap beside the tall table
    doc.add_paragraph(
        "First paragraph beside the tall table. " + LOREM
    )
    doc.add_paragraph(
        "Second paragraph beside the tall table. This paragraph should also "
        "wrap in the narrow column to the right of the floating table. " + LOREM
    )
    doc.add_paragraph(
        "Third paragraph. By now the table may have ended, so this text "
        "should resume full page width if the table is shorter than the "
        "combined text height. " + LOREM
    )
    doc.add_paragraph(
        "Fourth paragraph. This should definitely be at full width below "
        "the floating table. Normal layout resumes here."
    )

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    # --- Post-process: inject tblpPr into each table ---
    BORDERS = (
        '<w:tblBorders>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )

    # Table 1: left-aligned, anchored to text, offset below heading
    TBL1_POS = (
        '<w:tblpPr'
        ' w:vertAnchor="text"'
        ' w:horzAnchor="margin"'
        ' w:tblpX="0"'
        ' w:tblpY="400"'
        ' w:topFromText="72"'
        ' w:bottomFromText="72"'
        ' w:leftFromText="144"'
        ' w:rightFromText="144"/>'
    )

    # Table 2: right-aligned, anchored to margin
    TBL2_POS = (
        '<w:tblpPr'
        ' w:vertAnchor="text"'
        ' w:horzAnchor="margin"'
        ' w:tblpXSpec="right"'
        ' w:tblpY="400"'
        ' w:topFromText="72"'
        ' w:bottomFromText="72"'
        ' w:leftFromText="144"'
        ' w:rightFromText="144"/>'
    )

    # Table 3: centered, with generous wrap distances
    TBL3_POS = (
        '<w:tblpPr'
        ' w:vertAnchor="text"'
        ' w:horzAnchor="margin"'
        ' w:tblpXSpec="center"'
        ' w:tblpY="200"'
        ' w:topFromText="72"'
        ' w:bottomFromText="72"'
        ' w:leftFromText="144"'
        ' w:rightFromText="144"/>'
    )

    # Table 4: left-aligned tall table, spans multiple body paragraphs
    TBL4_POS = (
        '<w:tblpPr'
        ' w:vertAnchor="text"'
        ' w:horzAnchor="margin"'
        ' w:tblpX="0"'
        ' w:tblpY="400"'
        ' w:topFromText="72"'
        ' w:bottomFromText="72"'
        ' w:leftFromText="144"'
        ' w:rightFromText="144"/>'
    )

    # Fixed column widths so tables are predictably sized
    TBL_WIDTH_HALF = '<w:tblW w:w="4000" w:type="dxa"/>'

    positions = [TBL1_POS, TBL2_POS, TBL3_POS, TBL4_POS]

    with zipfile.ZipFile(tmp, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode()
        other_files = {}
        for item in zin.infolist():
            if item.filename != "word/document.xml":
                other_files[item.filename] = zin.read(item.filename)

    # Inject tblpPr, borders, and fixed width into each table's tblPr
    tbl_count = 0
    def inject_tblpr(match):
        nonlocal tbl_count
        if tbl_count < len(positions):
            pos = positions[tbl_count]
            tbl_count += 1
            # Remove existing tblStyle, inject our properties
            inner = match.group(0)
            inner = re.sub(r'<w:tblStyle w:val="[^"]*"/>', '', inner)
            inner = re.sub(r'<w:tblW[^/]*/>', TBL_WIDTH_HALF, inner)
            return inner.replace("</w:tblPr>", pos + BORDERS + "</w:tblPr>")
        return match.group(0)

    doc_xml = re.sub(r'<w:tblPr>.*?</w:tblPr>', inject_tblpr, doc_xml, flags=re.DOTALL)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_files.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    main()
