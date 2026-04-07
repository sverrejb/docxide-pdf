"""Generate case56: Shadow Variations test fixture.

Creates a DOCX with 6 images, each having different outerShdw configurations,
plus one image with no shadow as a baseline.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
import os
import io
import struct
from PIL import Image

def make_test_image(width_px, height_px, color_rgb):
    """Create a simple colored rectangle PNG image."""
    img = Image.new('RGB', (width_px, height_px), color_rgb)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()

def main():
    doc = Document()

    # Page title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("Shadow Variations Test (case56)")
    run.bold = True
    run.font.size = Pt(16)

    # Create test images with different colors so we can tell them apart
    images = [
        ("Standard shadow (23pt blur, 45deg, #333 @65%)", (100, 80), (50, 100, 200)),
        ("Subtle shadow (6pt blur, 90deg down, #000 @40%)", (120, 70), (200, 100, 50)),
        ("Large soft shadow (40pt blur, 45deg, #666 @50%)", (100, 80), (50, 180, 80)),
        ("Sharp shadow (no blur, 45deg, #000 @80%)", (110, 75), (180, 50, 50)),
        ("Blue colored shadow (12pt blur, 45deg, #00F @50%)", (100, 80), (200, 200, 50)),
        ("Close high-alpha shadow (12pt blur, 90deg, #000 @90%)", (100, 80), (150, 100, 200)),
        ("No shadow (baseline)", (100, 80), (180, 180, 180)),
    ]

    for label, (w, h), color in images:
        # Add label
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.size = Pt(10)

        # Add image
        img_data = make_test_image(w, h, color)
        img_stream = io.BytesIO(img_data)
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run().add_picture(img_stream, width=Inches(2.5))

        doc.add_paragraph()  # spacer

    # Save initial docx
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "input.docx")
    doc.save(output_path)

    # Now post-process the ZIP to inject outerShdw into each image's spPr
    shadow_configs = [
        # (blurRad, dist, dir, color, alpha) - EMU values
        (292100, 139700, 2700000, "333333", 65000),   # Standard
        (76200,  38100,  5400000, "000000", 40000),    # Subtle
        (508000, 203200, 2700000, "666666", 50000),    # Large soft
        (0,      63500,  2700000, "000000", 80000),    # Sharp (no blur)
        (152400, 101600, 2700000, "0000FF", 50000),    # Blue colored
        (152400, 25400,  5400000, "000000", 90000),    # Close high-alpha
        None,  # No shadow
    ]

    inject_shadows(output_path, shadow_configs)
    print(f"Generated {output_path}")

def inject_shadows(docx_path, shadow_configs):
    """Post-process DOCX ZIP to add outerShdw to each picture's spPr."""
    import xml.etree.ElementTree as ET

    # Read original ZIP
    with zipfile.ZipFile(docx_path, 'r') as zin:
        entries = {}
        for name in zin.namelist():
            entries[name] = zin.read(name)

    # Parse document.xml
    doc_xml = entries['word/document.xml']
    # Register namespaces to avoid ns0/ns1 prefixes
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    }
    for prefix, uri in namespaces.items():
        ET.register_namespace(prefix, uri)

    tree = ET.parse(io.BytesIO(doc_xml))
    root = tree.getroot()

    # Find all pic:pic elements (pictures)
    pics = root.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')

    for i, pic in enumerate(pics):
        if i >= len(shadow_configs) or shadow_configs[i] is None:
            continue

        blur_rad, dist, direction, color, alpha = shadow_configs[i]

        # Find or create spPr
        spPr = pic.find('{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')
        if spPr is None:
            spPr = ET.SubElement(pic, '{http://schemas.openxmlformats.org/drawingml/2006/picture}spPr')

        # Create effectLst with outerShdw
        effectLst = ET.SubElement(spPr, '{http://schemas.openxmlformats.org/drawingml/2006/main}effectLst')
        outerShdw = ET.SubElement(effectLst, '{http://schemas.openxmlformats.org/drawingml/2006/main}outerShdw')
        outerShdw.set('blurRad', str(blur_rad))
        outerShdw.set('dist', str(dist))
        outerShdw.set('dir', str(direction))
        outerShdw.set('algn', 'tl')
        outerShdw.set('rotWithShape', '0')

        srgbClr = ET.SubElement(outerShdw, '{http://schemas.openxmlformats.org/drawingml/2006/main}srgbClr')
        srgbClr.set('val', color)
        alphaElem = ET.SubElement(srgbClr, '{http://schemas.openxmlformats.org/drawingml/2006/main}alpha')
        alphaElem.set('val', str(alpha))

    # Serialize back
    buf = io.BytesIO()
    tree.write(buf, xml_declaration=True, encoding='UTF-8')
    entries['word/document.xml'] = buf.getvalue()

    # Write new ZIP
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)

if __name__ == '__main__':
    main()
