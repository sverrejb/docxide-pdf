"""Generate case78: picture cropping via a:srcRect.

A 400x300 px PNG with a 4x4 grid of distinct solid colours (black grid lines) is
placed four times. Crop values are in 1/100000 of the source dimension, matching
what Word writes. Each frame keeps the visible region's aspect ratio so a correct
render shows undistorted cells and a wrong one (full image squeezed into the
frame) is obvious.

1. Uncropped control.
2. Top/bottom crop t=4505 b=26576 (the values from brazilian_logistics_study).
3. All four sides cropped 25%: only the centre 2x2 cells remain.
4. Negative left crop l=-20000: Word pads the left 1/6 of the frame with blank.

python-docx has no crop API, so the a:srcRect element is inserted after a:blip
directly with lxml.
"""
import io
import os

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

SRC_W, SRC_H = 400, 300
CELLS = 4
COLOURS = [
    (230, 25, 75), (60, 180, 75), (255, 225, 25), (0, 130, 200),
    (245, 130, 48), (145, 30, 180), (70, 240, 240), (240, 50, 230),
    (210, 245, 60), (250, 190, 190), (0, 128, 128), (220, 190, 255),
    (170, 110, 40), (255, 250, 200), (128, 0, 0), (170, 255, 195),
]


def grid_png() -> bytes:
    img = Image.new("RGB", (SRC_W, SRC_H), "black")
    draw = ImageDraw.Draw(img)
    cw, ch = SRC_W / CELLS, SRC_H / CELLS
    for row in range(CELLS):
        for col in range(CELLS):
            x0, y0 = col * cw + 3, row * ch + 3
            x1, y1 = (col + 1) * cw - 3, (row + 1) * ch - 3
            draw.rectangle([x0, y0, x1, y1], fill=COLOURS[row * CELLS + col])
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


PNG = grid_png()


def add_cropped_picture(doc, label, l=0, t=0, r=0, b=0, frame_w_in=3.0):
    doc.add_paragraph().add_run(label).bold = True

    vis_w = 1 - (l + r) / 100000
    vis_h = 1 - (t + b) / 100000
    frame_h_in = frame_w_in * (SRC_H * vis_h) / (SRC_W * vis_w)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic = p.add_run().add_picture(io.BytesIO(PNG), width=Inches(frame_w_in), height=Inches(frame_h_in))

    if any((l, t, r, b)):
        src_rect = OxmlElement("a:srcRect")
        for name, val in (("l", l), ("t", t), ("r", r), ("b", b)):
            if val:
                src_rect.set(name, str(val))
        blip_fill = pic._inline.graphic.graphicData.pic.blipFill
        blip_fill.blip.addnext(src_rect)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(section, side, Inches(1))

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Picture cropping (a:srcRect)")
run.bold = True
run.font.size = Pt(16)

add_cropped_picture(doc, "1. Uncropped control, 400×300 px at 3 in wide")
add_cropped_picture(doc, "2. Top 4.505% and bottom 26.576% cropped (t=4505 b=26576)", t=4505, b=26576)
add_cropped_picture(doc, "3. All sides 25% cropped: centre 2×2 cells only", l=25000, t=25000, r=25000, b=25000, frame_w_in=2.0)
add_cropped_picture(doc, "4. Negative left crop l=-20000: blank strip on the left", l=-20000, frame_w_in=3.6)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "input.docx")
doc.save(out)
print(f"wrote {out}")
