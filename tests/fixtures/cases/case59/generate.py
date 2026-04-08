#!/usr/bin/env python3
"""case59: SmartArt advanced properties — comprehensive test of rendering features.

Tests SmartArt properties beyond case37's basic shapes:
  Row 1: Rotation & flips (0°, 45°, 90°, flipH)
  Row 2: Text formatting (bold, italic, multi-run, multi-paragraph)
  Row 3: Text positioning (top-left+insets, center, bottom-right, txXfrm offset)
  Row 4: Stroke styles & scheme colors (dashed, dotted, thick, schemeClr+hueOff)
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

OUT = Path("tests/fixtures/cases/case59/input.docx")

# XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DSP_NS = "http://schemas.microsoft.com/office/drawing/2008/diagram"
DGM_NS = "http://schemas.openxmlformats.org/drawingml/2006/diagram"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
REL_TYPE_DIAGRAM_DRAWING = "http://schemas.microsoft.com/office/2007/relationships/diagramDrawing"
REL_TYPE_DIAGRAM_DATA = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramData"
REL_TYPE_DIAGRAM_LAYOUT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramLayout"
REL_TYPE_DIAGRAM_STYLE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramStyle"
REL_TYPE_DIAGRAM_COLORS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramColors"

EMU_PER_INCH = 914400
DIAG_W = int(6.5 * EMU_PER_INCH)
DIAG_H = int(3.5 * EMU_PER_INCH)

SHAPE_W = 1100000
SHAPE_H = 550000
COL_SPACING = DIAG_W // 4


def text_color_for_bg(fill_color):
    r, g, b = int(fill_color[0:2], 16), int(fill_color[2:4], 16), int(fill_color[4:6], 16)
    return "FFFFFF" if (0.299 * r + 0.587 * g + 0.114 * b) < 140 else "000000"


def make_run_xml(text, sz=1000, bold=False, italic=False, color="000000"):
    b_attr = ' b="1"' if bold else ""
    i_attr = ' i="1"' if italic else ""
    return (
        f'<a:r>'
        f'<a:rPr lang="en-US" sz="{sz}"{b_attr}{i_attr} dirty="0">'
        f'<a:solidFill><a:srgbClr val="{color}"/></a:solidFill>'
        f'</a:rPr>'
        f'<a:t>{text}</a:t>'
        f'</a:r>'
    )


def make_para_xml(runs_xml, algn="l", spc_bef=0, spc_aft=0):
    spc = ""
    if spc_bef:
        spc += f'<a:spcBef><a:spcPts val="{spc_bef}"/></a:spcBef>'
    if spc_aft:
        spc += f'<a:spcAft><a:spcPts val="{spc_aft}"/></a:spcAft>'
    return f'<a:p><a:pPr algn="{algn}">{spc}</a:pPr>{runs_xml}</a:p>'


def dsp_shape(
    preset, x, y, cx, cy, fill_color="4472C4",
    label=None, text_size=1000,
    rot=0, flip_h=False, flip_v=False,
    anchor="t", algn="l",
    l_ins=36000, t_ins=36000, r_ins=36000, b_ins=36000,
    stroke_color="404040", stroke_w=12700, dash_style=None,
    avlst_inner="",
    text_body_xml=None,
    tx_xfrm=None,
    scheme_fill=None,
):
    """Build a dsp:sp element with full property support."""
    # Transform
    flip_attrs = ""
    if flip_h:
        flip_attrs += ' flipH="1"'
    if flip_v:
        flip_attrs += ' flipV="1"'
    rot_attr = f' rot="{rot}"' if rot else ""

    xfrm = (
        f'<a:xfrm{rot_attr}{flip_attrs}>'
        f'<a:off x="{x}" y="{y}"/>'
        f'<a:ext cx="{cx}" cy="{cy}"/>'
        f'</a:xfrm>'
    )

    # Geometry
    avlst = f"<a:avLst>{avlst_inner}</a:avLst>" if avlst_inner else "<a:avLst/>"

    # Fill
    if scheme_fill:
        val, hue_off, sat_off = scheme_fill
        transforms = ""
        if hue_off:
            transforms += f'<a:hueOff val="{hue_off}"/>'
        if sat_off:
            transforms += f'<a:satOff val="{sat_off}"/>'
        fill_xml = f'<a:solidFill><a:schemeClr val="{val}">{transforms}</a:schemeClr></a:solidFill>'
    else:
        fill_xml = f'<a:solidFill><a:srgbClr val="{fill_color}"/></a:solidFill>'

    # Stroke
    dash_xml = ""
    if dash_style:
        dash_xml = f'<a:prstDash val="{dash_style}"/>'
    ln_xml = (
        f'<a:ln w="{stroke_w}" cap="flat" cmpd="sng">'
        f'<a:solidFill><a:srgbClr val="{stroke_color}"/></a:solidFill>'
        f'{dash_xml}'
        f'</a:ln>'
    )

    # Text body
    if text_body_xml is None and label:
        tc = text_color_for_bg(fill_color) if not scheme_fill else "000000"
        run = make_run_xml(label, sz=text_size, color=tc)
        para = make_para_xml(run, algn=algn)
        text_body_xml = para

    text_xml = ""
    if text_body_xml:
        text_xml = (
            f'<dsp:txBody>'
            f'<a:bodyPr anchor="{anchor}" lIns="{l_ins}" tIns="{t_ins}" rIns="{r_ins}" bIns="{b_ins}"'
            f' spcFirstLastPara="0" vert="horz" wrap="square"/>'
            f'<a:lstStyle/>'
            f'{text_body_xml}'
            f'</dsp:txBody>'
        )

    # txXfrm
    tx_xfrm_xml = ""
    if tx_xfrm:
        tx, ty, tcx, tcy = tx_xfrm
        tx_xfrm_xml = (
            f'<dsp:txXfrm>'
            f'<a:off x="{tx}" y="{ty}"/>'
            f'<a:ext cx="{tcx}" cy="{tcy}"/>'
            f'</dsp:txXfrm>'
        )

    return (
        f'<dsp:sp>'
        f'<dsp:nvSpPr><dsp:cNvPr id="0" name=""/><dsp:cNvSpPr/></dsp:nvSpPr>'
        f'<dsp:spPr>'
        f'{xfrm}'
        f'<a:prstGeom prst="{preset}">{avlst}</a:prstGeom>'
        f'{fill_xml}'
        f'{ln_xml}'
        f'</dsp:spPr>'
        f'{text_xml}'
        f'{tx_xfrm_xml}'
        f'</dsp:sp>'
    )


def col_x(col):
    return COL_SPACING * col + (COL_SPACING - SHAPE_W) // 2


def build_drawing_xml():
    shapes = []

    # === Row 1: Rotation & Flips ===
    y1 = 100000
    # No rotation
    shapes.append(dsp_shape("roundRect", col_x(0), y1, SHAPE_W, SHAPE_H, "4472C4",
                            label="No Rotation", anchor="ctr", algn="ctr",
                            avlst_inner='<a:gd name="adj" fmla="val 16667"/>'))
    # 45° rotation
    shapes.append(dsp_shape("roundRect", col_x(1), y1, SHAPE_W, SHAPE_H, "ED7D31",
                            label="45° Rotated", rot=2700000, anchor="ctr", algn="ctr",
                            avlst_inner='<a:gd name="adj" fmla="val 16667"/>'))
    # 90° rotation (L-shape)
    shapes.append(dsp_shape("corner", col_x(2), y1, 600000, SHAPE_H, "70AD47",
                            label="90° Corner", rot=5400000, anchor="t", algn="l",
                            l_ins=22860, t_ins=22860, r_ins=22860, b_ins=22860,
                            avlst_inner='<a:gd name="adj1" fmla="val 16120"/><a:gd name="adj2" fmla="val 16110"/>'))
    # FlipH
    shapes.append(dsp_shape("rightArrow", col_x(3), y1, SHAPE_W, SHAPE_H, "7030A0",
                            label="Flipped H", flip_h=True, anchor="ctr", algn="ctr"))

    # === Row 2: Text Formatting ===
    y2 = 850000
    # Bold + large
    bold_run = make_run_xml("Bold Large", sz=1400, bold=True, color="FFFFFF")
    shapes.append(dsp_shape("rect", col_x(0), y2, SHAPE_W, SHAPE_H, "C00000",
                            text_body_xml=make_para_xml(bold_run, algn="ctr"),
                            anchor="ctr"))
    # Italic + colored
    italic_run = make_run_xml("Italic Blue", sz=1000, italic=True, color="0070C0")
    shapes.append(dsp_shape("rect", col_x(1), y2, SHAPE_W, SHAPE_H, "FFF2CC",
                            text_body_xml=make_para_xml(italic_run, algn="ctr"),
                            anchor="ctr"))
    # Multi-run (two runs, different formatting)
    multi_runs = (
        make_run_xml("Bold ", sz=1000, bold=True, color="FFFFFF") +
        make_run_xml("and Italic", sz=1000, italic=True, color="FFFF00")
    )
    shapes.append(dsp_shape("rect", col_x(2), y2, SHAPE_W, SHAPE_H, "2F5496",
                            text_body_xml=make_para_xml(multi_runs, algn="ctr"),
                            anchor="ctr"))
    # Multi-paragraph
    p1 = make_para_xml(make_run_xml("First Paragraph", sz=900, color="000000"), algn="l", spc_aft=600)
    p2 = make_para_xml(make_run_xml("Second Paragraph", sz=900, bold=True, color="C00000"), algn="l", spc_bef=600)
    shapes.append(dsp_shape("rect", col_x(3), y2, SHAPE_W, SHAPE_H, "D9E2F3",
                            text_body_xml=p1 + p2, anchor="t",
                            l_ins=72000, t_ins=72000))

    # === Row 3: Text Positioning ===
    y3 = 1600000
    # Top-left with large insets
    shapes.append(dsp_shape("rect", col_x(0), y3, SHAPE_W, SHAPE_H, "E2EFDA",
                            label="Top-Left", anchor="t", algn="l",
                            l_ins=108000, t_ins=108000, stroke_color="70AD47"))
    # Center-center
    shapes.append(dsp_shape("ellipse", col_x(1), y3, SHAPE_W, SHAPE_H, "FCE4D6",
                            label="Centered", anchor="ctr", algn="ctr",
                            stroke_color="ED7D31"))
    # Bottom-right with insets
    shapes.append(dsp_shape("rect", col_x(2), y3, SHAPE_W, SHAPE_H, "D6DCE4",
                            label="Bottom-Right", anchor="b", algn="r",
                            r_ins=108000, b_ins=108000, stroke_color="4472C4"))
    # Separate txXfrm — shape is tall rectangle, text rect is offset to lower half
    shapes.append(dsp_shape("rect", col_x(3), y3, SHAPE_W, SHAPE_H, "FBE5D6",
                            label="txXfrm Offset", anchor="t", algn="l",
                            l_ins=36000, t_ins=36000,
                            stroke_color="ED7D31",
                            tx_xfrm=(col_x(3) + 50000, y3 + 300000, SHAPE_W - 100000, 200000)))

    # === Row 4: Stroke Styles & Colors ===
    y4 = 2350000
    # Dashed stroke
    shapes.append(dsp_shape("rect", col_x(0), y4, SHAPE_W, SHAPE_H, "DEEBF7",
                            label="Dashed", anchor="ctr", algn="ctr",
                            stroke_color="2F5496", stroke_w=19050, dash_style="dash"))
    # Dotted stroke
    shapes.append(dsp_shape("ellipse", col_x(1), y4, SHAPE_W, SHAPE_H, "E2EFDA",
                            label="Dotted", anchor="ctr", algn="ctr",
                            stroke_color="548235", stroke_w=19050, dash_style="dot"))
    # Thick solid
    shapes.append(dsp_shape("hexagon", col_x(2), y4, SHAPE_W, SHAPE_H, "FFF2CC",
                            label="Thick 3pt", anchor="ctr", algn="ctr",
                            stroke_color="BF8F00", stroke_w=38100))
    # Scheme color with hueOff
    shapes.append(dsp_shape("rect", col_x(3), y4, SHAPE_W, SHAPE_H,
                            label="Scheme+HueOff", anchor="ctr", algn="ctr",
                            scheme_fill=("accent1", -3000000, -500),
                            stroke_color="404040"))

    all_shapes = "\n".join(shapes)
    return (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<dsp:drawing xmlns:dsp="{DSP_NS}" xmlns:a="{A_NS}">'
        f'<dsp:spTree>'
        f'<dsp:nvGrpSpPr><dsp:cNvPr id="0" name=""/><dsp:cNvGrpSpPr/></dsp:nvGrpSpPr>'
        f'<dsp:grpSpPr/>'
        f'{all_shapes}'
        f'</dsp:spTree>'
        f'</dsp:drawing>'
    )


def build_inline_drawing_xml():
    return (
        f'<w:r>'
        f'<w:drawing>'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{DIAG_W}" cy="{DIAG_H}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="1" name="Diagram 1"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="{A_NS}">'
        f'<a:graphicData uri="{DGM_NS}">'
        f'<dgm:relIds xmlns:dgm="{DGM_NS}" xmlns:r="{R_NS}"'
        f' r:dm="rIdDgmData" r:lo="rIdDgmLayout" r:qs="rIdDgmStyle" r:cs="rIdDgmColors"/>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
        f'</w:r>'
    )


def build_diagram_data_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<dgm:dataModel xmlns:dgm="{DGM_NS}">'
        '<dgm:ptLst>'
        '<dgm:pt modelId="0" type="doc"><dgm:prSet/><dgm:spPr/></dgm:pt>'
        '</dgm:ptLst>'
        '<dgm:cxnLst/><dgm:bg/><dgm:whole/>'
        '</dgm:dataModel>'
    )


def build_diagram_layout_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<dgm:layoutDef xmlns:dgm="{DGM_NS}"/>'
    )


def build_diagram_style_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<dgm:styleDef xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' uniqueId="urn:microsoft.com/office/officeart/2005/8/quickstyle/simple1"/>'
    )


def build_diagram_colors_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<dgm:colorsDef xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' uniqueId="urn:microsoft.com/office/officeart/2005/8/colors/accent1_2"/>'
    )


# Step 1: Create base document
doc = Document()
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p = doc.add_paragraph()
p.add_run("SMARTART_PLACEHOLDER")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Step 2: Post-process ZIP
drawing_xml = build_drawing_xml()
inline_xml = build_inline_drawing_xml()

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()
        rels_xml = zin.read("word/_rels/document.xml.rels").decode()
        content_types_xml = zin.read("[Content_Types].xml").decode()

        # Add namespace declarations
        ns_decls = (
            f' xmlns:wp="{WP_NS}"'
            f' xmlns:a="{A_NS}"'
            f' xmlns:r="{R_NS}"'
            f' xmlns:mc="{MC_NS}"'
        )
        if 'xmlns:wp=' not in doc_xml:
            doc_xml = doc_xml.replace('<w:document ', f'<w:document {ns_decls} ', 1)

        # Replace placeholder
        placeholder_pattern = r'<w:r>.*?<w:t>SMARTART_PLACEHOLDER</w:t>\s*</w:r>'
        doc_xml = re.sub(placeholder_pattern, inline_xml, doc_xml, count=1, flags=re.DOTALL)

        # Add diagram relationships
        diagram_rels = (
            f'<Relationship Id="rIdDiagram1" Type="{REL_TYPE_DIAGRAM_DRAWING}" Target="diagrams/drawing1.xml"/>'
            f'<Relationship Id="rIdDgmData" Type="{REL_TYPE_DIAGRAM_DATA}" Target="diagrams/data1.xml"/>'
            f'<Relationship Id="rIdDgmLayout" Type="{REL_TYPE_DIAGRAM_LAYOUT}" Target="diagrams/layout1.xml"/>'
            f'<Relationship Id="rIdDgmStyle" Type="{REL_TYPE_DIAGRAM_STYLE}" Target="diagrams/style1.xml"/>'
            f'<Relationship Id="rIdDgmColors" Type="{REL_TYPE_DIAGRAM_COLORS}" Target="diagrams/colors1.xml"/>'
        )
        rels_xml = rels_xml.replace('</Relationships>', f'{diagram_rels}</Relationships>')

        # Add content types
        diagram_cts = (
            '<Override PartName="/word/diagrams/drawing1.xml"'
            ' ContentType="application/vnd.ms-office.drawingml.diagramDrawing+xml"/>'
            '<Override PartName="/word/diagrams/data1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramData+xml"/>'
            '<Override PartName="/word/diagrams/layout1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramLayout+xml"/>'
            '<Override PartName="/word/diagrams/style1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramStyle+xml"/>'
            '<Override PartName="/word/diagrams/colors1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramColors+xml"/>'
        )
        content_types_xml = content_types_xml.replace('</Types>', f'{diagram_cts}</Types>')

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, content_types_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

        zout.writestr("word/diagrams/drawing1.xml", drawing_xml)
        zout.writestr("word/diagrams/data1.xml", build_diagram_data_xml())
        zout.writestr("word/diagrams/layout1.xml", build_diagram_layout_xml())
        zout.writestr("word/diagrams/style1.xml", build_diagram_style_xml())
        zout.writestr("word/diagrams/colors1.xml", build_diagram_colors_xml())

os.unlink(tmp)
print(f"Generated {OUT}")
