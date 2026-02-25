"""Generate a DOCX that exercises table features not yet supported."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Table 1: Horizontal merged cells ---
doc.add_heading("Horizontal Merged Cells", level=2)

t1 = doc.add_table(rows=3, cols=4)
t1.style = "Table Grid"

# Header row with merged cells
t1.cell(0, 0).merge(t1.cell(0, 1)).text = "Name"
t1.cell(0, 2).merge(t1.cell(0, 3)).text = "Contact"

t1.cell(1, 0).text = "First"
t1.cell(1, 1).text = "Last"
t1.cell(1, 2).text = "Email"
t1.cell(1, 3).text = "Phone"

t1.cell(2, 0).text = "Alice"
t1.cell(2, 1).text = "Smith"
t1.cell(2, 2).text = "alice@example.com"
t1.cell(2, 3).text = "+1 555-0100"

doc.add_paragraph()

# --- Table 2: Vertical merged cells ---
doc.add_heading("Vertical Merged Cells", level=2)

t2 = doc.add_table(rows=4, cols=3)
t2.style = "Table Grid"

t2.cell(0, 0).text = "Region"
t2.cell(0, 1).text = "Product"
t2.cell(0, 2).text = "Revenue"

t2.cell(1, 0).merge(t2.cell(2, 0)).text = "North"
t2.cell(1, 1).text = "Widgets"
t2.cell(1, 2).text = "$120,000"
t2.cell(2, 1).text = "Gadgets"
t2.cell(2, 2).text = "$85,000"

t2.cell(3, 0).text = "South"
t2.cell(3, 1).text = "Widgets"
t2.cell(3, 2).text = "$95,000"

doc.add_paragraph()

# --- Table 3: Vertical alignment + row heights ---
doc.add_heading("Cell Vertical Alignment and Row Heights", level=2)

t3 = doc.add_table(rows=2, cols=3)
t3.style = "Table Grid"

# Set explicit row height
for row in t3.rows:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "720")  # 720 twips = 0.5 inch
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

labels = ["Top aligned", "Center aligned", "Bottom aligned"]
aligns = [WD_ALIGN_VERTICAL.TOP, WD_ALIGN_VERTICAL.CENTER, WD_ALIGN_VERTICAL.BOTTOM]

for i, (label, align) in enumerate(zip(labels, aligns)):
    cell = t3.cell(0, i)
    cell.text = label
    cell.vertical_alignment = align

    cell2 = t3.cell(1, i)
    cell2.text = f"Row 2, Col {i+1}"
    cell2.vertical_alignment = align

doc.add_paragraph()

# --- Table 4: Mixed horizontal and vertical merges ---
doc.add_heading("Complex Merge Pattern", level=2)

t4 = doc.add_table(rows=4, cols=4)
t4.style = "Table Grid"

# Top-left 2x2 merge
t4.cell(0, 0).merge(t4.cell(1, 1)).text = "2x2 merged"

t4.cell(0, 2).text = "A"
t4.cell(0, 3).text = "B"
t4.cell(1, 2).text = "C"
t4.cell(1, 3).text = "D"

# Bottom row full merge
t4.cell(3, 0).merge(t4.cell(3, 3)).text = "Full-width footer row"

t4.cell(2, 0).text = "E"
t4.cell(2, 1).text = "F"
t4.cell(2, 2).text = "G"
t4.cell(2, 3).text = "H"

doc.add_paragraph()

# --- Table 5: Colored borders (thick, colored) ---
doc.add_heading("Border Colors and Widths", level=2)

t5 = doc.add_table(rows=3, cols=3)
t5.style = "Table Grid"

data = [
    ["Header 1", "Header 2", "Header 3"],
    ["Normal", "Normal", "Normal"],
    ["Footer", "Footer", "Footer"],
]
for r, row_data in enumerate(data):
    for c, text in enumerate(row_data):
        t5.cell(r, c).text = text

# Apply thick red border to header row cells
for c in range(3):
    tc = t5.cell(0, c)._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ["top", "bottom", "left", "right"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")  # 1.5pt (sz is in 1/8 pt)
        el.set(qn("w:color"), "FF0000")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)

# Apply thick blue bottom border to footer
for c in range(3):
    tc = t5.cell(2, c)._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")  # 3pt
    bottom.set(qn("w:color"), "0000FF")
    bottom.set(qn("w:space"), "0")
    borders.append(bottom)
    tcPr.append(borders)

out = "tests/fixtures/case15/input.docx"
doc.save(out)
print(f"Saved to {out}")
