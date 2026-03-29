#!/usr/bin/env python3
"""case53: Charts with extreme data (many categories, very small/large values).

Tests chart rendering edge cases:
1. Bar chart with 50 categories (stress test for label spacing and bar width)
2. Line chart with very small values (0.001–0.05 range)
3. Bar chart with very large values (millions)
4. Mixed-range chart: one series in thousands, another in single digits

Usage:
    uv run tests/fixtures/cases/case53/generate.py
"""

import math
import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

OUT = Path("tests/fixtures/cases/case53/input.docx")

CHART_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
DML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CHART_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
CT_CHART = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"


def build_series_xml(idx, label, color_hex, categories, values, fmt="General"):
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{c}</c:v></c:pt>' for i, c in enumerate(categories)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill><a:ln><a:noFill/></a:ln></c:spPr>
  <c:cat><c:strRef><c:f>cats</c:f>
    <c:strCache><c:ptCount val="{len(categories)}"/>{cat_pts}</c:strCache>
  </c:strRef></c:cat>
  <c:val><c:numRef><c:f>{idx}</c:f>
    <c:numCache><c:formatCode>{fmt}</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
  </c:numRef></c:val>
</c:ser>"""


def build_line_series_xml(idx, label, color_hex, categories, values, fmt="General"):
    """Line chart series (same structure but for lineChart)."""
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{c}</c:v></c:pt>' for i, c in enumerate(categories)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr><a:ln w="19050"><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill></a:ln></c:spPr>
  <c:marker><c:symbol val="circle"/><c:size val="5"/>
    <c:spPr><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill></c:spPr>
  </c:marker>
  <c:cat><c:strRef><c:f>cats</c:f>
    <c:strCache><c:ptCount val="{len(categories)}"/>{cat_pts}</c:strCache>
  </c:strRef></c:cat>
  <c:val><c:numRef><c:f>{idx}</c:f>
    <c:numCache><c:formatCode>{fmt}</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
  </c:numRef></c:val>
</c:ser>"""


def build_axis_xml(ax_id, cross_id, position, axis_type="cat", gridlines=False):
    grid_xml = ""
    if gridlines:
        grid_xml = """<c:majorGridlines><c:spPr><a:ln>
      <a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill>
    </a:ln></c:spPr></c:majorGridlines>"""
    tag = "c:valAx" if axis_type == "val" else "c:catAx"
    return f"""<{tag}>
  <c:axId val="{ax_id}"/>
  <c:scaling><c:orientation val="minMax"/></c:scaling>
  <c:delete val="0"/>
  <c:axPos val="{position}"/>
  {grid_xml}
  <c:majorTickMark val="out"/><c:minorTickMark val="none"/>
  <c:tickLblPos val="nextTo"/>
  <c:spPr><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
  <c:crossAx val="{cross_id}"/>
  <c:crossesAt val="0"/>
</{tag}>"""


def build_bar_chart_xml(bar_dir, grouping, series_list, categories, gap_width,
                        legend_pos=None, colors=None, fmt="General"):
    default_colors = ["4472C4", "ED7D31", "A5A5A5", "FFC000", "5B9BD5", "70AD47"]
    if colors is None:
        colors = default_colors

    series_xml = ""
    for i, (label, values) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_series_xml(i, label, color, categories, values, fmt)

    cat_ax_id, val_ax_id = 100, 200
    if bar_dir == "bar":
        cat_pos, val_pos = "l", "b"
    else:
        cat_pos, val_pos = "b", "l"

    cat_ax = build_axis_xml(cat_ax_id, val_ax_id, cat_pos, axis_type="cat")
    val_ax = build_axis_xml(val_ax_id, cat_ax_id, val_pos, axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:barChart>
        <c:barDir val="{bar_dir}"/>
        <c:grouping val="{grouping}"/>
        {series_xml}
        <c:gapWidth val="{gap_width}"/>
        <c:axId val="{cat_ax_id}"/><c:axId val="{val_ax_id}"/>
      </c:barChart>
      {cat_ax}
      {val_ax}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_line_chart_xml(series_list, categories, legend_pos=None, colors=None, fmt="General"):
    default_colors = ["4472C4", "ED7D31", "A5A5A5"]
    if colors is None:
        colors = default_colors

    series_xml = ""
    for i, (label, values) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_line_series_xml(i, label, color, categories, values, fmt)

    cat_ax_id, val_ax_id = 100, 200
    cat_ax = build_axis_xml(cat_ax_id, val_ax_id, "b", axis_type="cat")
    val_ax = build_axis_xml(val_ax_id, cat_ax_id, "l", axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:lineChart>
        <c:grouping val="standard"/>
        {series_xml}
        <c:axId val="{cat_ax_id}"/><c:axId val="{val_ax_id}"/>
      </c:lineChart>
      {cat_ax}
      {val_ax}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_drawing_xml(rel_id, cx_emu, cy_emu):
    return (
        f'<w:drawing xmlns:w="{W_NS}" xmlns:wp="{WP_NS}" '
        f'xmlns:a="{DML_NS}" xmlns:c="{CHART_NS}" xmlns:r="{REL_NS}">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx_emu}" cy="{cy_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{rel_id.replace("rId","")}" name="Chart {rel_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic><a:graphicData uri="{CHART_NS}">'
        f'<c:chart r:id="{rel_id}"/>'
        f'</a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing>'
    )


def generate():
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.add_heading("Extreme Data Chart Tests", level=1)

    doc.add_paragraph("Chart 1: 50-category bar chart (stress test)")
    p1 = doc.add_paragraph()
    p1.add_run("CHART_PLACEHOLDER_1")

    doc.add_paragraph("")

    doc.add_paragraph("Chart 2: Very small values (0.001-0.05 range)")
    p2 = doc.add_paragraph()
    p2.add_run("CHART_PLACEHOLDER_2")

    doc.add_paragraph("")

    doc.add_paragraph("Chart 3: Very large values (millions)")
    p3 = doc.add_paragraph()
    p3.add_run("CHART_PLACEHOLDER_3")

    doc.add_paragraph("")

    doc.add_paragraph("Chart 4: Mixed range (thousands vs single digits)")
    p4 = doc.add_paragraph()
    p4.add_run("CHART_PLACEHOLDER_4")

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    # Generate 50 US state abbreviations as categories
    states = [
        "AL", "AK", "AZ", "AR", "CA", "CO", "CT", "DE", "FL", "GA",
        "HI", "ID", "IL", "IN", "IA", "KS", "KY", "LA", "ME", "MD",
        "MA", "MI", "MN", "MS", "MO", "MT", "NE", "NV", "NH", "NJ",
        "NM", "NY", "NC", "ND", "OH", "OK", "OR", "PA", "RI", "SC",
        "SD", "TN", "TX", "UT", "VT", "VA", "WA", "WV", "WI", "WY",
    ]
    # Pseudo-random values using a simple formula
    state_values = [round(20 + 80 * abs(math.sin(i * 1.7)), 1) for i in range(50)]

    # Small values for chart 2
    small_cats = ["Day " + str(i + 1) for i in range(12)]
    small_vals_a = [round(0.003 + 0.002 * math.sin(i * 0.8), 4) for i in range(12)]
    small_vals_b = [round(0.015 + 0.01 * math.cos(i * 0.6), 4) for i in range(12)]

    # Large values for chart 3
    large_cats = ["Region " + chr(65 + i) for i in range(8)]
    large_vals = [
        12500000, 8300000, 15700000, 6200000,
        22100000, 9800000, 11400000, 18600000,
    ]

    # Mixed range for chart 4
    mixed_cats = ["Product " + chr(65 + i) for i in range(6)]
    mixed_revenue = [4500, 3200, 5800, 2100, 6700, 3900]
    mixed_rating = [4.2, 3.8, 4.7, 3.1, 4.9, 4.0]

    charts = {
        1: build_bar_chart_xml(
            bar_dir="col", grouping="clustered", gap_width=50,
            legend_pos=None,
            categories=states,
            series_list=[("Population Index", state_values)],
            colors=["4472C4"],
        ),
        2: build_line_chart_xml(
            legend_pos="r",
            categories=small_cats,
            series_list=[
                ("Measurement A", small_vals_a),
                ("Measurement B", small_vals_b),
            ],
            colors=["4472C4", "ED7D31"],
            fmt="0.0000",
        ),
        3: build_bar_chart_xml(
            bar_dir="col", grouping="clustered", gap_width=100,
            legend_pos=None,
            categories=large_cats,
            series_list=[("Revenue ($)", large_vals)],
            colors=["2E75B6"],
            fmt="#,##0",
        ),
        4: build_bar_chart_xml(
            bar_dir="col", grouping="clustered", gap_width=100,
            legend_pos="r",
            categories=mixed_cats,
            series_list=[
                ("Revenue ($K)", mixed_revenue),
                ("Rating (1-5)", mixed_rating),
            ],
            colors=["4472C4", "ED7D31"],
        ),
    }

    chart_sizes = {
        1: (5943600, 2743200),  # wider for 50 categories
        2: (4572000, 2743200),
        3: (4572000, 2743200),
        4: (4572000, 2743200),
    }

    with zipfile.ZipFile(tmp, "r") as zin:
        with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
            rels_xml = zin.read("word/_rels/document.xml.rels").decode()
            ct_xml = zin.read("[Content_Types].xml").decode()

            existing_rids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_xml)]
            next_rid = max(existing_rids, default=0) + 1

            chart_rids = {}
            for chart_num in sorted(charts.keys()):
                chart_rids[chart_num] = f"rId{next_rid}"
                next_rid += 1

            new_rels = ""
            for chart_num, rid in chart_rids.items():
                new_rels += (
                    f'<Relationship Id="{rid}" '
                    f'Type="{CHART_REL_TYPE}" '
                    f'Target="charts/chart{chart_num}.xml"/>'
                )
            rels_xml = rels_xml.replace("</Relationships>", new_rels + "</Relationships>")

            new_ct = ""
            for chart_num in charts:
                new_ct += (
                    f'<Override PartName="/word/charts/chart{chart_num}.xml" '
                    f'ContentType="{CT_CHART}"/>'
                )
            ct_xml = ct_xml.replace("</Types>", new_ct + "</Types>")

            doc_xml = zin.read("word/document.xml").decode()
            for chart_num, rid in chart_rids.items():
                placeholder = f"CHART_PLACEHOLDER_{chart_num}"
                cx, cy = chart_sizes[chart_num]
                drawing = build_drawing_xml(rid, cx, cy)
                run_replacement = f'<w:r>{drawing}</w:r>'
                for pattern in [
                    f'<w:r><w:rPr></w:rPr><w:t>{placeholder}</w:t></w:r>',
                    f'<w:r><w:t>{placeholder}</w:t></w:r>',
                ]:
                    if pattern in doc_xml:
                        doc_xml = doc_xml.replace(pattern, run_replacement)
                        break
                else:
                    doc_xml = re.sub(
                        rf'<w:r[^>]*>.*?<w:t[^>]*>{placeholder}</w:t>.*?</w:r>',
                        run_replacement, doc_xml, flags=re.DOTALL,
                    )

            for item in zin.infolist():
                if item.filename == "word/_rels/document.xml.rels":
                    zout.writestr(item, rels_xml)
                elif item.filename == "[Content_Types].xml":
                    zout.writestr(item, ct_xml)
                elif item.filename == "word/document.xml":
                    zout.writestr(item, doc_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))

            for chart_num, chart_xml in charts.items():
                zout.writestr(f"word/charts/chart{chart_num}.xml", chart_xml)

    os.unlink(tmp)
    print(f"Generated {OUT}")


if __name__ == "__main__":
    generate()
