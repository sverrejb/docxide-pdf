#!/usr/bin/env python3
"""case72: w:lvlRestart (§17.9.10) — continuous numbering that never restarts.

By default a numbering level restarts at its `start` value whenever a more
significant (lower-ilvl) level advances. `<w:lvlRestart w:val="0"/>` overrides
that: the level counts continuously across the whole document, ignoring parent
advances. The corpus only ever carries lvlRestart on unused or never-triggering
levels (Word boilerplate), so the feature is genuinely unexercised. This fixture
isolates the visible case with a two-level list:

  ilvl 0: decimal, lvlText "%1."          -> "1.", "2."  (section headings)
  ilvl 1: decimal, <w:lvlRestart w:val="0"/>, lvlText "(%2)"

Document order: Section1, sub, sub, Section2, sub, sub. With lvlRestart=0 the
sub-items run continuously (1)(2)(3)(4) — the (3)(4) carry past the Section2
heading instead of restarting. A renderer applying the default restart rule
would print (1)(2)(1)(2). The parenthesized sub form stays visually distinct
from the "1."/"2." section form so the continuity is unmistakable.
"""

import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case72/input.docx")

NUM_ID = 100
ABSTRACT_ID = 100

# ilvl 1 carries lvlRestart=0 -> never resets when ilvl 0 advances.
ABSTRACT_XML = f"""<w:abstractNum w:abstractNumId="{ABSTRACT_ID}">
  <w:multiLevelType w:val="multilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlText w:val="%1."/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>
  </w:lvl>
  <w:lvl w:ilvl="1">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlRestart w:val="0"/>
    <w:lvlText w:val="(%2)"/>
    <w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
  </w:lvl>
</w:abstractNum>"""

NUM_XML = f'<w:num w:numId="{NUM_ID}"><w:abstractNumId w:val="{ABSTRACT_ID}"/></w:num>'


def set_num(paragraph, ilvl):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(NUM_ID))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def main():
    doc = Document()

    title = doc.add_paragraph()
    r = title.add_run("lvlRestart=0 continuous numbering")
    r.bold = True
    r.font.size = Pt(14)

    # expect: 1. / (1) (2) / 2. / (3) (4)  — sub counter never restarts
    items = [
        (0, "First section"),
        (1, "Continuous item"),
        (1, "Continuous item"),
        (0, "Second section"),
        (1, "Continuous item"),
        (1, "Continuous item"),
    ]
    for ilvl, text in items:
        p = doc.add_paragraph(text)
        set_num(p, ilvl)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    num_xml = files["word/numbering.xml"].decode("utf-8")
    # CT_Numbering child order is abstractNum* then num*: drop the abstractNum in
    # with the other abstractNums, the num just before the closing tag.
    assert "<w:num " in num_xml, "default template should ship at least one w:num"
    num_xml = num_xml.replace("<w:num ", ABSTRACT_XML + "<w:num ", 1)
    num_xml = num_xml.replace("</w:numbering>", NUM_XML + "</w:numbering>", 1)
    files["word/numbering.xml"] = num_xml.encode("utf-8")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
