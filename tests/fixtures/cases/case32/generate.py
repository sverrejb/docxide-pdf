"""Generate a DOCX with a floating table and inline tblBorders for case32.

Tests:
- Floating table positioning (w:tblpPr with center alignment, vertical offset)
- Inline w:tblBorders (single borders on tblPr, not via named style)
- Vertical flow: text after floating table not pushed down by table height
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT = Path("tests/fixtures/cases/case32/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p1 = doc.add_paragraph(
    "This paragraph appears above the floating table. "
    "The table should be centered on the margin and offset below this anchor paragraph. "
    "Normal body text continues to flow here."
)

# Placeholder table that we'll modify via ZIP post-processing
table = doc.add_table(rows=3, cols=3)
cells = [
    ("Item", "Qty", "Price"),
    ("Apples", "5", "$3.00"),
    ("Bread", "2", "$4.50"),
]
for ri, row_data in enumerate(cells):
    for ci, text in enumerate(row_data):
        table.rows[ri].cells[ci].text = text

p2 = doc.add_paragraph(
    "This paragraph appears after the floating table in document order. "
    "Since the table is floating, this text should flow normally without being "
    "pushed down by the table height. It should appear near the top of the page, "
    "right after the first paragraph."
)

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Post-process: inject tblpPr (floating) and tblBorders (inline) into the table
TBL_P_PR = (
    f'<w:tblpPr {W_NS.join([""])}' # namespace handled at document level
    f' w:vertAnchor="text"'
    f' w:horzAnchor="margin"'
    f' w:tblpXSpec="center"'
    f' w:tblpY="200"/>'
)

TBL_BORDERS = (
    '<w:tblBorders>'
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '</w:tblBorders>'
)

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        # Inject tblpPr and tblBorders into tblPr.
        # python-docx generates <w:tblPr> with style and width — insert after the opening tag.
        # Find the first <w:tblPr> and inject our properties right after the existing children.
        # We'll insert before </w:tblPr>.
        doc_xml = doc_xml.replace(
            "</w:tblPr>",
            TBL_P_PR + TBL_BORDERS + "</w:tblPr>",
            1,  # only first table
        )

        # Also remove the tblStyle reference so borders come only from inline tblBorders
        doc_xml = re.sub(
            r'<w:tblStyle w:val="[^"]*"/>',
            "",
            doc_xml,
            count=1,
        )

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)
print(f"Generated {OUT}")
