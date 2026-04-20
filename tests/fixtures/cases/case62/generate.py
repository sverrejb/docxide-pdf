"""
Case 62: run-level shading (w:rPr/w:shd)

Exercises the run-shading path (separate from w:highlight and from
paragraph/cell shading) across the scenarios most likely to stress the
implementation:

  1. Visible pale-yellow shd on a run (baseline — does it paint at all?)
  2. highlight + shd on the same run, different colors (highlight draws
     on top of shd, both must be carried)
  3. Character-style shading inherited by a run (no inline w:shd on the run)
  4. w:shd w:val="nil" on a run whose character style sets shading.
     Word does NOT treat val="nil" as an override that clears inherited
     shading — the character style's pink wins. Both the baseline run
     and the "override" run must render pink.
  5. Shaded run that wraps across a line break (shading must cover both
     line fragments, not just the first)
  6. Adjacent runs: first two share a shading color (merge optimization
     should produce one rectangle), then a third run with a different
     color (no bleed, no seam)
  7. Shaded runs at different font sizes in the same line (merge-and-flush
     takes max(font_size) — rect must cover the taller glyphs)

python-docx can't express w:shd at the run level directly, so we build
the paragraphs with plain runs and then post-process word/document.xml
and word/styles.xml via lxml to inject the properties.
"""

from docx import Document
from docx.shared import Pt
import zipfile, os, shutil
from lxml import etree

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns = {'w': NS}
W = f'{{{NS}}}'


def w(el):
    return f'{W}{el}'


def make_rpr_with_shd(*, val='clear', fill=None, highlight=None, sz_halfpt=None,
                     rstyle=None, with_nil_shd=False):
    """Build a <w:rPr> element with optional shd / highlight / rStyle / sz."""
    rpr = etree.Element(w('rPr'))
    if rstyle is not None:
        r_style = etree.SubElement(rpr, w('rStyle'))
        r_style.set(w('val'), rstyle)
    if sz_halfpt is not None:
        sz = etree.SubElement(rpr, w('sz'))
        sz.set(w('val'), str(sz_halfpt))
    if highlight is not None:
        hl = etree.SubElement(rpr, w('highlight'))
        hl.set(w('val'), highlight)
    if with_nil_shd:
        shd = etree.SubElement(rpr, w('shd'))
        shd.set(w('val'), 'nil')
    elif fill is not None:
        shd = etree.SubElement(rpr, w('shd'))
        shd.set(w('val'), val)
        shd.set(w('color'), 'auto')
        shd.set(w('fill'), fill)
    return rpr


def build_paragraph(parent_body, tag_before, runs):
    """Create a <w:p> with the given runs.
    runs is a list of (text, rpr_kwargs_or_None) tuples.
    tag_before is an existing element; the new <w:p> is inserted after it.
    """
    p = etree.Element(w('p'))
    for text, rpr_kwargs in runs:
        r = etree.SubElement(p, w('r'))
        if rpr_kwargs is not None:
            r.append(make_rpr_with_shd(**rpr_kwargs))
        t = etree.SubElement(r, w('t'))
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
    tag_before.addnext(p)
    return p


# ---- Step 1: build a plain doc with headings as anchors ----------------------
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

doc.add_paragraph('Case 62: run-level shading (w:rPr/w:shd)', style='Heading 1')

# Each heading is an anchor after which we splice in the test paragraph(s).
# We tag them with unique marker text so the post-processor can find them.
ANCHORS = [
    ('S1', 'Section 1: visible pale-yellow shading'),
    ('S2', 'Section 2: highlight + shading, different colors'),
    ('S3', 'Section 3: shading inherited from character style'),
    ('S4', 'Section 4: val="nil" does not override inherited shading'),
    ('S5', 'Section 5: shaded run wraps across a line break'),
    ('S6', 'Section 6: adjacent runs — merge same, split different'),
    ('S7', 'Section 7: shaded runs at different font sizes'),
]
for marker, heading in ANCHORS:
    # Put the marker as the first word so we can grep for it uniquely.
    doc.add_paragraph(f'{marker} {heading}', style='Heading 2')

# Save a scratch copy, then surgically edit.
script_dir = os.path.dirname(os.path.abspath(__file__))
tmp_path = os.path.join(script_dir, '_tmp.docx')
final_path = os.path.join(script_dir, 'input.docx')
doc.save(tmp_path)


# ---- Step 2: rewrite document.xml and styles.xml ----------------------------

def process_document(xml_bytes):
    root = etree.fromstring(xml_bytes)
    body = root.find('w:body', ns)

    # Find each anchor heading by its marker text
    def find_heading(marker):
        for p in body.findall('w:p', ns):
            for t in p.findall('.//w:t', ns):
                if t.text and t.text.startswith(marker + ' '):
                    return p
        raise RuntimeError(f'anchor {marker} not found')

    # Section 1 — visible pale yellow
    build_paragraph(body, find_heading('S1'), [
        ('Normal text, then ', None),
        ('this run has pale yellow shading', {'fill': 'FFF4A3'}),
        (', then normal again.', None),
    ])

    # Section 2 — highlight + shading, different colors
    build_paragraph(body, find_heading('S2'), [
        ('Normal, then ', None),
        ('this run is yellow highlight on a cyan shading fill',
         {'fill': '00FFFF', 'highlight': 'yellow'}),
        (', normal.', None),
    ])

    # Section 3 — inherit from character style (no inline shd on this run)
    build_paragraph(body, find_heading('S3'), [
        ('Normal, ', None),
        ('this run uses char style PinkShd — no inline w:shd',
         {'rstyle': 'PinkShd'}),
        (', normal.', None),
    ])

    # Section 4 — val="nil" alongside rStyle should NOT clear inherited
    # shading (Word's actual behavior). Both shaded runs must render pink.
    build_paragraph(body, find_heading('S4'), [
        ('Baseline: ', None),
        ('inherits pink from style', {'rstyle': 'PinkShd'}),
        ('. Attempted override: ', None),
        ('this run has rStyle=PinkShd AND w:shd val="nil" — Word keeps the pink',
         {'rstyle': 'PinkShd', 'with_nil_shd': True}),
        ('.', None),
    ])

    # Section 5 — shading wraps across line break
    # Long paragraph where one middle run is long enough to span a wrap.
    long_shaded = (
        'this is a deliberately long shaded span that will wrap across a '
        'line break so the shading rectangle must be drawn on two lines '
        'and not only the first line of the paragraph'
    )
    build_paragraph(body, find_heading('S5'), [
        ('Start. ', None),
        (long_shaded, {'fill': 'B0E0E6'}),  # powder blue
        (' End.', None),
    ])

    # Section 6 — adjacent runs: same color merges, different color does not
    build_paragraph(body, find_heading('S6'), [
        ('Normal. ', None),
        ('[first shaded run green] ', {'fill': 'C8FACC'}),
        ('[second shaded run green — should merge with first] ',
         {'fill': 'C8FACC'}),
        ('[third shaded run ORANGE — must not bleed green into it]',
         {'fill': 'FFD8A8'}),
        (' normal.', None),
    ])

    # Section 7 — mixed font sizes with shared shading color
    # Two contiguous runs sharing fill but with font sizes 10 and 18 —
    # merge uses max(font_size), rectangle must cover the taller glyphs.
    build_paragraph(body, find_heading('S7'), [
        ('Shading across sizes: ', None),
        ('small 10pt', {'fill': 'FFE4B5', 'sz_halfpt': 20}),
        (' ', {'fill': 'FFE4B5'}),
        ('BIG 18pt', {'fill': 'FFE4B5', 'sz_halfpt': 36}),
        (' end.', None),
    ])

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def process_styles(xml_bytes):
    """Inject a character style `PinkShd` with w:rPr/w:shd fill=F8BBD0."""
    root = etree.fromstring(xml_bytes)
    style = etree.SubElement(root, w('style'))
    style.set(w('type'), 'character')
    style.set(w('styleId'), 'PinkShd')
    name = etree.SubElement(style, w('name'))
    name.set(w('val'), 'PinkShd')
    rpr = make_rpr_with_shd(fill='F8BBD0')
    style.append(rpr)
    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


with zipfile.ZipFile(tmp_path, 'r') as zin:
    with zipfile.ZipFile(final_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/document.xml':
                data = process_document(data)
            elif item.filename == 'word/styles.xml':
                data = process_styles(data)
            zout.writestr(item, data)

os.remove(tmp_path)
print(f'Created {final_path}')
