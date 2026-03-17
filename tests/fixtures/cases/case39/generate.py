#!/usr/bin/env python3
"""case39: TOC internal link navigation.

Tests that clicking a TOC entry navigates to the corresponding heading in the PDF.

Document structure:
- TOC section at top with two hyperlink entries pointing to anchor bookmarks
- Heading 1 "Introduction" with bookmark _Toc001
- Body text paragraph
- Heading 2 "Methods" with bookmark _Toc002
- Body text paragraph

Bookmarks and anchor hyperlinks are injected via ZIP post-processing
because python-docx does not expose these natively.
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

OUT = Path("tests/fixtures/cases/case39/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def toc_entry_xml(display_text, anchor_name):
    """TOC paragraph with an anchor hyperlink (no pPr — keep it simple)."""
    return (
        f'<w:p>'
        f'<w:hyperlink w:anchor="{anchor_name}">'
        f'<w:r>'
        f'<w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{display_text}</w:t>'
        f'</w:r>'
        f'</w:hyperlink>'
        f'</w:p>'
    )


def heading_with_bookmark_xml(text, heading_style, bookmark_name, bookmark_id):
    """Heading paragraph with bookmarkStart/End wrapping the run."""
    return (
        f'<w:p>'
        f'<w:pPr><w:pStyle w:val="{heading_style}"/></w:pPr>'
        f'<w:bookmarkStart w:id="{bookmark_id}" w:name="{bookmark_name}"/>'
        f'<w:r><w:t>{text}</w:t></w:r>'
        f'<w:bookmarkEnd w:id="{bookmark_id}"/>'
        f'</w:p>'
    )


# Step 1: Build base document with placeholder markers
doc = Document()
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

doc.add_paragraph("TOC_ENTRY_1")
doc.add_paragraph("TOC_ENTRY_2")
doc.add_paragraph("HEADING1_MARK")
doc.add_paragraph("Body text for the Introduction section.")
doc.add_paragraph("HEADING2_MARK")
doc.add_paragraph("Body text for the Methods section.")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Step 2: Post-process ZIP — replace placeholder paragraphs with real XML
replacements = [
    (
        r'<w:p><w:r><w:t>TOC_ENTRY_1</w:t></w:r></w:p>',
        toc_entry_xml("Introduction", "_Toc001"),
    ),
    (
        r'<w:p><w:r><w:t>TOC_ENTRY_2</w:t></w:r></w:p>',
        toc_entry_xml("Methods", "_Toc002"),
    ),
    (
        r'<w:p><w:r><w:t>HEADING1_MARK</w:t></w:r></w:p>',
        heading_with_bookmark_xml("Introduction", "Heading1", "_Toc001", 1),
    ),
    (
        r'<w:p><w:r><w:t>HEADING2_MARK</w:t></w:r></w:p>',
        heading_with_bookmark_xml("Methods", "Heading2", "_Toc002", 2),
    ),
]

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        for pattern, replacement in replacements:
            doc_xml = re.sub(pattern, replacement, doc_xml, count=1)

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)
print(f"Generated {OUT}")
