#!/usr/bin/env python3
"""case68: w:pgBorders (Page Borders, OOXML 17.6.10) — plain-line border box.

Isolates the section-level page-border feature: a stroked rectangle drawn
around every page. Tests the four CT_Border children (top/left/bottom/right,
17.6.2/.7/.15/.21) with a plain `single` line style plus the `@offsetFrom`
placement attribute. Art (image) borders are deliberately out of scope.
"""

import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case68/input.docx")

# 17.6.10: offsetFrom="page" measures @space from the page edge (points).
# Each child border is a CT_Border: @val style, @sz eighths-of-a-point,
# @space points (0-31), @color. sz="24" = 3pt line; space="24" = 24pt inset.
PG_BORDERS = (
    '<w:pgBorders w:offsetFrom="page">'
    '<w:top w:val="single" w:sz="24" w:space="24" w:color="1F4E79"/>'
    '<w:left w:val="single" w:sz="24" w:space="24" w:color="1F4E79"/>'
    '<w:bottom w:val="single" w:sz="24" w:space="24" w:color="1F4E79"/>'
    '<w:right w:val="single" w:sz="24" w:space="24" w:color="1F4E79"/>'
    '</w:pgBorders>'
)


def main():
    doc = Document()
    run = doc.add_paragraph().add_run("Page Border Test")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph(
        "This page is enclosed by a w:pgBorders box (single line, 3pt, "
        "inset 24pt from the page edge)."
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode("utf-8")
        other = {n: zin.read(n) for n in zin.namelist() if n != "word/document.xml"}

    # pgBorders must precede w:cols in the CT_SectPr child sequence.
    if "<w:cols" in doc_xml:
        doc_xml = doc_xml.replace("<w:cols", PG_BORDERS + "<w:cols", 1)
    else:
        doc_xml = doc_xml.replace("</w:sectPr>", PG_BORDERS + "</w:sectPr>", 1)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
