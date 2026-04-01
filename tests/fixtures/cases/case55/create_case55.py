"""Test fixture for table conditional formatting (tblLook + tblStylePr).

Creates tables using built-in Word styles that have conditional formatting:
- Medium Shading 2 Accent 1: blue header row, alternating row shading
- Light Grid Accent 3: green header, grid borders
- Plain table with no style (control)
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
p = doc.add_paragraph("Table Conditional Formatting Test")
p.style = doc.styles['Heading 1']

# Table 1: Medium Shading 2 Accent 1 (blue header, banded rows)
doc.add_paragraph("Table 1: Medium Shading 2 - Accent 1")
table1 = doc.add_table(rows=6, cols=4, style='Medium Shading 2 Accent 1')
headers = ["Name", "Department", "Score", "Status"]
for i, h in enumerate(headers):
    table1.rows[0].cells[i].text = h
data = [
    ["Alice Johnson", "Engineering", "95", "Pass"],
    ["Bob Smith", "Marketing", "82", "Pass"],
    ["Carol White", "Engineering", "91", "Pass"],
    ["David Brown", "Sales", "67", "Fail"],
    ["Eve Davis", "Marketing", "88", "Pass"],
]
for ri, row_data in enumerate(data):
    for ci, val in enumerate(row_data):
        table1.rows[ri + 1].cells[ci].text = val

doc.add_paragraph("")  # spacer

# Table 2: Light Grid Accent 3 (green header, grid lines)
doc.add_paragraph("Table 2: Light Grid - Accent 3")
table2 = doc.add_table(rows=5, cols=3, style='Light Grid Accent 3')
headers2 = ["Item", "Quantity", "Price"]
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
data2 = [
    ["Widget A", "100", "$12.50"],
    ["Widget B", "250", "$8.75"],
    ["Widget C", "50", "$25.00"],
    ["Widget D", "175", "$15.00"],
]
for ri, row_data in enumerate(data2):
    for ci, val in enumerate(row_data):
        table2.rows[ri + 1].cells[ci].text = val

doc.add_paragraph("")  # spacer

# Table 3: Plain Table Grid (control — no conditional formatting)
doc.add_paragraph("Table 3: Table Grid (control)")
table3 = doc.add_table(rows=4, cols=3, style='Table Grid')
for ri in range(4):
    for ci in range(3):
        table3.rows[ri].cells[ci].text = f"R{ri+1}C{ci+1}"

doc.save("tests/fixtures/cases/case55/input.docx")
print("Created case55/input.docx")
