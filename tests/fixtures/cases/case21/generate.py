"""Generate case21: Two Equal Columns with Column Break

Tests:
- Two equal-width columns (w:cols w:num="2")
- Column break (w:br w:type="column")
- Text flowing from first column to second column
- Paragraph spanning both columns (not really — just filling both)
"""

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Title
p = doc.add_paragraph("Two-Column Layout")
p.style = doc.styles["Heading 1"]

# Paragraphs that will flow into column 1
doc.add_paragraph(
    "This is the first column of text. It demonstrates a basic two-column layout "
    "where text flows naturally from the first column to the second column when the "
    "first column is full."
)

doc.add_paragraph(
    "Additional content in the first column. This paragraph adds more text to show "
    "that the column break works correctly and text appears in the right column."
)

# Column break paragraph — we'll inject w:br type="column" via post-processing
doc.add_paragraph("COLUMN_BREAK_MARKER")

# Text for column 2
doc.add_paragraph(
    "This text should appear in the second column after the explicit column break. "
    "Column breaks force remaining content to the next column."
)

doc.add_paragraph(
    "Final paragraph in the second column. This verifies that text positioning "
    "is correct in the right column."
)

# Save, then post-process to set columns and inject column break
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/document.xml":
            tree = etree.fromstring(data)
            nsmap = {"w": WML}

            # Replace w:cols in sectPr
            body = tree.find("w:body", nsmap)
            sect_pr = body.find("w:sectPr", nsmap)
            for old_cols in sect_pr.findall("w:cols", nsmap):
                sect_pr.remove(old_cols)
            cols = etree.SubElement(sect_pr, qn("w:cols"))
            cols.set(qn("w:num"), "2")
            cols.set(qn("w:space"), "720")  # 0.5 inch gap

            # Find the marker paragraph and replace its text with a column break
            for p in body.findall(".//w:p", nsmap):
                for r in p.findall(".//w:r", nsmap):
                    t = r.find("w:t", nsmap)
                    if t is not None and t.text == "COLUMN_BREAK_MARKER":
                        # Remove the text element, add a break
                        r.remove(t)
                        br = etree.SubElement(r, qn("w:br"))
                        br.set(qn("w:type"), "column")

            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        elif item.filename == "word/settings.xml":
            tree = etree.fromstring(data)
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
