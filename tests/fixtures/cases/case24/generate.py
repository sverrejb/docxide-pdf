"""Generate case24: Text Expansion/Compression (w:w) and keepLines

Tests:
- w:w (text scale percentage) at various values (50%, 100%, 150%, 200%)
- w:keepLines (prevent paragraph splitting across pages)
- Combination of text scaling with other formatting
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_scaled_run(paragraph, text, scale_pct, bold=False):
    """Add a run with explicit text scaling (w:w)."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    rpr = run._element.get_or_add_rPr()
    if scale_pct != 100:
        w_elem = OxmlElement("w:w")
        w_elem.set(qn("w:val"), f"{scale_pct}%")
        rpr.append(w_elem)
    return run


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Title
doc.add_heading("Text Expansion/Compression & Keep Lines", level=1)

# Section 1: Text scaling examples
doc.add_heading("Text Scaling (w:w)", level=2)

p = doc.add_paragraph()
add_scaled_run(p, "Normal (100%). ", 100)
add_scaled_run(p, "Expanded 150%. ", 150)
add_scaled_run(p, "Expanded 200%. ", 200)

p = doc.add_paragraph()
add_scaled_run(p, "Compressed 50%. ", 50)
add_scaled_run(p, "Normal again. ", 100)
add_scaled_run(p, "Compressed 66%. ", 66)

p = doc.add_paragraph()
add_scaled_run(p, "This entire paragraph has 150% text expansion applied to every word. "
    "This tests how text scaling affects line breaking and text width calculations "
    "across a longer block of text.", 150)

p = doc.add_paragraph()
add_scaled_run(p, "This paragraph uses 80% compression to fit more text per line. "
    "It should be noticeably narrower than normal text while still being readable.", 80)

# Section 2: Mixed scaling
doc.add_heading("Mixed Scaling", level=2)

p = doc.add_paragraph()
add_scaled_run(p, "Wide ", 200, bold=True)
add_scaled_run(p, "and narrow ", 50)
add_scaled_run(p, "text in the same line. ", 100)
add_scaled_run(p, "Back to wide.", 200)

# Section 3: keepLines
doc.add_heading("Keep Lines Test", level=2)

# Add enough filler to push to near page break
for i in range(30):
    doc.add_paragraph(
        f"Filler paragraph {i+1} to push content toward the page break. "
        "This text occupies space so the keepLines paragraph is near the bottom."
    )

# This paragraph should NOT be split across pages
p = doc.add_paragraph(
    "THIS PARAGRAPH HAS keepLines ENABLED. It should not be split across pages. "
    "If there is not enough room on the current page, the entire paragraph should "
    "move to the next page rather than being split. This is important for maintaining "
    "readability of cohesive text blocks. Adding more text to make it multi-line so "
    "that the keepLines property has an observable effect on the output."
)

# Save
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(
    out_buf, "w", zipfile.ZIP_DEFLATED
) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/document.xml":
            tree = etree.fromstring(data)
            nsmap = {"w": WML}
            # Find the last paragraph and add keepLines to its pPr
            body = tree.find("w:body", nsmap)
            paragraphs = body.findall("w:p", nsmap)
            last_p = paragraphs[-1]
            ppr = last_p.find("w:pPr", nsmap)
            if ppr is None:
                ppr = etree.SubElement(last_p, qn("w:pPr"))
                last_p.insert(0, ppr)
            keep_lines = etree.SubElement(ppr, qn("w:keepLines"))
            data = etree.tostring(
                tree, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        elif item.filename == "word/settings.xml":
            tree = etree.fromstring(data)
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(
                tree, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
