#!/usr/bin/env python3
"""case47: Floating image bottom gap diagnostic.

One image per page, varying image height. Left-aligned with Square wrapping,
followed by enough text to wrap past the image bottom.

Section A (pages 1-10): image anchored to heading paragraph, text starts at image top.
Section B (pages 11-16): text paragraph BEFORE the image anchor, so text wraps from
above the image down past its bottom — tests the look-ahead wrapping path.

All images use distB=72 twips (3.6pt) bottomFromText.
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Emu

OUT = Path("tests/fixtures/cases/case47/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"

EMU_PER_PT = 12700

LOREM = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
    "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
    "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris "
    "nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in "
    "reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla "
    "pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
    "culpa qui officia deserunt mollit anim id est laborum. "
)

# Section A: image at heading, text wraps beside from the top
HEIGHTS_A = [50, 70, 80, 100, 108, 120, 140, 160, 180, 200]
# Section B: text starts above image, wraps down past it
HEIGHTS_B = [50, 80, 108, 140, 180, 200]
IMAGE_WIDTH_PT = 150
DIST_B = 72  # twips = 3.6pt


def make_anchor_xml(rid, width_pt, height_pt, page_idx, v_offset_emu=0):
    """Generate wp:anchor XML for a left-aligned floating image with Square wrap."""
    cx = int(width_pt * EMU_PER_PT)
    cy = int(height_pt * EMU_PER_PT)
    dist_t = int(DIST_B * EMU_PER_PT / 20)  # twips to EMU
    dist_b = dist_t
    dist_l = int(114300)  # ~9pt
    dist_r = int(114300)

    return f"""<w:drawing xmlns:w="{W_NS}">
      <wp:anchor distT="{dist_t}" distB="{dist_b}" distL="{dist_l}" distR="{dist_r}"
          simplePos="0" relativeHeight="{251658240 + page_idx}" behindDoc="0"
          locked="0" layoutInCell="1" allowOverlap="1"
          xmlns:wp="{WP_NS}">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="paragraph"><wp:posOffset>{v_offset_emu}</wp:posOffset></wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapSquare wrapText="largest"/>
        <wp:docPr id="{100 + page_idx}" name="Image {page_idx}"/>
        <a:graphic xmlns:a="{A_NS}">
          <a:graphicData uri="{PIC_NS}">
            <pic:pic xmlns:pic="{PIC_NS}">
              <pic:nvPicPr>
                <pic:cNvPr id="{100 + page_idx}" name="img{page_idx}.png"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{rid}" xmlns:r="{R_NS}"/>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>"""


def make_test_image(width, height):
    """Create a simple colored PNG."""
    from PIL import Image
    img = Image.new("RGB", (width, height), (100, 180, 100))
    import io
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def main():
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    from docx.enum.text import WD_BREAK

    all_pages = []  # (page_num, h_pt, section, has_text_before)

    # Section A: image anchored to heading, text starts beside image
    for i, h_pt in enumerate(HEIGHTS_A):
        page_num = i + 1
        all_pages.append((page_num, h_pt, "A"))
        p = doc.add_paragraph()
        run = p.add_run(f"Page {page_num}: A {h_pt}pt")
        run.bold = True
        doc.add_paragraph(LOREM * 4)
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    # Section B: text paragraph before the image anchor, wraps from above
    for i, h_pt in enumerate(HEIGHTS_B):
        page_num = len(HEIGHTS_A) + i + 1
        all_pages.append((page_num, h_pt, "B"))
        # The image will be anchored to THIS heading, but text before it wraps
        p = doc.add_paragraph()
        run = p.add_run(f"Page {page_num}: B {h_pt}pt (text before)")
        run.bold = True
        # Body text that starts BEFORE the image and wraps past it
        doc.add_paragraph(LOREM * 5)
        if i < len(HEIGHTS_B) - 1:
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    img_data = make_test_image(int(IMAGE_WIDTH_PT * 2), 400)

    with zipfile.ZipFile(tmp, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode()
        rels_xml = zin.read("word/_rels/document.xml.rels").decode()
        other_files = {}
        for item in zin.infolist():
            if item.filename not in ("word/document.xml", "word/_rels/document.xml.rels"):
                other_files[item.filename] = zin.read(item.filename)

    rid = "rIdImg1"
    img_path = "word/media/image1.png"

    rels_xml = rels_xml.replace(
        "</Relationships>",
        f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image1.png"/></Relationships>',
    )

    ct_path = "[Content_Types].xml"
    if ct_path in other_files:
        ct = other_files[ct_path].decode() if isinstance(other_files[ct_path], bytes) else other_files[ct_path]
        if "png" not in ct.lower():
            ct = ct.replace("</Types>", '<Default Extension="png" ContentType="image/png"/></Types>')
            other_files[ct_path] = ct.encode()

    # Inject anchors into heading paragraphs
    page_idx = [0]
    def inject_image(match):
        idx = page_idx[0]
        if idx >= len(all_pages):
            return match.group(0)
        page_idx[0] += 1
        _, h_pt, section = all_pages[idx]
        # Section B images offset down so text starts above
        v_offset = int(2 * 914400) if section == "B" else 0  # 2 inches down for B
        anchor = make_anchor_xml(rid, IMAGE_WIDTH_PT, h_pt, idx, v_offset_emu=v_offset)
        return match.group(0).replace("</w:r></w:p>", anchor + "</w:r></w:p>", 1)

    doc_xml = re.sub(
        r'<w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Page \d+:.*?</w:t></w:r></w:p>',
        inject_image,
        doc_xml,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_files.items():
            if isinstance(data, str):
                data = data.encode()
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))
        zout.writestr("word/_rels/document.xml.rels", rels_xml.encode("utf-8"))
        zout.writestr(img_path, img_data)

    os.unlink(tmp)
    total = len(all_pages)
    print(f"Generated {OUT} ({total} pages: {len(HEIGHTS_A)} section A + {len(HEIGHTS_B)} section B)")


if __name__ == "__main__":
    main()
