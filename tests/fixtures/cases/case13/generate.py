"""Generate a ~200-page DOCX with headings and body paragraphs for case13."""

from docx import Document
from docx.shared import Pt, Inches

CHAPTERS = [
    "Introduction",
    "Background and Context",
    "Literature Review",
    "Methodology",
    "Data Collection",
    "Statistical Framework",
    "Preliminary Findings",
    "Core Analysis",
    "Regional Variations",
    "Temporal Patterns",
    "Comparative Assessment",
    "Economic Implications",
    "Environmental Factors",
    "Social Dimensions",
    "Policy Considerations",
    "Implementation Strategies",
    "Case Studies",
    "Stakeholder Perspectives",
    "Risk Analysis",
    "Future Projections",
    "Recommendations",
    "Conclusion",
    "Appendix A: Data Tables",
    "Appendix B: Supplementary Analysis",
    "Appendix C: Methodology Notes",
]

SECTIONS = [
    "Overview",
    "Key Findings",
    "Detailed Analysis",
    "Discussion",
]

PARAGRAPHS = [
    (
        "The analysis presented in this section draws upon a comprehensive dataset "
        "spanning multiple years of observation. Quantitative measures were triangulated "
        "with qualitative assessments to ensure robustness of the conclusions. The "
        "methodology follows established protocols widely accepted in the field, with "
        "several adaptations to account for the unique characteristics of the study "
        "population."
    ),
    (
        "Preliminary results indicate statistically significant trends across the "
        "primary outcome variables. Effect sizes ranged from moderate to large, with "
        "confidence intervals suggesting stable estimates. Subgroup analyses revealed "
        "heterogeneity in treatment response, warranting further investigation into "
        "moderating factors that may influence observed outcomes."
    ),
    (
        "Cross-referencing with existing literature confirms the general direction of "
        "our findings while highlighting several novel contributions. The divergence "
        "observed in secondary endpoints merits additional research, particularly "
        "regarding the interaction between demographic variables and contextual factors "
        "that were not fully captured in previous studies."
    ),
    (
        "Resource allocation patterns exhibited marked seasonal variation, with peak "
        "utilization occurring during the third and fourth quarters. Efficiency metrics "
        "improved across all operational categories, driven by process optimization "
        "initiatives implemented during the preceding fiscal year. These gains are "
        "expected to compound as the organization scales its operations."
    ),
    (
        "Stakeholder engagement surveys yielded response rates above the threshold "
        "required for statistical validity. Sentiment analysis of open-ended responses "
        "identified recurring themes around service quality, accessibility, and long-term "
        "sustainability. These qualitative insights complement the quantitative data and "
        "provide actionable direction for strategic planning."
    ),
    (
        "The financial model projects steady growth under baseline assumptions, with "
        "sensitivity analyses indicating resilience to moderate economic shocks. Capital "
        "expenditure requirements remain within previously approved budgetary envelopes, "
        "though contingency provisions may need to be revisited should certain risk "
        "scenarios materialize during the implementation phase."
    ),
    (
        "Comparative benchmarking against peer organizations reveals competitive "
        "positioning in most key performance indicators. Areas identified for improvement "
        "include digital transformation maturity and workforce development investment, "
        "both of which have been prioritized in the forthcoming strategic plan. External "
        "partnerships are being explored to accelerate progress in these domains."
    ),
    (
        "The regulatory environment continues to evolve, with several proposed changes "
        "currently under review. Impact assessments conducted by the compliance team "
        "suggest manageable adaptation costs under the most likely scenarios. Proactive "
        "engagement with regulatory bodies has positioned the organization favorably for "
        "the anticipated transition period."
    ),
]

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

para_idx = 0

for ch_num, chapter in enumerate(CHAPTERS):
    doc.add_heading(chapter, level=1)

    for sec_num, section_title in enumerate(SECTIONS):
        doc.add_heading(f"{section_title}", level=2)

        # 16-20 paragraphs per section, cycling through the pool
        n_paras = 16 + (ch_num + sec_num) % 5
        for _ in range(n_paras):
            doc.add_paragraph(PARAGRAPHS[para_idx % len(PARAGRAPHS)])
            para_idx += 1

doc.save("tests/fixtures/case13/input.docx")
print(f"Generated tests/fixtures/case13/input.docx ({para_idx} paragraphs)")
