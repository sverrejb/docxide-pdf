#!/usr/bin/env python3
"""case37: SmartArt with diverse shapes — inline diagram using dsp: namespace.

Tests the SmartArt-specific parsing path (word/diagrams/drawing1.xml with dsp:sp elements)
with multiple preset shape types. Exercises the geometry engine through the SmartArt code path
rather than the floating wps:wsp path tested in case34-36.

Layout: 2-row process diagram
  Row 1: Start(ellipse) → Research(roundRect) → Review(diamond) → Approve(pentagon) → Complete(ellipse)
  Row 2: Planning(hexagon), Execute(chevron), Report(trapezoid), Process(flowChartProcess)
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

OUT = Path("tests/fixtures/cases/case37/input.docx")

# XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DSP_NS = "http://schemas.microsoft.com/office/drawing/2008/diagram"
DGM_NS = "http://schemas.openxmlformats.org/drawingml/2006/diagram"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
REL_TYPE_DIAGRAM_DRAWING = (
    "http://schemas.microsoft.com/office/2007/relationships/diagramDrawing"
)
REL_TYPE_DIAGRAM_DATA = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramData"
)
REL_TYPE_DIAGRAM_LAYOUT = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramLayout"
)
REL_TYPE_DIAGRAM_STYLE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramStyle"
)
REL_TYPE_DIAGRAM_COLORS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramColors"
)

EMU_PER_INCH = 914400

# Diagram dimensions
DIAG_W = int(6.5 * EMU_PER_INCH)  # 5943600
DIAG_H = int(2.8 * EMU_PER_INCH)  # 2560320

# Row 1: Process flow shapes
SHAPE_W = 820000   # ~0.90"
SHAPE_H = 580000   # ~0.63"
ARROW_W = 170000   # ~0.19"
ARROW_H = 200000   # ~0.22"
GAP = 50000        # gap between shape edge and arrow

# Row 1 vertical center
ROW1_Y = 150000
ARROW_Y = ROW1_Y + (SHAPE_H - ARROW_H) // 2  # vertically centered

# Row 2
ROW2_Y = 1050000
ROW2_SHAPE_W = 880000
ROW2_SHAPE_H = 620000

# Compute row 1 x positions
def row1_positions():
    """Return list of (x, is_arrow) for row 1 elements."""
    positions = []
    x = 80000
    for i in range(5):
        positions.append((x, False))  # main shape
        x += SHAPE_W
        if i < 4:
            x += GAP
            positions.append((x, True))  # arrow connector
            x += ARROW_W + GAP
    return positions


ROW1 = row1_positions()

# Row 1 main shapes (index into ROW1 for non-arrow positions)
MAIN_SHAPES = [
    # (preset, color, label, adjustments_xml)
    ("ellipse",    "70AD47", "Start",    ""),
    ("roundRect",  "4472C4", "Research", '<a:gd name="adj" fmla="val 25000"/>'),
    ("diamond",    "ED7D31", "Review",   ""),
    ("pentagon",   "7030A0", "Approve",  ""),
    ("ellipse",    "C00000", "Complete", ""),
]

# Row 1 arrow connectors
ARROW_COLOR = "BFBFBF"

# Row 2 diverse shapes
ROW2_SHAPES = [
    ("hexagon",          "00B0F0", "Planning", ""),
    ("chevron",          "FFC000", "Execute",  ""),
    ("trapezoid",        "FF6699", "Report",   ""),
    ("flowChartProcess", "8FAADC", "Process",  ""),
]


def dsp_shape_xml(preset, x, y, cx, cy, fill_color, label, text_size, avlst_inner="",
                  stroke_color="404040", stroke_w=12700):
    """Build a dsp:sp element for one shape."""
    r, g, b = int(fill_color[0:2], 16), int(fill_color[2:4], 16), int(fill_color[4:6], 16)
    luma = 0.299 * r + 0.587 * g + 0.114 * b
    text_color = "FFFFFF" if luma < 140 else "000000"

    avlst = f"<a:avLst>{avlst_inner}</a:avLst>" if avlst_inner else "<a:avLst/>"

    text_xml = ""
    if label:
        text_xml = (
            f'<dsp:txBody>'
            f'<a:bodyPr anchor="ctr" lIns="36000" tIns="36000" rIns="36000" bIns="36000"/>'
            f'<a:p>'
            f'<a:pPr algn="ctr"/>'
            f'<a:r>'
            f'<a:rPr lang="en-US" sz="{text_size}" dirty="0">'
            f'<a:solidFill><a:srgbClr val="{text_color}"/></a:solidFill>'
            f'</a:rPr>'
            f'<a:t>{label}</a:t>'
            f'</a:r>'
            f'</a:p>'
            f'</dsp:txBody>'
        )

    return (
        f'<dsp:sp>'
        f'<dsp:spPr>'
        f'<a:xfrm>'
        f'<a:off x="{x}" y="{y}"/>'
        f'<a:ext cx="{cx}" cy="{cy}"/>'
        f'</a:xfrm>'
        f'<a:prstGeom prst="{preset}">'
        f'{avlst}'
        f'</a:prstGeom>'
        f'<a:solidFill>'
        f'<a:srgbClr val="{fill_color}"/>'
        f'</a:solidFill>'
        f'<a:ln w="{stroke_w}">'
        f'<a:solidFill>'
        f'<a:srgbClr val="{stroke_color}"/>'
        f'</a:solidFill>'
        f'</a:ln>'
        f'</dsp:spPr>'
        f'{text_xml}'
        f'</dsp:sp>'
    )


def build_drawing_xml():
    """Build the complete word/diagrams/drawing1.xml content."""
    shapes = []

    # Row 1: Main shapes and arrow connectors
    main_idx = 0
    for x, is_arrow in ROW1:
        if is_arrow:
            shapes.append(dsp_shape_xml(
                "rightArrow", x, ARROW_Y, ARROW_W, ARROW_H,
                ARROW_COLOR, "", 800,
                stroke_color="808080", stroke_w=6350,
            ))
        else:
            preset, color, label, adj = MAIN_SHAPES[main_idx]
            shapes.append(dsp_shape_xml(
                preset, x, ROW1_Y, SHAPE_W, SHAPE_H,
                color, label, 1000, avlst_inner=adj,
            ))
            main_idx += 1

    # Row 2: Diverse shapes evenly spaced
    spacing = DIAG_W // len(ROW2_SHAPES)
    for i, (preset, color, label, adj) in enumerate(ROW2_SHAPES):
        x = spacing * i + (spacing - ROW2_SHAPE_W) // 2
        shapes.append(dsp_shape_xml(
            preset, x, ROW2_Y, ROW2_SHAPE_W, ROW2_SHAPE_H,
            color, label, 1000, avlst_inner=adj,
        ))

    all_shapes = "\n".join(shapes)
    return (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<dsp:drawing xmlns:dsp="{DSP_NS}" xmlns:a="{A_NS}">'
        f'<dsp:spTree>'
        f'{all_shapes}'
        f'</dsp:spTree>'
        f'</dsp:drawing>'
    )


def build_inline_drawing_xml():
    """Build the wp:inline element that goes inside w:drawing in document.xml."""
    return (
        f'<w:r>'
        f'<w:drawing>'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{DIAG_W}" cy="{DIAG_H}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="1" name="Diagram 1"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="{A_NS}">'
        f'<a:graphicData uri="{DGM_NS}">'
        f'<dgm:relIds xmlns:dgm="{DGM_NS}" xmlns:r="{R_NS}"'
        f' r:dm="rIdDgmData" r:lo="rIdDgmLayout" r:qs="rIdDgmStyle" r:cs="rIdDgmColors"/>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
        f'</w:r>'
    )


def build_diagram_data_xml():
    """Minimal dgm:dataModel so Word can open the file."""
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<dgm:dataModel xmlns:dgm="{DGM_NS}">'
        '<dgm:ptLst>'
        '<dgm:pt modelId="0" type="doc"><dgm:prSet/><dgm:spPr/></dgm:pt>'
        '</dgm:ptLst>'
        '<dgm:cxnLst/>'
        '<dgm:bg/>'
        '<dgm:whole/>'
        '</dgm:dataModel>'
    )


def build_diagram_layout_xml():
    """Minimal layout definition."""
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<dgm:layoutDef xmlns:dgm="{DGM_NS}"/>'
    )


def build_diagram_style_xml():
    """Minimal quick style definition."""
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<dgm:styleDef xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' uniqueId="urn:microsoft.com/office/officeart/2005/8/quickstyle/simple1"/>'
    )


def build_diagram_colors_xml():
    """Minimal color transform definition."""
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<dgm:colorsDef xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' uniqueId="urn:microsoft.com/office/officeart/2005/8/colors/accent1_2"/>'
    )


# Step 1: Create base document with python-docx
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p = doc.add_paragraph()
run = p.add_run("SMARTART_PLACEHOLDER")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Step 2: Post-process ZIP
drawing_xml = build_drawing_xml()
inline_xml = build_inline_drawing_xml()

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()
        rels_xml = zin.read("word/_rels/document.xml.rels").decode()
        content_types_xml = zin.read("[Content_Types].xml").decode()

        # Add namespace declarations to document.xml
        ns_decls = (
            f' xmlns:wp="{WP_NS}"'
            f' xmlns:a="{A_NS}"'
            f' xmlns:r="{R_NS}"'
            f' xmlns:mc="{MC_NS}"'
        )
        if 'xmlns:wp=' not in doc_xml:
            doc_xml = doc_xml.replace(
                '<w:document ',
                f'<w:document {ns_decls} ',
                1,
            )

        # Replace placeholder with inline diagram drawing
        placeholder_pattern = r'<w:r>.*?<w:t>SMARTART_PLACEHOLDER</w:t>\s*</w:r>'
        doc_xml = re.sub(
            placeholder_pattern,
            inline_xml,
            doc_xml,
            count=1,
            flags=re.DOTALL,
        )

        # Add all diagram relationships to document.xml.rels
        diagram_rels = (
            f'<Relationship Id="rIdDiagram1"'
            f' Type="{REL_TYPE_DIAGRAM_DRAWING}"'
            f' Target="diagrams/drawing1.xml"/>'
            f'<Relationship Id="rIdDgmData"'
            f' Type="{REL_TYPE_DIAGRAM_DATA}"'
            f' Target="diagrams/data1.xml"/>'
            f'<Relationship Id="rIdDgmLayout"'
            f' Type="{REL_TYPE_DIAGRAM_LAYOUT}"'
            f' Target="diagrams/layout1.xml"/>'
            f'<Relationship Id="rIdDgmStyle"'
            f' Type="{REL_TYPE_DIAGRAM_STYLE}"'
            f' Target="diagrams/style1.xml"/>'
            f'<Relationship Id="rIdDgmColors"'
            f' Type="{REL_TYPE_DIAGRAM_COLORS}"'
            f' Target="diagrams/colors1.xml"/>'
        )
        rels_xml = rels_xml.replace(
            '</Relationships>',
            f'{diagram_rels}</Relationships>',
        )

        # Add content types for all diagram parts
        diagram_cts = (
            '<Override PartName="/word/diagrams/drawing1.xml"'
            ' ContentType="application/vnd.ms-office.drawingml.diagramDrawing+xml"/>'
            '<Override PartName="/word/diagrams/data1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramData+xml"/>'
            '<Override PartName="/word/diagrams/layout1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramLayoutDefinition+xml"/>'
            '<Override PartName="/word/diagrams/style1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramStyle+xml"/>'
            '<Override PartName="/word/diagrams/colors1.xml"'
            ' ContentType="application/vnd.openxmlformats-officedocument.drawingml.diagramColors+xml"/>'
        )
        content_types_xml = content_types_xml.replace(
            '</Types>',
            f'{diagram_cts}</Types>',
        )

        # Write all entries
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            elif item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, content_types_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

        # Add all diagram part files
        zout.writestr("word/diagrams/drawing1.xml", drawing_xml)
        zout.writestr("word/diagrams/data1.xml", build_diagram_data_xml())
        zout.writestr("word/diagrams/layout1.xml", build_diagram_layout_xml())
        zout.writestr("word/diagrams/style1.xml", build_diagram_style_xml())
        zout.writestr("word/diagrams/colors1.xml", build_diagram_colors_xml())

os.unlink(tmp)
print(f"Generated {OUT}")
