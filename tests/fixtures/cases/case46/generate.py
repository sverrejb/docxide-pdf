#!/usr/bin/env python3
"""case46: Floating table bottom gap diagnostic.

One table per page, varying both row count and row height.
Each table is followed by enough text to wrap past the table bottom.
All tables are left-aligned with identical bottomFromText="72" (3.6pt).

Pages:
 1: 1 row × 14pt    5: 2 rows × 14pt     9: 3 rows × 14pt
 2: 1 row × 20pt    6: 2 rows × 20pt    10: 3 rows × 20pt
 3: 1 row × 30pt    7: 2 rows × 30pt    11: 3 rows × 30pt
 4: 1 row × 40pt    8: 2 rows × 40pt    12: 3 rows × 40pt
13: 5 rows default  14: 8 rows default
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK

OUT = Path("tests/fixtures/cases/case46/input.docx")

LOREM = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
    "Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
    "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris "
    "nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in "
    "reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla "
    "pariatur. Excepteur sint occaecat cupidatat non proident, sunt in "
    "culpa qui officia deserunt mollit anim id est laborum. "
)

CONFIGS = [
    (1, 14), (1, 20), (1, 30), (1, 40),
    (2, 14), (2, 20), (2, 30), (2, 40),
    (3, 14), (3, 20), (3, 30), (3, 40),
    (5, None), (8, None),
]


def main():
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for i, (n_rows, row_h) in enumerate(CONFIGS):
        label = f"{n_rows}r"
        if row_h:
            label += f" × {row_h}pt"
        else:
            label += " × default"

        p = doc.add_paragraph()
        run = p.add_run(f"Page {i+1}: {label}")
        run.bold = True

        t = doc.add_table(rows=n_rows, cols=2)
        for ri in range(n_rows):
            t.rows[ri].cells[0].text = f"R{ri+1}"
            t.rows[ri].cells[1].text = f"V{ri+1}"

        # Enough text to wrap well past the table bottom
        doc.add_paragraph(LOREM * 4)

        # Page break after each (except the last)
        if i < len(CONFIGS) - 1:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    BORDERS = (
        '<w:tblBorders>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )

    TBL_POS = (
        '<w:tblpPr'
        ' w:vertAnchor="text"'
        ' w:horzAnchor="margin"'
        ' w:tblpX="0"'
        ' w:tblpY="200"'
        ' w:topFromText="72"'
        ' w:bottomFromText="72"'
        ' w:leftFromText="144"'
        ' w:rightFromText="144"/>'
    )

    TBL_WIDTH = '<w:tblW w:w="3600" w:type="dxa"/>'

    with zipfile.ZipFile(tmp, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode()
        other_files = {}
        for item in zin.infolist():
            if item.filename != "word/document.xml":
                other_files[item.filename] = zin.read(item.filename)

    # Inject tblpPr + borders + width
    config_idx = [0]
    def inject_tblpr(match):
        config_idx[0] += 1
        inner = match.group(0)
        inner = re.sub(r'<w:tblStyle w:val="[^"]*"/>', '', inner)
        inner = re.sub(r'<w:tblW[^/]*/>', TBL_WIDTH, inner)
        return inner.replace("</w:tblPr>", TBL_POS + BORDERS + "</w:tblPr>")

    doc_xml = re.sub(r'<w:tblPr>.*?</w:tblPr>', inject_tblpr, doc_xml, flags=re.DOTALL)

    # Inject row heights
    cfg_idx = [0]
    row_in_tbl = [0]

    def inject_row_height(match):
        if cfg_idx[0] >= len(CONFIGS):
            return match.group(0)
        n_rows, row_h = CONFIGS[cfg_idx[0]]
        row_in_tbl[0] += 1
        if row_in_tbl[0] > n_rows:
            cfg_idx[0] += 1
            row_in_tbl[0] = 1
            if cfg_idx[0] >= len(CONFIGS):
                return match.group(0)
            n_rows, row_h = CONFIGS[cfg_idx[0]]
        if row_h is not None:
            twips = int(row_h * 20)
            return match.group(0) + f'<w:trPr><w:trHeight w:val="{twips}" w:hRule="exact"/></w:trPr>'
        return match.group(0)

    doc_xml = re.sub(r'<w:tr>', inject_row_height, doc_xml)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_files.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp)
    print(f"Generated {OUT} ({config_idx[0]} tables, {len(CONFIGS)} pages)")


if __name__ == "__main__":
    main()
