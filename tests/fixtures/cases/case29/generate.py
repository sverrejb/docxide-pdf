"""Generate a DOCX with multiple chart variations for case29.

Charts are injected by post-processing the python-docx output:
1. Create the document with placeholder paragraphs using python-docx
2. Open the resulting ZIP and inject chart XML files + drawing references

Chart variations:
- Chart 1: Vertical clustered bar (3 series, 4 categories, legend right)
- Chart 2: Horizontal bar (2 series, 5 categories, legend bottom)
- Chart 3: Vertical bar with single series, no legend
- Chart 4: Vertical bar with large values (thousands), legend right
"""

import copy
import os
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("tests/fixtures/cases/case29/input.docx")

CHART_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
DML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CHART_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart"
CT_CHART = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"


def build_series_xml(idx, label, color_hex, categories, values):
    cat_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{c}</c:v></c:pt>' for i, c in enumerate(categories)
    )
    val_pts = "".join(
        f'<c:pt idx="{i}"><c:v>{v}</c:v></c:pt>' for i, v in enumerate(values)
    )
    return f"""<c:ser>
  <c:idx val="{idx}"/><c:order val="{idx}"/>
  <c:tx><c:strRef><c:f>label{idx}</c:f>
    <c:strCache><c:ptCount val="1"/><c:pt idx="0"><c:v>{label}</c:v></c:pt></c:strCache>
  </c:strRef></c:tx>
  <c:spPr><a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill><a:ln><a:noFill/></a:ln></c:spPr>
  <c:cat><c:strRef><c:f>cats</c:f>
    <c:strCache><c:ptCount val="{len(categories)}"/>{cat_pts}</c:strCache>
  </c:strRef></c:cat>
  <c:val><c:numRef><c:f>{idx}</c:f>
    <c:numCache><c:formatCode>General</c:formatCode><c:ptCount val="{len(values)}"/>{val_pts}</c:numCache>
  </c:numRef></c:val>
</c:ser>"""


def build_axis_xml(ax_id, cross_id, position, axis_type="cat", gridlines=False, line_color="b3b3b3"):
    grid_xml = ""
    if gridlines:
        grid_xml = f"""<c:majorGridlines><c:spPr><a:ln>
      <a:solidFill><a:srgbClr val="{line_color}"/></a:solidFill>
    </a:ln></c:spPr></c:majorGridlines>"""
    tag = "c:valAx" if axis_type == "val" else "c:catAx"
    return f"""<{tag}>
  <c:axId val="{ax_id}"/>
  <c:scaling><c:orientation val="minMax"/></c:scaling>
  <c:delete val="0"/>
  <c:axPos val="{position}"/>
  {grid_xml}
  <c:majorTickMark val="out"/><c:minorTickMark val="none"/>
  <c:tickLblPos val="nextTo"/>
  <c:spPr><a:ln><a:solidFill><a:srgbClr val="{line_color}"/></a:solidFill></a:ln></c:spPr>
  <c:crossAx val="{cross_id}"/>
  <c:crossesAt val="0"/>
</{tag}>"""


def build_chart_xml(
    bar_dir, grouping, series_list, categories, gap_width,
    legend_pos=None, colors=None
):
    default_colors = ["4472C4", "ED7D31", "A5A5A5", "FFC000", "5B9BD5", "70AD47"]
    if colors is None:
        colors = default_colors

    series_xml = ""
    for i, (label, values) in enumerate(series_list):
        color = colors[i % len(colors)]
        series_xml += build_series_xml(i, label, color, categories, values)

    cat_ax_id, val_ax_id = 100, 200
    if bar_dir == "bar":
        cat_pos, val_pos = "l", "b"
    else:
        cat_pos, val_pos = "b", "l"

    cat_ax = build_axis_xml(cat_ax_id, val_ax_id, cat_pos, axis_type="cat")
    val_ax = build_axis_xml(val_ax_id, cat_ax_id, val_pos, axis_type="val", gridlines=True)

    legend_xml = ""
    if legend_pos:
        legend_xml = f"""<c:legend>
  <c:legendPos val="{legend_pos}"/>
  <c:overlay val="0"/>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:legend>"""

    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{CHART_NS}" xmlns:a="{DML_NS}" xmlns:r="{REL_NS}">
  <c:lang val="en-US"/>
  <c:chart>
    <c:plotArea>
      <c:layout/>
      <c:barChart>
        <c:barDir val="{bar_dir}"/>
        <c:grouping val="{grouping}"/>
        {series_xml}
        <c:gapWidth val="{gap_width}"/>
        <c:axId val="{cat_ax_id}"/><c:axId val="{val_ax_id}"/>
      </c:barChart>
      {cat_ax}
      {val_ax}
      <c:spPr><a:noFill/><a:ln><a:solidFill><a:srgbClr val="b3b3b3"/></a:solidFill></a:ln></c:spPr>
    </c:plotArea>
    {legend_xml}
    <c:plotVisOnly val="1"/>
  </c:chart>
  <c:spPr><a:noFill/><a:ln><a:noFill/></a:ln></c:spPr>
</c:chartSpace>"""


def build_drawing_xml(rel_id, cx_emu, cy_emu):
    return (
        f'<w:drawing xmlns:w="{W_NS}" xmlns:wp="{WP_NS}" '
        f'xmlns:a="{DML_NS}" xmlns:c="{CHART_NS}" xmlns:r="{REL_NS}">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx_emu}" cy="{cy_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{rel_id.replace("rId","")}" name="Chart {rel_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic><a:graphicData uri="{CHART_NS}">'
        f'<c:chart r:id="{rel_id}"/>'
        f'</a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing>'
    )


# ── Step 1: Build document with python-docx ──

doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

doc.add_heading("Chart Test Cases", level=1)

# Placeholder paragraphs — we'll inject chart drawings into these
doc.add_paragraph("Chart 1: Vertical clustered bar chart (3 series, 4 categories)")
p1 = doc.add_paragraph()  # chart 1 goes here
p1.add_run("CHART_PLACEHOLDER_1")

doc.add_paragraph("")

doc.add_paragraph("Chart 2: Horizontal bar chart (2 series, 5 categories)")
p2 = doc.add_paragraph()
p2.add_run("CHART_PLACEHOLDER_2")

doc.add_paragraph("")

doc.add_paragraph("Chart 3: Single series vertical bar, no legend")
p3 = doc.add_paragraph()
p3.add_run("CHART_PLACEHOLDER_3")

doc.add_paragraph("")

doc.add_paragraph("Chart 4: Large values (thousands), 2 series")
p4 = doc.add_paragraph()
p4.add_run("CHART_PLACEHOLDER_4")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# ── Step 2: Post-process the ZIP to inject charts ──

charts = {
    1: build_chart_xml(
        bar_dir="col", grouping="clustered", gap_width=100,
        legend_pos="r",
        categories=["Q1", "Q2", "Q3", "Q4"],
        series_list=[
            ("Revenue", [12.5, 18.3, 15.7, 22.1]),
            ("Expenses", [10.2, 14.8, 13.1, 16.5]),
            ("Profit", [2.3, 3.5, 2.6, 5.6]),
        ],
    ),
    2: build_chart_xml(
        bar_dir="bar", grouping="clustered", gap_width=80,
        legend_pos="b",
        categories=["Engineering", "Marketing", "Sales", "Support", "Admin"],
        series_list=[
            ("2024", [45, 22, 38, 15, 12]),
            ("2025", [52, 28, 41, 18, 14]),
        ],
        colors=["2E75B6", "BF504D"],
    ),
    3: build_chart_xml(
        bar_dir="col", grouping="clustered", gap_width=150,
        legend_pos=None,
        categories=["Mon", "Tue", "Wed", "Thu", "Fri"],
        series_list=[
            ("Visitors", [340, 520, 480, 610, 390]),
        ],
        colors=["548235"],
    ),
    4: build_chart_xml(
        bar_dir="col", grouping="clustered", gap_width=100,
        legend_pos="r",
        categories=["North", "South", "East", "West"],
        series_list=[
            ("Budget", [4500, 3200, 5100, 2800]),
            ("Actual", [4200, 3800, 4700, 3100]),
        ],
        colors=["4472C4", "ED7D31"],
    ),
}

# Chart sizes in EMUs (1 inch = 914400 EMU, 1 pt = 12700 EMU)
chart_sizes = {
    1: (4572000, 2743200),  # ~5.0 x 3.0 inches
    2: (5486400, 2743200),  # ~6.0 x 3.0 inches
    3: (4572000, 2286000),  # ~5.0 x 2.5 inches
    4: (4572000, 2743200),  # ~5.0 x 3.0 inches
}

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        # Track which rIds are already used
        rels_xml = zin.read("word/_rels/document.xml.rels").decode()
        ct_xml = zin.read("[Content_Types].xml").decode()

        # Find highest existing rId
        import re
        existing_rids = [int(x) for x in re.findall(r'Id="rId(\d+)"', rels_xml)]
        next_rid = max(existing_rids, default=0) + 1

        chart_rids = {}
        for chart_num in sorted(charts.keys()):
            rid = f"rId{next_rid}"
            chart_rids[chart_num] = rid
            next_rid += 1

        # Patch rels: add chart relationships
        new_rels = ""
        for chart_num, rid in chart_rids.items():
            new_rels += (
                f'<Relationship Id="{rid}" '
                f'Type="{CHART_REL_TYPE}" '
                f'Target="charts/chart{chart_num}.xml"/>'
            )
        rels_xml = rels_xml.replace("</Relationships>", new_rels + "</Relationships>")

        # Patch content types: add chart content types
        new_ct = ""
        for chart_num in charts:
            new_ct += (
                f'<Override PartName="/word/charts/chart{chart_num}.xml" '
                f'ContentType="{CT_CHART}"/>'
            )
        ct_xml = ct_xml.replace("</Types>", new_ct + "</Types>")

        # Patch document.xml: replace placeholders with drawing elements
        doc_xml = zin.read("word/document.xml").decode()
        for chart_num, rid in chart_rids.items():
            placeholder = f"CHART_PLACEHOLDER_{chart_num}"
            cx, cy = chart_sizes[chart_num]
            drawing = build_drawing_xml(rid, cx, cy)
            # Replace the run containing the placeholder with one containing the drawing
            run_pattern = (
                f'<w:r><w:rPr></w:rPr><w:t>{placeholder}</w:t></w:r>'
            )
            run_replacement = f'<w:r>{drawing}</w:r>'
            if run_pattern in doc_xml:
                doc_xml = doc_xml.replace(run_pattern, run_replacement)
            else:
                # Try without rPr
                run_pattern2 = f'<w:r><w:t>{placeholder}</w:t></w:r>'
                if run_pattern2 in doc_xml:
                    doc_xml = doc_xml.replace(run_pattern2, run_replacement)
                else:
                    # Broader regex replacement
                    doc_xml = re.sub(
                        rf'<w:r[^>]*>.*?<w:t[^>]*>{placeholder}</w:t>.*?</w:r>',
                        run_replacement,
                        doc_xml,
                        flags=re.DOTALL,
                    )

        # Write all entries
        for item in zin.infolist():
            if item.filename == "word/_rels/document.xml.rels":
                zout.writestr(item, rels_xml)
            elif item.filename == "[Content_Types].xml":
                zout.writestr(item, ct_xml)
            elif item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

        # Add chart XML files
        for chart_num, chart_xml in charts.items():
            zout.writestr(f"word/charts/chart{chart_num}.xml", chart_xml)

os.unlink(tmp)
print(f"Generated {OUT}")
