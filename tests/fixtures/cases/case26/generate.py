"""Generate case26: Mixed page sizes and orientations in one document

Tests:
- Sections with different page sizes within the same document
- Portrait-to-landscape-to-portrait transitions
- Different margins per section
- Content that verifies correct text width per section
"""

from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document()

# Section 1: Standard US Letter portrait with narrow margins
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.orientation = WD_ORIENT.PORTRAIT
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

doc.add_heading("Section 1: Letter Portrait, Narrow Margins", level=1)
doc.add_paragraph(
    "This is a standard US Letter page (8.5 x 11 inches) with narrow margins "
    "(0.75 inch left/right, 0.5 inch top/bottom). The text area is 7 inches "
    "wide. This first section establishes the baseline portrait layout."
)
doc.add_paragraph(
    "A ruler line for visual reference: "
    + "|" + "-" * 40 + "|" + "-" * 40 + "|"
)
doc.add_paragraph(
    "The quick brown fox jumps over the lazy dog. Pack my box with five dozen "
    "liquor jugs. How vexingly quick daft zebras jump. The five boxing wizards "
    "jump quickly."
)

# Section 2: Landscape Letter with wide margins
section = doc.add_section()
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.orientation = WD_ORIENT.LANDSCAPE
section.top_margin = Inches(1.5)
section.bottom_margin = Inches(1.5)
section.left_margin = Inches(1.5)
section.right_margin = Inches(1.5)

doc.add_heading("Section 2: Letter Landscape, Wide Margins", level=1)
doc.add_paragraph(
    "This section switches to landscape orientation with wide margins "
    "(1.5 inches all around). The text area is 8 inches wide by 5.5 inches "
    "tall. This tests the transition from portrait to landscape within the "
    "same document."
)
doc.add_paragraph(
    "More content to fill the landscape page. The quick brown fox jumps over "
    "the lazy dog. Pack my box with five dozen liquor jugs. How vexingly quick "
    "daft zebras jump."
)

# Section 3: A4 portrait
section = doc.add_section()
section.page_width = Mm(210)
section.page_height = Mm(297)
section.orientation = WD_ORIENT.PORTRAIT
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

doc.add_heading("Section 3: A4 Portrait", level=1)
doc.add_paragraph(
    "Now we switch to A4 portrait (210 x 297mm). A4 is the international "
    "standard paper size. It is slightly narrower than US Letter (about 0.3 "
    "inches narrower) but taller (about 0.7 inches taller). The text width "
    "should be noticeably different from the previous sections."
)

# Section 4: Back to Letter portrait with default margins
section = doc.add_section()
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.orientation = WD_ORIENT.PORTRAIT
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

doc.add_heading("Section 4: Back to Letter Portrait", level=1)
doc.add_paragraph(
    "Final section returns to US Letter portrait with standard 1-inch margins. "
    "This verifies that switching back from A4 works correctly and the page "
    "dimensions reset properly."
)
doc.add_paragraph(
    "The quick brown fox jumps over the lazy dog. Pack my box with five dozen "
    "liquor jugs. How vexingly quick daft zebras jump. The five boxing wizards "
    "jump quickly. Sphinx of black quartz, judge my vow."
)

# Save with compat mode 15
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(
    out_buf, "w", zipfile.ZIP_DEFLATED
) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/settings.xml":
            tree = etree.fromstring(data)
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(
                tree, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
