#!/usr/bin/env python3
"""Generate test fixture for hyphenation and Unicode line breaking.

Exercises:
- w:autoHyphenation in document settings
- Long English words that need hyphenation
- Justified text (shows hyphenation quality)
- w:suppressAutoHyphens on one paragraph
- German compound words with w:lang="de-DE"
- Narrow margins to force frequent line breaks

Usage:
    uv run tests/fixtures/cases/case48/generate.py
"""

import zipfile
import shutil
import tempfile
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path("tests/fixtures/cases/case48")


def generate():
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.75)
        section.right_margin = Inches(1.75)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hyphenation Test")
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    run.bold = True

    # Body paragraph with long words (justified)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        "The concept of antidisestablishmentarianism has been discussed extensively "
        "in philosophical and political literature. Similarly, the word "
        "supercalifragilisticexpialidocious has captured the imagination of generations. "
        "Modern electroencephalographically monitored studies have shown that "
        "incomprehensibilities arise from miscommunication. The characteristically "
        "disproportionate representation of psychopharmacological interventions "
        "demonstrates the unconstitutionality of such practices."
    )
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Second body paragraph with more natural text
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        "Professional documentation requires careful consideration of typographical "
        "conventions. Hyphenation improves the visual appearance of justified text by "
        "reducing excessive inter-word spacing. Without hyphenation, lines with long "
        "words may have uncomfortably large gaps between words, or words may overflow "
        "the right margin. This is particularly noticeable in narrow columns."
    )
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Paragraph with suppressAutoHyphens (for comparison)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        "This paragraph has automatic hyphenation suppressed. The word "
        "internationalization should not be hyphenated here even though it would "
        "benefit from breaking. Compare the inter-word spacing with the paragraphs "
        "above to see the difference that hyphenation makes in justified text."
    )
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # German paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text = (
        "Die Donaudampfschifffahrtsgesellschaftskapitaenswitwe ging zum "
        "Rindfleischetikettierungsueberwachungsaufgabenuebertragungsgesetz. "
        "Die Geschwindigkeitsbeschraenkung auf der Bundesautobahn wurde "
        "aufgrund von Verkehrssicherheitsbedenken eingefuehrt."
    )
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Save initial docx
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # ZIP post-processing: inject autoHyphenation + suppressAutoHyphens + lang
    out_path = OUT_DIR / "input.docx"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(buf.read())
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        with zipfile.ZipFile(str(out_path), "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == "word/settings.xml":
                    # Inject w:autoHyphenation before </w:settings>
                    text_data = data.decode("utf-8")
                    text_data = text_data.replace(
                        "</w:settings>",
                        '<w:autoHyphenation/></w:settings>',
                    )
                    data = text_data.encode("utf-8")

                elif item.filename == "word/document.xml":
                    text_data = data.decode("utf-8")

                    # Find the 3rd w:pPr (paragraph with suppressAutoHyphens)
                    # Count paragraph occurrences to find the right one
                    # The 4th <w:p  (0-indexed: paragraph 3) should get suppressAutoHyphens
                    ppr_count = 0
                    result = []
                    i = 0
                    while i < len(text_data):
                        if text_data[i:].startswith("<w:pPr"):
                            ppr_count += 1
                            # Find the end of this pPr opening tag
                            end = text_data.index(">", i)
                            if text_data[end - 1] == "/":
                                # Self-closing <w:pPr/>
                                if ppr_count == 4:
                                    result.append(
                                        "<w:pPr><w:suppressAutoHyphens/></w:pPr>"
                                    )
                                    i = end + 1
                                    continue
                                else:
                                    result.append(text_data[i : end + 1])
                                    i = end + 1
                                    continue
                            else:
                                result.append(text_data[i : end + 1])
                                if ppr_count == 4:
                                    result.append("<w:suppressAutoHyphens/>")
                                i = end + 1
                                continue
                        result.append(text_data[i])
                        i += 1
                    text_data = "".join(result)

                    # Add w:lang to the last paragraph's runs (German)
                    # Find the last <w:rPr> that contains Times New Roman
                    # and add w:lang before the closing </w:rPr>
                    # Strategy: find the last paragraph's runs and add lang
                    last_p_start = text_data.rfind("<w:p ")
                    if last_p_start == -1:
                        last_p_start = text_data.rfind("<w:p>")

                    # Within the last paragraph, add w:lang to all rPr elements
                    before = text_data[:last_p_start]
                    after = text_data[last_p_start:]
                    after = after.replace(
                        "</w:rPr>",
                        '<w:lang w:val="de-DE"/></w:rPr>',
                    )
                    text_data = before + after

                    data = text_data.encode("utf-8")

                zout.writestr(item, data)

    Path(tmp_path).unlink()
    print(f"Generated {out_path}")


if __name__ == "__main__":
    generate()
