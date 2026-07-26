"""Generate case77: distributed and kashida paragraph alignment (w:jc, §17.18.44)

Isolates the ST_Jc values that are *not* plain left/center/right/both:

  - distribute      "Distribute All Characters Equally" — like justify, but the
                    last line is stretched too and the slack goes between
                    CHARACTERS rather than only between word gaps.
  - thaiDistribute  Thai flavor of the same; we render it identically.
  - mediumKashida   Arabic kashida justification. Real kashida elongates glyphs,
                    which needs Arabic shaping; with Latin text Word should fall
                    back to ordinary justification, which is what we render.

Paragraph 1 is the control: the same text as paragraph 2 but `both`, so the
reference shows the one difference that matters — a justified paragraph leaves
its last line short, a distributed one stretches it to the right margin.

Paragraph 3 is the single-short-line case, where distribute is most visible:
a two-word line spreads edge to edge across the full column.

Paragraph 6 covers the CJK branch. Our Tc divisor differs by one gap between
CJK justify (keeps the grid's trailing cell gap) and distribute (ends flush at
the margin), so a CJK distributed line pins that rule against Word.

No eastAsia font is forced — Word picks its own CJK fallback, and pinning a
Windows-only face (MS Mincho) would just diverge on a Mac reference.
"""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case77/input.docx")

BODY_FONT = "Arial"
BODY_PT = 12
LABEL_PT = 9

LONG = (
    "The quick brown fox jumps over the lazy dog while the industrious beaver "
    "constructs an elaborate dam across the narrow stream near the old mill. "
    "Short tail."
)

# (jc value, label, body text)
SAMPLES = [
    ("both", "1. both (justify) — control: last line stays short", LONG),
    ("distribute", "2. distribute — every line stretched, last one included", LONG),
    ("distribute", "3. distribute — single short line spreads edge to edge", "Spread me"),
    ("thaiDistribute", "4. thaiDistribute — same treatment as distribute", "Thai distribute"),
    ("mediumKashida", "5. mediumKashida on Latin text — plain justify", LONG),
    ("distribute", "6. distribute — CJK characters", "日本語の均等割り付け"),
]


def set_jc(paragraph, val):
    """Set w:jc directly so the exact ST_Jc token lands in the XML."""
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:jc")):
        pPr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), val)
    pPr.append(jc)


def add_para(doc, text, size_pt, jc=None, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    # w:rFonts @ascii alone leaves CJK to the eastAsia slot; copy the ascii face
    # to cs so Word doesn't fall back to a different face for the Latin samples.
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        rFonts.set(qn("w:cs"), BODY_FONT)
    if jc:
        set_jc(p, jc)
    return p


def main():
    doc = Document()

    # Narrow the column a little so the distributed lines have obvious slack.
    section = doc.sections[0]
    section.left_margin = section.right_margin = Pt(90)

    add_para(doc, "Distributed & Kashida Alignment (w:jc)", 14, bold=True)

    for jc, label, text in SAMPLES:
        add_para(doc, label, LABEL_PT)
        add_para(doc, text, BODY_PT, jc=jc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
