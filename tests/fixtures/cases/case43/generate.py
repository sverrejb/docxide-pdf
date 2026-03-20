#!/usr/bin/env python3
"""case43: WordArt — modern DrawingML WordArt with text effects and warping.

Tests WordArt rendering at multiple levels:
1. Flat WordArt with text outline (textPlain warp = no distortion)
2. WordArt with text shadow
3. WordArt with wave warp (textWave1)
4. WordArt with arch warp (textArchUp)
5. WordArt with slant (textSlantUp)
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT = Path("tests/fixtures/cases/case43/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WP14_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
GRAPHIC_DATA_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

EMU_PER_PT = 12700
EMU_PER_INCH = 914400


def wordart_xml(
    text,
    warp_preset,
    y_offset_emu,
    color="4472C4",
    font_size_hps=72,
    outline_color=None,
    outline_width_emu=9525,
    shadow=False,
    adj_values=None,
    width_emu=int(5.5 * EMU_PER_INCH),
    height_emu=int(1.0 * EMU_PER_INCH),
):
    """Generate mc:AlternateContent XML for a WordArt textbox."""

    font_size_pt = font_size_hps // 2

    # w14:textOutline (xmlns:w14 declared on document root, not here)
    outline_xml = ""
    if outline_color:
        outline_xml = (
            f'<w14:textOutline w14:w="{outline_width_emu}" w14:cap="flat">'
            f'<w14:solidFill><w14:srgbClr w14:val="{outline_color}"/></w14:solidFill>'
            f'</w14:textOutline>'
        )

    # w14:shadow
    shadow_xml = ""
    if shadow:
        shadow_xml = (
            f'<w14:shadow w14:blurRad="38100" w14:dist="25400" w14:dir="3600000">'
            f'<w14:srgbClr w14:val="808080">'
            f'<w14:alpha w14:val="60000"/>'
            f'</w14:srgbClr>'
            f'</w14:shadow>'
        )

    # Adjustment values
    av_xml = "<a:avLst/>"
    if adj_values:
        gds = "".join(
            f'<a:gd name="{n}" fmla="val {v}"/>' for n, v in adj_values
        )
        av_xml = f"<a:avLst>{gds}</a:avLst>"

    return f"""<mc:AlternateContent xmlns:mc="{MC_NS}">
  <mc:Choice Requires="wps">
    <w:drawing xmlns:w="{W_NS}">
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
          simplePos="0" relativeHeight="251658240" behindDoc="0"
          locked="0" layoutInCell="1" allowOverlap="1"
          xmlns:wp="{WP_NS}" xmlns:wp14="{WP14_NS}">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column"><wp:align>center</wp:align></wp:positionH>
        <wp:positionV relativeFrom="paragraph"><wp:posOffset>{y_offset_emu}</wp:posOffset></wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{abs(hash(text)) % 100000}" name="WordArt {text[:10]}"/>
        <a:graphic xmlns:a="{A_NS}">
          <a:graphicData uri="{GRAPHIC_DATA_URI}">
            <wps:wsp xmlns:wps="{WPS_NS}">
              <wps:cNvSpPr txBox="1"/>
              <wps:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:noFill/>
                <a:ln><a:noFill/></a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent xmlns:w="{W_NS}">
                  <w:p>
                    <w:pPr><w:jc w:val="center"/></w:pPr>
                    <w:r>
                      <w:rPr>
                        <w:b/>
                        <w:color w:val="{color}"/>
                        <w:sz w:val="{font_size_hps}"/>
                        <w:szCs w:val="{font_size_hps}"/>
                        {outline_xml}
                        {shadow_xml}
                      </w:rPr>
                      <w:t>{text}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr fromWordArt="1" wrap="none" lIns="0" tIns="0" rIns="0" bIns="0">
                <a:prstTxWarp prst="{warp_preset}">{av_xml}</a:prstTxWarp>
                <a:spAutoFit/>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </mc:Choice>
</mc:AlternateContent>"""


def main():
    doc = Document()

    # Title paragraph
    title = doc.add_paragraph()
    run = title.add_run("WordArt Test Cases")
    run.bold = True
    run.font.size = Pt(16)

    # Add spacing paragraphs and WordArt anchors
    wordarts = [
        # (text, preset, y_offset_inches, color, outline, shadow, adj)
        ("FLAT OUTLINE", "textPlain", 0.8, "2E75B6", "1F4E79", False, None),
        ("SHADOW TEXT", "textPlain", 2.2, "C00000", None, True, None),
        ("WAVE EFFECT", "textWave1", 3.8, "ED7D31", "C44F1A", False, None),
        ("ARCH TEXT", "textArchUp", 5.5, "70AD47", None, False, [("adj", "10800000")]),
        ("SLANT UP", "textSlantUp", 8.0, "7030A0", "4B0082", False, None),
    ]

    # Add enough blank paragraphs for spacing
    for _ in range(20):
        doc.add_paragraph("")

    # Save base document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    # Post-process: inject WordArt XML into the first paragraph
    with zipfile.ZipFile(tmp_path, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode("utf-8")
        other_files = {}
        for name in zin.namelist():
            if name != "word/document.xml":
                other_files[name] = zin.read(name)

    # Add w14 namespace to document root if not present
    if f'xmlns:w14="{W14_NS}"' not in doc_xml:
        doc_xml = doc_xml.replace(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xmlns:w14="{W14_NS}"',
        )
    # Add mc:Ignorable for w14 if not already listed
    if "mc:Ignorable" in doc_xml and "w14" not in re.search(r'mc:Ignorable="([^"]*)"', doc_xml).group(1):
        doc_xml = re.sub(
            r'mc:Ignorable="([^"]*)"',
            r'mc:Ignorable="\1 w14"',
            doc_xml,
        )

    # Inject WordArt into the first <w:p> (title paragraph)
    all_wordart = ""
    for text, preset, y_inches, color, outline, shadow, adj in wordarts:
        y_emu = int(y_inches * EMU_PER_INCH)
        all_wordart += wordart_xml(
            text=text,
            warp_preset=preset,
            y_offset_emu=y_emu,
            color=color,
            outline_color=outline,
            shadow=shadow,
            adj_values=adj,
        )

    # Insert all WordArt before the first </w:p>
    doc_xml = doc_xml.replace("</w:p>", all_wordart + "</w:p>", 1)

    # Write output
    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_files.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
