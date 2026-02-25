#!/usr/bin/env python3
"""Generate case14 test fixture: clickable hyperlinks."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


def add_hyperlink(paragraph, url, text, font_name="Aptos", font_size=Pt(12)):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    return run


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(12)

    # Hyperlink style
    if "Hyperlink" not in [s.name for s in doc.styles]:
        hl_style = doc.styles.add_style("Hyperlink", 2)  # character style
        hl_style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        hl_style.font.underline = True

    # 1. Simple hyperlink
    p = doc.add_paragraph("This paragraph has a ")
    add_hyperlink(p, "https://example.com", "simple hyperlink")
    r = p.add_run(" in the middle of it.")

    # 2. Multiple hyperlinks in one paragraph
    p = doc.add_paragraph("Visit ")
    add_hyperlink(p, "https://example.com", "Example")
    p.add_run(" or ")
    add_hyperlink(p, "https://example.org", "Example Org")
    p.add_run(" for more information.")

    # 3. Standalone hyperlink
    p = doc.add_paragraph()
    add_hyperlink(p, "https://www.rust-lang.org", "https://www.rust-lang.org")

    # 4. Normal text after hyperlinks
    doc.add_paragraph(
        "This is a normal paragraph with no hyperlinks, to verify spacing is unaffected."
    )

    out = Path(__file__).parent / "input.docx"
    doc.save(str(out))
    print(f"Saved {out}")
    print(f"\nOpen in Word, then File > Save As > PDF to create reference.pdf")


if __name__ == "__main__":
    main()
