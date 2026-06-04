"""Generate case65: run borders (w:rPr/w:bdr).

Exercises the run-border feature added in PR #1:
- single border, default 0.5pt black
- colored borders (red, blue)
- varying widths via w:sz (eighths of a point)
- varying spacing via w:space (points of padding around the text)
- adjacent runs sharing an identical border merge into one box
- adjacent runs with differing borders draw separate boxes
- a run border inherited from a character style
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_bordered_run(paragraph, text, *, color=None, sz=None, space=None, bold=False):
    """Add a run carrying a w:rPr/w:bdr element.

    sz is in eighths of a point (Word's unit); space is in points.
    """
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    bdr = OxmlElement("w:bdr")
    bdr.set(qn("w:val"), "single")
    if sz is not None:
        bdr.set(qn("w:sz"), str(sz))
    if space is not None:
        bdr.set(qn("w:space"), str(space))
    if color is not None:
        bdr.set(qn("w:color"), color)
    rpr.append(bdr)
    return run


doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.orientation = None
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(12)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(10)

# Define a character style that carries a run border, to exercise the
# style-inheritance path (CharacterStyle.border in styles.rs).
boxed = doc.styles.add_style("Boxed", WD_STYLE_TYPE.CHARACTER)
boxed.font.name = "Arial"
boxed_rpr = boxed.element.get_or_add_rPr()
style_bdr = OxmlElement("w:bdr")
style_bdr.set(qn("w:val"), "single")
style_bdr.set(qn("w:sz"), "12")  # 1.5pt
style_bdr.set(qn("w:space"), "2")
style_bdr.set(qn("w:color"), "008000")  # green
boxed_rpr.append(style_bdr)

title = doc.add_heading("Run Borders", level=1)

# 1. Default border (no sz/space/color -> 0.5pt black, no padding).
p = doc.add_paragraph()
p.add_run("Default: ").font.name = "Arial"
add_bordered_run(p, "0.5pt black border")

# 2. Colored borders.
p = doc.add_paragraph()
p.add_run("Colors: ").font.name = "Arial"
add_bordered_run(p, "red", color="FF0000")
p.add_run("  ").font.name = "Arial"
add_bordered_run(p, "blue", color="0000FF")

# 3. Varying widths (sz is eighths of a point: 4=0.5pt, 16=2pt, 32=4pt).
p = doc.add_paragraph()
p.add_run("Widths: ").font.name = "Arial"
add_bordered_run(p, "0.5pt", sz=4)
p.add_run("  ").font.name = "Arial"
add_bordered_run(p, "2pt", sz=16)
p.add_run("  ").font.name = "Arial"
add_bordered_run(p, "4pt", sz=32)

# 4. Varying padding (space in points) around the boxed text.
p = doc.add_paragraph()
p.add_run("Padding: ").font.name = "Arial"
add_bordered_run(p, "0pt", space=0, color="800080")
p.add_run("  ").font.name = "Arial"
add_bordered_run(p, "2pt", space=2, color="800080")
p.add_run("  ").font.name = "Arial"
add_bordered_run(p, "5pt", space=5, color="800080")

# 5. Adjacent identical borders merge into a single box; a differing
#    border immediately after starts a new box.
p = doc.add_paragraph()
p.add_run("Merging: ").font.name = "Arial"
add_bordered_run(p, "these ", color="FF0000", sz=8)
add_bordered_run(p, "three ", color="FF0000", sz=8)
add_bordered_run(p, "merge", color="FF0000", sz=8)
add_bordered_run(p, " but this is separate", color="0000FF", sz=8)

# 6. Border applied via a character style.
p = doc.add_paragraph()
p.add_run("Style: ").font.name = "Arial"
p.add_run("green box from char style", style="Boxed")

out = "tests/fixtures/cases/case65/input.docx"
doc.save(out)
print(f"Created {out}")
