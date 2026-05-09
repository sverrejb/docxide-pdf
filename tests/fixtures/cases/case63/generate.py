#!/usr/bin/env python3
"""Generate case63 test fixture: Word comments.

DOCX comments (`w:comment` in word/comments.xml, referenced via
`w:commentRangeStart`/`w:commentRangeEnd`/`w:commentReference` in document.xml)
are exported by Word's PDF export as PDF Text (sticky-note) annotations with
/Contents set to the comment text. Acrobat shows a note icon that reveals the
comment on hover or click.

python-docx has no native comment support, so this script:
  1. Uses python-docx to build the base document and insert the comment
     range markers + commentReference runs directly into paragraph XML.
  2. Post-processes the saved .docx (a ZIP) to add word/comments.xml, the
     relationship entry, and the content-type override.
"""
import shutil
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COMMENTS_XML_TMPL = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
{comments}
</w:comments>
"""


def _comment_xml(cid, author, initials, date, text):
    return (
        f'  <w:comment w:id="{cid}" w:author="{author}" '
        f'w:initials="{initials}" w:date="{date}">\n'
        f"    <w:p><w:r><w:t>{text}</w:t></w:r></w:p>\n"
        f"  </w:comment>"
    )


def add_commented_run(paragraph, text, comment_id):
    p = paragraph._p

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    p.append(start)

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    p.append(run)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    p.append(end)

    ref_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "CommentReference")
    rPr.append(rStyle)
    ref_run.append(rPr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)
    p.append(ref_run)


def inject_comments_part(docx_path, comments):
    tmp_path = docx_path.with_suffix(".docx.tmp")

    comments_xml = COMMENTS_XML_TMPL.format(
        comments="\n".join(
            _comment_xml(c["id"], c["author"], c["initials"], c["date"], c["text"])
            for c in comments
        )
    )

    comments_rel_type = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    )
    comments_content_type = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.comments+xml"
    )

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp_path, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                data = data.replace(
                    b"</Relationships>",
                    (
                        f'<Relationship Id="rIdComments" Type="{comments_rel_type}" '
                        f'Target="comments.xml"/></Relationships>'
                    ).encode("utf-8"),
                )
            elif item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"</Types>",
                    (
                        f'<Override PartName="/word/comments.xml" '
                        f'ContentType="{comments_content_type}"/></Types>'
                    ).encode("utf-8"),
                )
            zout.writestr(item, data)
        zout.writestr("word/comments.xml", comments_xml)

    shutil.move(tmp_path, docx_path)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(12)

    doc.add_paragraph(
        "Word comments export as PDF Text (sticky-note) annotations. "
        "Open reference.pdf in Acrobat to see the note icons."
    )

    comments = [
        {
            "id": 1,
            "author": "Reviewer",
            "initials": "R",
            "date": "2024-01-01T00:00:00Z",
            "text": "This phrase has a comment attached to it.",
        },
        {
            "id": 2,
            "author": "Reviewer",
            "initials": "R",
            "date": "2024-01-01T00:00:00Z",
            "text": "Second comment — with non-ASCII: Besøk ✓",
        },
        {
            "id": 3,
            "author": "Reviewer",
            "initials": "R",
            "date": "2024-01-01T00:00:00Z",
            "text": (
                "A longer comment to check that multi-sentence notes survive "
                "Word's PDF export and render in the annotation pane."
            ),
        },
    ]

    p = doc.add_paragraph("1. ")
    add_commented_run(p, "This phrase", comment_id=1)
    p.add_run(" has a short comment attached.")

    p = doc.add_paragraph("2. ")
    add_commented_run(p, "Another phrase", comment_id=2)
    p.add_run(" carries a comment with unicode.")

    p = doc.add_paragraph("3. ")
    add_commented_run(p, "This one", comment_id=3)
    p.add_run(" has a longer multi-sentence comment.")

    out = Path(__file__).parent / "input.docx"
    doc.save(str(out))
    inject_comments_part(out, comments)

    print(f"Saved {out}")
    print("Next: open in Word, File > Save As > PDF to create reference.pdf")


if __name__ == "__main__":
    main()
