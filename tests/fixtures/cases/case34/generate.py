#!/usr/bin/env python3
"""case34: Preset shape gallery — grid of 20 common preset shapes with solid fills.

Tests geometry engine rendering of diverse preset shapes via a:prstGeom.
Each shape is a floating wps:wsp anchor with a:solidFill and shape label text.
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT = Path("tests/fixtures/cases/case34/input.docx")

# XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WP14_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
GRAPHIC_DATA_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

# 20 preset shapes with distinct fill colors
SHAPES = [
    ("roundRect",           "4472C4", "Round Rect"),
    ("diamond",             "ED7D31", "Diamond"),
    ("chevron",             "A5A5A5", "Chevron"),
    ("pentagon",            "FFC000", "Pentagon"),
    ("hexagon",             "5B9BD5", "Hexagon"),
    ("triangle",            "70AD47", "Triangle"),
    ("rightArrow",          "264478", "Right Arrow"),
    ("heart",               "C00000", "Heart"),
    ("star5",               "BF8F00", "Star 5"),
    ("cloud",               "9DC3E6", "Cloud"),
    ("donut",               "F4B183", "Donut"),
    ("flowChartProcess",    "8FAADC", "Process"),
    ("flowChartDecision",   "C55A11", "Decision"),
    ("flowChartTerminator", "548235", "Terminator"),
    ("plus",                "7030A0", "Plus"),
    ("trapezoid",           "FF6699", "Trapezoid"),
    ("parallelogram",       "00B0F0", "Parallelogram"),
    ("frame",               "002060", "Frame"),
    ("leftArrow",           "92D050", "Left Arrow"),
    ("upDownArrow",         "FF0000", "Up Down Arrow"),
]

# Layout constants (in EMUs: 1 inch = 914400 EMU, 1 pt = 12700 EMU)
EMU_PER_PT = 12700
EMU_PER_INCH = 914400

COLS = 4
ROWS = 5
SHAPE_W_EMU = int(1.3 * EMU_PER_INCH)   # 1.3 inches wide
SHAPE_H_EMU = int(1.0 * EMU_PER_INCH)   # 1.0 inch tall

# Content area: 6.5" wide, 9" tall (letter page with 1" margins)
CONTENT_W = 6.5 * EMU_PER_INCH
CONTENT_H = 9.0 * EMU_PER_INCH
COL_PITCH = CONTENT_W / COLS
ROW_PITCH = CONTENT_H / ROWS


def shape_anchor_xml(idx, preset, color, label):
    """Build wp:anchor XML for a single floating shape."""
    row = idx // COLS
    col = idx % COLS
    # Center each shape in its grid cell
    x_emu = int(col * COL_PITCH + (COL_PITCH - SHAPE_W_EMU) / 2)
    y_emu = int(row * ROW_PITCH + (ROW_PITCH - SHAPE_H_EMU) / 2)

    # Contrasting text color: white for dark fills, black for light fills
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    luma = 0.299 * r + 0.587 * g + 0.114 * b
    text_color = "FFFFFF" if luma < 140 else "000000"

    return (
        f'<w:r>'
        f'<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f' simplePos="0" relativeHeight="{idx}" behindDoc="0"'
        f' locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="margin">'
        f'<wp:posOffset>{x_emu}</wp:posOffset>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="margin">'
        f'<wp:posOffset>{y_emu}</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{idx + 1}" name="Shape {idx + 1}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="{A_NS}">'
        f'<a:graphicData uri="{GRAPHIC_DATA_URI}">'
        f'<wps:wsp xmlns:wps="{WPS_NS}">'
        f'<wps:cNvSpPr/>'
        f'<wps:spPr>'
        f'<a:xfrm>'
        f'<a:off x="0" y="0"/>'
        f'<a:ext cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'</a:xfrm>'
        f'<a:prstGeom prst="{preset}">'
        f'<a:avLst/>'
        f'</a:prstGeom>'
        f'<a:solidFill>'
        f'<a:srgbClr val="{color}"/>'
        f'</a:solidFill>'
        f'<a:ln w="12700">'
        f'<a:solidFill>'
        f'<a:srgbClr val="000000"/>'
        f'</a:solidFill>'
        f'</a:ln>'
        f'</wps:spPr>'
        f'<wps:txbx>'
        f'<w:txbxContent xmlns:w="{W_NS}">'
        f'<w:p>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:sz w:val="16"/>'
        f'<w:color w:val="{text_color}"/>'
        f'</w:rPr>'
        f'<w:t>{label}</w:t>'
        f'</w:r>'
        f'</w:p>'
        f'</w:txbxContent>'
        f'</wps:txbx>'
        f'<wps:bodyPr anchor="ctr" anchorCtr="0" lIns="0" tIns="0" rIns="0" bIns="0"/>'
        f'</wps:wsp>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
        f'</w:r>'
    )


# Step 1: Create base document with python-docx
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Single paragraph that will hold all the floating shape anchors
p = doc.add_paragraph()
run = p.add_run("SHAPE_PLACEHOLDER")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Step 2: Post-process ZIP to inject shape drawings
all_shapes_xml = "".join(
    shape_anchor_xml(i, preset, color, label)
    for i, (preset, color, label) in enumerate(SHAPES)
)

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        # Add namespace declarations to the root element
        # Find the <w:document ...> opening tag and add our namespaces
        ns_decls = (
            f' xmlns:wp="{WP_NS}"'
            f' xmlns:a="{A_NS}"'
            f' xmlns:wps="{WPS_NS}"'
            f' xmlns:r="{R_NS}"'
            f' xmlns:mc="{MC_NS}"'
        )
        # Add namespaces if not already present
        if 'xmlns:wp=' not in doc_xml:
            doc_xml = doc_xml.replace(
                '<w:document ',
                f'<w:document {ns_decls} ',
                1,
            )

        # Replace the placeholder run with shape anchors
        placeholder_pattern = r'<w:r>.*?<w:t>SHAPE_PLACEHOLDER</w:t>\s*</w:r>'
        doc_xml = re.sub(
            placeholder_pattern,
            all_shapes_xml,
            doc_xml,
            count=1,
            flags=re.DOTALL,
        )

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)
print(f"Generated {OUT}")
