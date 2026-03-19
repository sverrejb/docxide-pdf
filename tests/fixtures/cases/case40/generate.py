"""Generate a DOCX with a half-width floating table and text wrapping for case40.

Tests:
- Floating table positioned to the left, half page width
- Body text wrapping alongside the table (rightFromText gap)
- Text continuing below the table after wrapping
- Multiple floating table positions: left-aligned and right-aligned
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT = Path("tests/fixtures/cases/case40/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- Section 1: Left-aligned floating table with text wrapping ---

p1 = doc.add_paragraph(
    "This anchor paragraph precedes a left-aligned floating table that is roughly "
    "half the page width. The body text below should wrap to the right of the table."
)

# Table 1: left-aligned, ~3 inches wide (half of 6.5" text width)
table1 = doc.add_table(rows=5, cols=2)
rows_data1 = [
    ("Category", "Value"),
    ("Revenue", "$120,000"),
    ("Expenses", "$85,000"),
    ("Net Income", "$35,000"),
    ("Growth Rate", "12.5%"),
]
for ri, row_data in enumerate(rows_data1):
    for ci, text in enumerate(row_data):
        table1.rows[ri].cells[ci].text = text

# This text should wrap to the right of the floating table
p2 = doc.add_paragraph(
    "This paragraph should appear to the right of the floating table, not below it. "
    "Word wraps body text around floating tables when there is sufficient horizontal "
    "space. The table is positioned on the left side of the margin with a gap of "
    "0.1 inches between the table edge and this flowing text. "
    "This text continues to fill the space beside the table. "
    "When the text extends below the bottom of the table, it should revert to full "
    "page width. Additional sentences ensure we have enough content to flow past "
    "the table boundary and demonstrate the return to normal full-width layout. "
    "The quick brown fox jumps over the lazy dog. Pack my box with five dozen "
    "liquor jugs. How vexingly quick daft zebras jump."
)

p3 = doc.add_paragraph(
    "This paragraph should definitely be below the floating table, at full page "
    "width, since the table has ended. Normal layout resumes here."
)

# --- Section 2: Right-aligned floating table ---

doc.add_page_break()

p4 = doc.add_paragraph(
    "This section has a right-aligned floating table. The body text should wrap "
    "to the left of the table."
)

table2 = doc.add_table(rows=4, cols=2)
rows_data2 = [
    ("Item", "Status"),
    ("Design", "Complete"),
    ("Development", "In Progress"),
    ("Testing", "Pending"),
]
for ri, row_data in enumerate(rows_data2):
    for ci, text in enumerate(row_data):
        table2.rows[ri].cells[ci].text = text

p5 = doc.add_paragraph(
    "This text wraps to the left of the right-aligned floating table. The table "
    "is anchored to the right side of the text margin. Body text fills the space "
    "on the left side. When there is enough vertical content, the text flows past "
    "the table and returns to full width. Lorem ipsum dolor sit amet, consectetur "
    "adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna "
    "aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris."
)

p6 = doc.add_paragraph(
    "After the right-aligned table, text returns to full page width as expected."
)

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Post-process: inject tblpPr into both tables
# Table 1: left-aligned, half-width, with text gap
TBL1_P_PR = (
    '<w:tblpPr'
    ' w:leftFromText="0"'
    ' w:rightFromText="144"'
    ' w:topFromText="72"'
    ' w:bottomFromText="72"'
    ' w:vertAnchor="text"'
    ' w:horzAnchor="margin"'
    ' w:tblpXSpec="left"'
    ' w:tblpY="0"/>'
)

TBL_BORDERS = (
    '<w:tblBorders>'
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '</w:tblBorders>'
)

# Table 1 column widths: ~3 inches = 4320 twips total (2160 each)
TBL1_GRID = (
    '<w:tblGrid>'
    '<w:gridCol w:w="2160"/>'
    '<w:gridCol w:w="2160"/>'
    '</w:tblGrid>'
)
TBL1_WIDTH = '<w:tblW w:type="dxa" w:w="4320"/>'

# Table 2: right-aligned, half-width
TBL2_P_PR = (
    '<w:tblpPr'
    ' w:leftFromText="144"'
    ' w:rightFromText="0"'
    ' w:topFromText="72"'
    ' w:bottomFromText="72"'
    ' w:vertAnchor="text"'
    ' w:horzAnchor="margin"'
    ' w:tblpXSpec="right"'
    ' w:tblpY="0"/>'
)

TBL2_GRID = (
    '<w:tblGrid>'
    '<w:gridCol w:w="2160"/>'
    '<w:gridCol w:w="2160"/>'
    '</w:tblGrid>'
)
TBL2_WIDTH = '<w:tblW w:type="dxa" w:w="4320"/>'

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        # Find all </w:tblPr> and inject into first and second tables
        parts = doc_xml.split("</w:tblPr>")
        if len(parts) >= 3:
            # Table 1: inject positioning + borders, replace width and grid
            parts[0] = re.sub(
                r'<w:tblW [^/]*/>', TBL1_WIDTH, parts[0], count=1
            )
            parts[0] = re.sub(
                r'<w:tblStyle w:val="[^"]*"/>', "", parts[0], count=1
            )
            parts[0] += TBL1_P_PR + TBL_BORDERS + "</w:tblPr>"
            # Replace grid for table 1
            parts[1] = re.sub(
                r'<w:tblGrid>.*?</w:tblGrid>', TBL1_GRID, parts[1], count=1
            )

            # Table 2: inject positioning + borders, replace width and grid
            parts[1] = re.sub(
                r'<w:tblW [^/]*/>', TBL2_WIDTH, parts[1], count=1
            )
            parts[1] = re.sub(
                r'<w:tblStyle w:val="[^"]*"/>', "", parts[1], count=1
            )
            parts[1] += TBL2_P_PR + TBL_BORDERS + "</w:tblPr>"
            parts[2] = re.sub(
                r'<w:tblGrid>.*?</w:tblGrid>', TBL2_GRID, parts[2], count=1
            )

            doc_xml = "".join(parts)

        # Also set cell widths to match grid (2160 twips each)
        doc_xml = re.sub(
            r'<w:tcW w:type="dxa" w:w="\d+"/>',
            '<w:tcW w:type="dxa" w:w="2160"/>',
            doc_xml,
        )

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)
print(f"Generated {OUT}")
