"""Generate case57: 2D Picture Effects test fixture.

Creates a DOCX with images demonstrating implementable picture effects:
soft edge, reflection, glow, inner shadow, and combinations.
Uses real cat photos from placekitten.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
import os
import io
import xml.etree.ElementTree as ET
import urllib.request

DML = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"

def download_image(url):
    """Download an image from URL, return bytes."""
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req) as resp:
        return resp.read()

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "input.docx")

    # Download cat images (different sizes for variety)
    print("Downloading cat images...")
    cat_urls = [
        "https://picsum.photos/400/300",
        "https://picsum.photos/380/280",
        "https://picsum.photos/420/310",
        "https://picsum.photos/390/290",
        "https://picsum.photos/410/300",
        "https://picsum.photos/400/280",
        "https://picsum.photos/380/300",
        "https://picsum.photos/420/290",
        "https://picsum.photos/400/310",
        "https://picsum.photos/390/300",
    ]
    images = []
    for i, url in enumerate(cat_urls):
        print(f"  Downloading image {i+1}/10...")
        images.append(download_image(url))

    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("2D Picture Effects (case57)")
    run.bold = True
    run.font.size = Pt(16)

    labels = [
        "1. Soft Edge (small, 6pt radius)",
        "2. Soft Edge (large, 25pt radius)",
        "3. Reflection (standard)",
        "4. Reflection (tight, offset)",
        "5. Glow (blue, 10pt)",
        "6. Glow (gold, 18pt)",
        "7. Inner Shadow (45deg)",
        "8. Combined: outer shadow + soft edge",
        "9. Combined: white frame + reflection",
        "10. No effects (baseline)",
    ]

    for i, label in enumerate(labels):
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run().add_picture(io.BytesIO(images[i]), width=Inches(2.5))

        doc.add_paragraph()

    doc.save(output_path)
    print("Injecting picture effects...")
    inject_effects(output_path)
    print(f"Generated {output_path}")


def inject_effects(docx_path):
    """Post-process DOCX ZIP to add picture effects to each image's spPr."""
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': DML,
        'pic': PIC,
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    }
    for prefix, uri in namespaces.items():
        ET.register_namespace(prefix, uri)

    with zipfile.ZipFile(docx_path, 'r') as zin:
        entries = {}
        for name in zin.namelist():
            entries[name] = zin.read(name)

    tree = ET.parse(io.BytesIO(entries['word/document.xml']))
    root = tree.getroot()
    pics = root.findall(f'.//{{{PIC}}}pic')

    effect_configs = [
        # 1. Soft Edge small
        {"effectLst": [("softEdge", {"rad": "76200"}, [])]},
        # 2. Soft Edge large
        {"effectLst": [("softEdge", {"rad": "317500"}, [])]},
        # 3. Reflection standard
        {"effectLst": [("reflection", {
            "blurRad": "6350", "stA": "50000", "stPos": "0",
            "endA": "300", "endPos": "100000",
            "dist": "0", "dir": "5400000", "fadeDir": "5400000",
            "sx": "100000", "sy": "-100000", "algn": "bl", "rotWithShape": "0"
        }, [])]},
        # 4. Reflection tight
        {"effectLst": [("reflection", {
            "blurRad": "12700", "stA": "28000", "stPos": "0",
            "endA": "0", "endPos": "100000",
            "dist": "29997", "dir": "5400000", "fadeDir": "5400000",
            "sx": "100000", "sy": "-100000", "algn": "bl", "rotWithShape": "0"
        }, [])]},
        # 5. Glow blue
        {"effectLst": [("glow", {"rad": "127000"}, [
            ("srgbClr", {"val": "4472C4"}, [("alpha", {"val": "60000"}, [])])
        ])]},
        # 6. Glow gold
        {"effectLst": [("glow", {"rad": "228600"}, [
            ("srgbClr", {"val": "FFC000"}, [("alpha", {"val": "40000"}, [])])
        ])]},
        # 7. Inner Shadow
        {"effectLst": [("innerShdw", {"blurRad": "63500", "dist": "50800", "dir": "2700000"}, [
            ("srgbClr", {"val": "000000"}, [("alpha", {"val": "50000"}, [])])
        ])]},
        # 8. Combined: outer shadow + soft edge
        {"effectLst": [
            ("outerShdw", {"blurRad": "190500", "dist": "101600", "dir": "2700000", "algn": "tl", "rotWithShape": "0"}, [
                ("srgbClr", {"val": "000000"}, [("alpha", {"val": "55000"}, [])])
            ]),
            ("softEdge", {"rad": "101600"}, []),
        ]},
        # 9. Combined: white frame + reflection
        {"line": {"w": "38100", "color": "FFFFFF"},
         "effectLst": [("reflection", {
            "blurRad": "6350", "stA": "50000", "stPos": "0",
            "endA": "300", "endPos": "100000",
            "dist": "0", "dir": "5400000", "fadeDir": "5400000",
            "sx": "100000", "sy": "-100000", "algn": "bl", "rotWithShape": "0"
        }, [])]},
        # 10. No effects
        None,
    ]

    for i, pic in enumerate(pics):
        if i >= len(effect_configs) or effect_configs[i] is None:
            continue

        config = effect_configs[i]
        spPr = pic.find(f'{{{PIC}}}spPr')
        if spPr is None:
            spPr = ET.SubElement(pic, f'{{{PIC}}}spPr')

        # Add line/border if specified
        if "line" in config:
            ln = ET.SubElement(spPr, f'{{{DML}}}ln')
            ln.set('w', config["line"]["w"])
            fill = ET.SubElement(ln, f'{{{DML}}}solidFill')
            clr = ET.SubElement(fill, f'{{{DML}}}srgbClr')
            clr.set('val', config["line"]["color"])

        # Add effect list
        if "effectLst" in config:
            effectLst = ET.SubElement(spPr, f'{{{DML}}}effectLst')
            for effect_name, attrs, children in config["effectLst"]:
                add_element(effectLst, effect_name, attrs, children)

    buf = io.BytesIO()
    tree.write(buf, xml_declaration=True, encoding='UTF-8')
    entries['word/document.xml'] = buf.getvalue()

    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)


def add_element(parent, tag, attrs, children):
    """Recursively build an XML element tree under the DML namespace."""
    elem = ET.SubElement(parent, f'{{{DML}}}{tag}')
    for k, v in attrs.items():
        elem.set(k, v)
    for child_tag, child_attrs, grandchildren in children:
        add_element(elem, child_tag, child_attrs, grandchildren)
    return elem


if __name__ == '__main__':
    main()
