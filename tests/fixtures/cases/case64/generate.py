#!/usr/bin/env python3
"""Generate case64 test fixture: multi-page document with comments only on
the middle page.

Page layout exercised:
  - Page 1: no comments (body text only)
  - Page 2: several comments (the right-side comment pane should render here)
  - Page 3: no comments

Word's PDF export draws the comment pane and applies the body content
scaling on EVERY page once the document contains any comments, even on
pages without any comment anchored to them. This fixture lets us verify
that pages 1 and 3 still get the pane geometry / body scaling, and that
the page-2 callouts align with their highlighted phrases.

See case63 for the comment-injection mechanics; this file reuses the same
post-processing helper.
"""
import shutil
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COMMENTS_XML_TMPL = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
{comments}
</w:comments>
"""


def _comment_xml(cid, author, initials, date, text):
    return (
        f'  <w:comment w:id="{cid}" w:author="{author}" '
        f'w:initials="{initials}" w:date="{date}">\n'
        f"    <w:p><w:r><w:t>{text}</w:t></w:r></w:p>\n"
        f"  </w:comment>"
    )


def add_commented_run(paragraph, text, comment_id):
    p = paragraph._p

    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    p.append(start)

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    p.append(run)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    p.append(end)

    ref_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "CommentReference")
    rPr.append(rStyle)
    ref_run.append(rPr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)
    p.append(ref_run)


def inject_comments_part(docx_path, comments):
    tmp_path = docx_path.with_suffix(".docx.tmp")

    comments_xml = COMMENTS_XML_TMPL.format(
        comments="\n".join(
            _comment_xml(c["id"], c["author"], c["initials"], c["date"], c["text"])
            for c in comments
        )
    )

    comments_rel_type = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    )
    comments_content_type = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.comments+xml"
    )

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp_path, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                data = data.replace(
                    b"</Relationships>",
                    (
                        f'<Relationship Id="rIdComments" Type="{comments_rel_type}" '
                        f'Target="comments.xml"/></Relationships>'
                    ).encode("utf-8"),
                )
            elif item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"</Types>",
                    (
                        f'<Override PartName="/word/comments.xml" '
                        f'ContentType="{comments_content_type}"/></Types>'
                    ).encode("utf-8"),
                )
            zout.writestr(item, data)
        zout.writestr("word/comments.xml", comments_xml)

    shutil.move(tmp_path, docx_path)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(12)

    # Page 1 — no comments. Plain paragraphs that establish the body
    # column width and verify the pane still renders with no callouts.
    doc.add_heading("Page 1: introduction (no comments)", level=1)
    doc.add_paragraph(
        "This first page has no comments attached to it. Word's PDF export "
        "still draws the gray comment pane on the right side of every page "
        "once the document contains any comments anywhere, so the body on "
        "this page should be scaled and indented the same way as pages with "
        "actual callouts — only the callouts themselves are absent here."
    )
    doc.add_paragraph(
        "The point of this paragraph is to make sure the body text wraps "
        "correctly when no comment anchors exist on the page. The pane area "
        "should still appear with its gray background, but with no boxes or "
        "connector lines."
    )

    add_page_break(doc)

    # Page 2 — six comments anchored to phrases across several paragraphs.
    doc.add_heading("Page 2: heavy commenting", level=1)
    comments = [
        {
            "id": 1, "author": "Reviewer A", "initials": "RA",
            "date": "2024-01-01T00:00:00Z",
            "text": "Short note on the opening phrase.",
        },
        {
            "id": 2, "author": "Reviewer A", "initials": "RA",
            "date": "2024-01-01T00:00:00Z",
            "text": (
                "A longer remark that should wrap onto multiple lines inside "
                "the callout. It exists to exercise the wrapping algorithm "
                "for comment text."
            ),
        },
        {
            "id": 3, "author": "Reviewer B", "initials": "RB",
            "date": "2024-01-01T00:00:00Z",
            "text": "Unicode check — naïve façade, Besøk ✓",
        },
        {
            "id": 4, "author": "Reviewer B", "initials": "RB",
            "date": "2024-01-01T00:00:00Z",
            "text": "Question: is this the canonical phrasing?",
        },
        {
            "id": 5, "author": "Reviewer A", "initials": "RA",
            "date": "2024-01-01T00:00:00Z",
            "text": (
                "Multi-sentence remark. The first sentence sets up context. "
                "The second sentence asks a follow-up. The third sentence "
                "wraps things up so we get a tall callout."
            ),
        },
        {
            "id": 6, "author": "Reviewer C", "initials": "RC",
            "date": "2024-01-01T00:00:00Z",
            "text": "Final note for this page — keep it tight.",
        },
    ]

    p = doc.add_paragraph()
    add_commented_run(p, "The opening sentence", comment_id=1)
    p.add_run(
        " sets the tone for the entire page. It anchors the first comment "
        "on the page and verifies that the callout aligns with this row."
    )

    p = doc.add_paragraph()
    p.add_run("The next paragraph contains ")
    add_commented_run(p, "a phrase that is longer than usual", comment_id=2)
    p.add_run(
        " so the callout body can demonstrate multi-line wrapping. The "
        "surrounding text must continue to flow normally despite the "
        "comment range markers split across multiple runs."
    )

    p = doc.add_paragraph()
    p.add_run("This paragraph includes ")
    add_commented_run(p, "non-ASCII characters", comment_id=3)
    p.add_run(" to confirm that comment text and labels survive correctly.")

    p = doc.add_paragraph()
    p.add_run("A short question is attached to ")
    add_commented_run(p, "this particular wording", comment_id=4)
    p.add_run(" to keep the pane busy.")

    p = doc.add_paragraph()
    add_commented_run(p, "A multi-sentence remark", comment_id=5)
    p.add_run(
        " is anchored here. We want to see how the pane stacks a tall "
        "callout next to a shorter neighbour, and that the connector line "
        "still leaves the highlight cleanly."
    )

    p = doc.add_paragraph()
    p.add_run("Final paragraph on this page ends with ")
    add_commented_run(p, "the last commented phrase", comment_id=6)
    p.add_run(".")

    add_page_break(doc)

    # Page 3 — no comments again. Used to verify that the comment pane
    # continues to render on a tail page even though no anchors point to it.
    doc.add_heading("Page 3: conclusion (no comments)", level=1)
    doc.add_paragraph(
        "This third page also has no comments. The intent here is to "
        "confirm that the comment pane stays on every page of the "
        "document — Word draws it uniformly — and that the body scaling "
        "is applied consistently regardless of whether any anchor lands "
        "on this particular page."
    )
    doc.add_paragraph(
        "There should be no callouts visible on this page, no connector "
        "lines, and no highlights. Only the gray pane and the scaled body "
        "text should be present."
    )

    out = Path(__file__).parent / "input.docx"
    doc.save(str(out))
    inject_comments_part(out, comments)

    print(f"Saved {out}")
    print("Next: open in Word, File > Save As > PDF to create reference.pdf")


if __name__ == "__main__":
    main()
