"""Generate case58: 3D / Deferred Picture Effects test fixture.

Creates a DOCX with images demonstrating hard-to-implement effects:
bevel, metal frame, 3D rotation, preset shadows. These are deferred
but having the fixture lets us track future progress.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
import os
import io
import xml.etree.ElementTree as ET
import urllib.request

DML = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"

def download_image(url):
    """Download an image from URL, return bytes."""
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req) as resp:
        return resp.read()

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "input.docx")

    print("Downloading images...")
    urls = [
        "https://picsum.photos/400/300",
        "https://picsum.photos/380/280",
        "https://picsum.photos/420/310",
        "https://picsum.photos/390/290",
        "https://picsum.photos/410/300",
        "https://picsum.photos/400/280",
        "https://picsum.photos/380/300",
        "https://picsum.photos/420/290",
    ]
    images = []
    for i, url in enumerate(urls):
        print(f"  Downloading image {i+1}/{len(urls)}...")
        images.append(download_image(url))

    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("3D / Deferred Picture Effects (case58)")
    run.bold = True
    run.font.size = Pt(16)

    labels = [
        "1. Bevel (circle, top only)",
        "2. Bevel (relaxedInset)",
        "3. Bevel (artDeco)",
        "4. Metal Frame (metal material + bevel + contour)",
        "5. 3D Rotation (perspective front)",
        "6. Soft Edge Material (softEdge + convex bevel)",
        "7. Preset Shadow (shdw14)",
        "8. No effects (baseline)",
    ]

    for i, label in enumerate(labels):
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run().add_picture(io.BytesIO(images[i]), width=Inches(2.5))

        doc.add_paragraph()

    doc.save(output_path)
    print("Injecting 3D effects...")
    inject_effects(output_path)
    print(f"Generated {output_path}")


def a(tag):
    return f'{{{DML}}}{tag}'


def make_scene3d(parent, camera_prst="orthographicFront", rig="threePt", rig_dir="t",
                 camera_rot=None):
    """Add a standard scene3d element."""
    scene = ET.SubElement(parent, a('scene3d'))
    camera = ET.SubElement(scene, a('camera'))
    camera.set('prst', camera_prst)
    if camera_rot:
        rot = ET.SubElement(camera, a('rot'))
        rot.set('lat', camera_rot.get('lat', '0'))
        rot.set('lon', camera_rot.get('lon', '0'))
        rot.set('rev', camera_rot.get('rev', '0'))
    light = ET.SubElement(scene, a('lightRig'))
    light.set('rig', rig)
    light.set('dir', rig_dir)
    return scene


def make_sp3d(parent, material=None, contour_w=None, contour_color=None,
              bevel_t=None, bevel_b=None):
    """Add sp3d element with optional bevel, material, contour."""
    sp3d = ET.SubElement(parent, a('sp3d'))
    if material:
        sp3d.set('prstMaterial', material)
    if contour_w:
        sp3d.set('contourW', contour_w)
    if bevel_t:
        bt = ET.SubElement(sp3d, a('bevelT'))
        for k, v in bevel_t.items():
            bt.set(k, v)
    if bevel_b:
        bb = ET.SubElement(sp3d, a('bevelB'))
        for k, v in bevel_b.items():
            bb.set(k, v)
    if contour_color:
        cc = ET.SubElement(sp3d, a('contourClr'))
        clr = ET.SubElement(cc, a('srgbClr'))
        clr.set('val', contour_color)
    return sp3d


def inject_effects(docx_path):
    """Post-process DOCX ZIP to add 3D effects to each image's spPr."""
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': DML,
        'pic': PIC,
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    }
    for prefix, uri in namespaces.items():
        ET.register_namespace(prefix, uri)

    with zipfile.ZipFile(docx_path, 'r') as zin:
        entries = {}
        for name in zin.namelist():
            entries[name] = zin.read(name)

    tree = ET.parse(io.BytesIO(entries['word/document.xml']))
    root = tree.getroot()
    pics = root.findall(f'.//{{{PIC}}}pic')

    for i, pic in enumerate(pics):
        spPr = pic.find(f'{{{PIC}}}spPr')
        if spPr is None:
            spPr = ET.SubElement(pic, f'{{{PIC}}}spPr')

        if i == 0:
            # 1. Bevel (circle, top only)
            make_scene3d(spPr)
            make_sp3d(spPr, bevel_t={"w": "127000", "h": "50800", "prst": "circle"})

        elif i == 1:
            # 2. Bevel (relaxedInset)
            make_scene3d(spPr)
            make_sp3d(spPr, bevel_t={"w": "190500", "h": "63500", "prst": "relaxedInset"})

        elif i == 2:
            # 3. Bevel (artDeco)
            make_scene3d(spPr)
            make_sp3d(spPr, bevel_t={"w": "139700", "h": "139700", "prst": "artDeco"})

        elif i == 3:
            # 4. Metal Frame
            make_scene3d(spPr, rig="harsh", rig_dir="t")
            make_sp3d(spPr, material="metal", contour_w="25400",
                      bevel_t={"w": "88900", "h": "88900", "prst": "circle"},
                      contour_color="C0C0C0")

        elif i == 4:
            # 5. 3D Rotation (perspective)
            make_scene3d(spPr, camera_prst="perspectiveFront",
                         camera_rot={"lat": "300000", "lon": "4200000", "rev": "0"})
            make_sp3d(spPr)

        elif i == 5:
            # 6. Soft edge material + convex bevel
            make_scene3d(spPr)
            make_sp3d(spPr, material="softEdge",
                      bevel_t={"w": "63500", "h": "25400", "prst": "convex"})

        elif i == 6:
            # 7. Preset Shadow (shdw14)
            effectLst = ET.SubElement(spPr, a('effectLst'))
            prstShdw = ET.SubElement(effectLst, a('prstShdw'))
            prstShdw.set('prst', 'shdw14')
            clr = ET.SubElement(prstShdw, a('srgbClr'))
            clr.set('val', '000000')
            alpha = ET.SubElement(clr, a('alpha'))
            alpha.set('val', '50000')

        # i == 7: no effects (baseline)

    buf = io.BytesIO()
    tree.write(buf, xml_declaration=True, encoding='UTF-8')
    entries['word/document.xml'] = buf.getvalue()

    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)


if __name__ == '__main__':
    main()
