"""Generate case20: Advanced Nested Lists

Tests:
- Multi-level labels (%1.%2, %1.%2.%3) with cross-level substitution
- upperRoman and upperLetter formats
- Bullet lists with different symbols per level
- Two independent numbered lists (separate counters)
- Deep nesting (4 levels) with various format combos
- Interleaved bullet and numbered sections
"""

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import pathlib
import io

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ABS_BASE = 100
NUM_BASE = 100


def make_abstract_num_xml(abstract_num_id, levels):
    lvls = []
    for ilvl, start, num_fmt, lvl_text, left, hanging in levels:
        lvls.append(
            f'<w:lvl w:ilvl="{ilvl}">'
            f'<w:start w:val="{start}"/>'
            f'<w:numFmt w:val="{num_fmt}"/>'
            f'<w:lvlText w:val="{lvl_text}"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="{left}" w:hanging="{hanging}"/></w:pPr>'
            f'</w:lvl>'
        )
    return (
        f'<w:abstractNum w:abstractNumId="{abstract_num_id}" xmlns:w="{WML}">'
        f'<w:multiLevelType w:val="hybridMultilevel"/>'
        + "".join(lvls)
        + "</w:abstractNum>"
    )


def make_num_xml(num_id, abstract_num_id):
    return (
        f'<w:num w:numId="{num_id}" xmlns:w="{WML}">'
        f'<w:abstractNumId w:val="{abstract_num_id}"/>'
        f'</w:num>'
    )


def add_list_para(doc, text, num_id, ilvl):
    p = doc.add_paragraph(text)
    ppr = p._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_pr.append(ilvl_el)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    ppr.append(num_pr)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# --- Numbering definitions ---
# 0: Multi-level outline: I. / A. / 1. / a)
# 1: Bullet list with 3 different chars
# 2: Cross-level labels: %1. / %1.%2 / %1.%2.%3
# 3: Second independent numbered list (decimal, separate counters)

NUM_OUTLINE = NUM_BASE
NUM_BULLETS = NUM_BASE + 1
NUM_CROSSLEVEL = NUM_BASE + 2
NUM_INDEPENDENT = NUM_BASE + 3

# --- Document content ---

# Section 1: Outline-style list (I. A. 1. a))
doc.add_heading("Outline Format", level=2)
add_list_para(doc, "Introduction", NUM_OUTLINE, 0)
add_list_para(doc, "Background", NUM_OUTLINE, 1)
add_list_para(doc, "Historical context", NUM_OUTLINE, 2)
add_list_para(doc, "Early developments", NUM_OUTLINE, 3)
add_list_para(doc, "Later developments", NUM_OUTLINE, 3)
add_list_para(doc, "Current state", NUM_OUTLINE, 2)
add_list_para(doc, "Motivation", NUM_OUTLINE, 1)
add_list_para(doc, "Methods", NUM_OUTLINE, 0)
add_list_para(doc, "Data collection", NUM_OUTLINE, 1)
add_list_para(doc, "Primary sources", NUM_OUTLINE, 2)
add_list_para(doc, "Secondary sources", NUM_OUTLINE, 2)
add_list_para(doc, "Analysis", NUM_OUTLINE, 1)
add_list_para(doc, "Results", NUM_OUTLINE, 0)

# Section 2: Bullet list with varied symbols
doc.add_heading("Bullet Variations", level=2)
add_list_para(doc, "Main point one", NUM_BULLETS, 0)
add_list_para(doc, "Detail with dash", NUM_BULLETS, 1)
add_list_para(doc, "Sub-detail with arrow", NUM_BULLETS, 2)
add_list_para(doc, "Another sub-detail", NUM_BULLETS, 2)
add_list_para(doc, "Another detail", NUM_BULLETS, 1)
add_list_para(doc, "Main point two", NUM_BULLETS, 0)
add_list_para(doc, "Single detail", NUM_BULLETS, 1)

# Section 3: Cross-level numbering (1. / 1.a / 1.a.i)
doc.add_heading("Cross-Level Numbering", level=2)
add_list_para(doc, "Chapter one", NUM_CROSSLEVEL, 0)
add_list_para(doc, "Section one-a", NUM_CROSSLEVEL, 1)
add_list_para(doc, "Clause one-a-i", NUM_CROSSLEVEL, 2)
add_list_para(doc, "Clause one-a-ii", NUM_CROSSLEVEL, 2)
add_list_para(doc, "Section one-b", NUM_CROSSLEVEL, 1)
add_list_para(doc, "Chapter two", NUM_CROSSLEVEL, 0)
add_list_para(doc, "Section two-a", NUM_CROSSLEVEL, 1)
add_list_para(doc, "Clause two-a-i", NUM_CROSSLEVEL, 2)
add_list_para(doc, "Chapter three", NUM_CROSSLEVEL, 0)

# Section 4: Independent second list (proves separate counter tracking)
doc.add_heading("Independent List", level=2)
add_list_para(doc, "Alpha list item one", NUM_INDEPENDENT, 0)
add_list_para(doc, "Alpha list item two", NUM_INDEPENDENT, 0)
add_list_para(doc, "Alpha list item three", NUM_INDEPENDENT, 0)

# Save, then post-process to inject custom numbering
tmp_buf = io.BytesIO()
doc.save(tmp_buf)
tmp_buf.seek(0)

out_buf = io.BytesIO()
with zipfile.ZipFile(tmp_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/numbering.xml":
            tree = etree.fromstring(data)
            nsmap = {"w": WML}
            first_num = tree.find("w:num", nsmap)

            defs = [
                # Outline: I. / A. / 1. / a)
                (ABS_BASE, [
                    (0, 1, "upperRoman",  "%1.", 720, 360),
                    (1, 1, "upperLetter", "%2.", 1440, 360),
                    (2, 1, "decimal",     "%3.", 2160, 360),
                    (3, 1, "lowerLetter", "%4)", 2880, 360),
                ]),
                # Bullets: bullet / dash / arrow
                (ABS_BASE + 1, [
                    (0, 1, "bullet", "\u2022", 720, 360),
                    (1, 1, "bullet", "\u2013", 1440, 360),
                    (2, 1, "bullet", "\u203A", 2160, 360),
                ]),
                # Cross-level: %1. / %1.%2 / %1.%2.%3
                (ABS_BASE + 2, [
                    (0, 1, "decimal",     "%1.", 720, 360),
                    (1, 1, "lowerLetter", "%1.%2", 1440, 720),
                    (2, 1, "lowerRoman",  "%1.%2.%3", 2160, 1080),
                ]),
                # Independent decimal
                (ABS_BASE + 3, [
                    (0, 1, "decimal", "%1.", 720, 360),
                ]),
            ]
            for abs_id, levels in defs:
                el = etree.fromstring(make_abstract_num_xml(abs_id, levels))
                if first_num is not None:
                    tree.insert(list(tree).index(first_num), el)
                else:
                    tree.append(el)

            for i in range(4):
                tree.append(etree.fromstring(make_num_xml(NUM_BASE + i, ABS_BASE + i)))

            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        elif item.filename == "word/settings.xml":
            tree = etree.fromstring(data)
            for compat_setting in tree.iter("{%s}compatSetting" % WML):
                if compat_setting.get(qn("w:name")) == "compatibilityMode":
                    compat_setting.set(qn("w:val"), "15")
            data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        zout.writestr(item, data)

out_path = pathlib.Path(__file__).parent / "input.docx"
out_path.write_bytes(out_buf.getvalue())
print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")
