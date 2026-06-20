#!/usr/bin/env python3
"""case66: legacy run text-effect toggles — w:outline / w:shadow / w:emboss /
w:imprint / w:effect (§17.3.2.23 / .31 / .13 / .18 + ST_TextEffect).

These are the Word 97-era boolean run properties, NOT the modern w14 DrawingML
text effects (glow/shadow/outline) handled in src/docx/wordart.rs. Each is its
own paragraph so the reference PDF isolates how Word prints that one toggle:
  - outline : glyphs drawn as hollow strokes (stroke text render mode)
  - shadow  : offset drop shadow behind the glyphs
  - emboss  : text raised out of the page (light highlight + dark shadow)
  - imprint : text engraved into the page (dark/light reversed)
  - effect  : animated text effect (shimmer); no print form — base text only

python-docx exposes none of these, so we inject raw <w:r> runs by ZIP
post-processing (same pattern as case43).
"""

import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case66/input.docx")

# Big bold dark text so the stroke/shadow/raise effects are visible.
SZ = 56  # half-points = 28pt

EFFECTS = [
    ("outline", "<w:outline/>", "Outline effect"),
    ("shadow", "<w:shadow/>", "Shadow effect"),
    ("emboss", "<w:emboss/>", "Emboss effect"),
    ("imprint", "<w:imprint/>", "Imprint effect"),
    ("effect", '<w:effect w:val="shimmer"/>', "Shimmer effect"),
]


def effect_para(toggle_xml, text):
    return (
        f"<w:p><w:r><w:rPr>"
        f"<w:b/><w:sz w:val=\"{SZ}\"/><w:szCs w:val=\"{SZ}\"/>{toggle_xml}"
        f"</w:rPr><w:t xml:space=\"preserve\">{text}</w:t></w:r></w:p>"
    )


def main():
    doc = Document()
    run = doc.add_paragraph().add_run("Run Text-Effect Toggles")
    run.bold = True
    run.font.size = Pt(14)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode("utf-8")
        other = {n: zin.read(n) for n in zin.namelist() if n != "word/document.xml"}

    paras = "".join(effect_para(x, t) for _, x, t in EFFECTS)
    # Final sectPr is a direct child of <w:body>; insert effect paragraphs before it.
    assert "<w:sectPr" in doc_xml, "expected a body-level sectPr to anchor injection"
    doc_xml = doc_xml.replace("<w:sectPr", paras + "<w:sectPr", 1)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
