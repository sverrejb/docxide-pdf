"""Diagnostic fixture for character width precision testing.

Contains blocks of repeated short words at known font sizes,
allowing measurement of how many characters Word fits per line
vs our output. Left-aligned to avoid justified spacing confounds.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# US Letter, 1-inch margins
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.0

# Block 1: Calibri 12pt — repeating "mm " (3-char word + space)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("mm " * 200)
run.font.name = 'Calibri'
run.font.size = Pt(12)

# Block 2: Calibri 12pt — repeating "test " (5-char word)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("test " * 200)
run.font.name = 'Calibri'
run.font.size = Pt(12)

# Block 3: TNR 12pt — repeating "test "
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("test " * 200)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Block 4: TNR 14pt — repeating "test " (matches Russian essay font size)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("test " * 200)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

# Block 5: Arial 12pt — repeating "test "
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("test " * 200)
run.font.name = 'Arial'
run.font.size = Pt(12)

# Block 6: Calibri 11pt — natural sentence text
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("The quick brown fox jumps over the lazy dog. " * 30)
run.font.name = 'Calibri'
run.font.size = Pt(11)

doc.save("tests/fixtures/cases/case54/input.docx")
print("Created case54/input.docx")
