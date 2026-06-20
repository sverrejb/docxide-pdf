#!/usr/bin/env python3
"""case67: w:customXml transparent wrapper at all four placements
(§17.5.1.1 block / .3 cell / .5 row / .6 inline).

w:customXml wraps content to attach schema-based semantics but has NO print
form of its own — like w:sdt, the renderer must descend through it and render
the children. docxside has zero customXml handling today, so any content
wrapped *solely* in customXml is silently dropped. This fixture puts text that
exists ONLY inside customXml wrappers at each of the four placement levels, so
the reference PDF reveals whether each placement's content survives:
  - block-level : a whole paragraph wrapped in <w:customXml>
  - inline-level: a run wrapped in <w:customXml> mid-paragraph
  - cell-level  : a <w:tc> wrapped in <w:customXml> inside a row
  - row-level   : a whole <w:tr> wrapped in <w:customXml> inside the table

python-docx exposes none of this, so we inject the raw XML by ZIP
post-processing (same pattern as case66).
"""

import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path("tests/fixtures/cases/case67/input.docx")

SZ = 28  # half-points = 14pt

# Single black borders so the table (and thus the row/cell content) is visible.
BORDERS = (
    "<w:tblBorders>"
    '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    "</w:tblBorders>"
)


def t(text):
    return (
        f'<w:r><w:rPr><w:sz w:val="{SZ}"/><w:szCs w:val="{SZ}"/></w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t></w:r>'
    )


def cell(text):
    return (
        '<w:tc><w:tcPr><w:tcW w:w="2400" w:type="dxa"/></w:tcPr>'
        f"<w:p>{t(text)}</w:p></w:tc>"
    )


def build_content():
    # Block-level: an entire paragraph that exists only inside customXml.
    block = (
        '<w:customXml w:element="BlockTag">'
        f"<w:p>{t('Block-level customXml content')}</w:p>"
        "</w:customXml>"
    )

    # Inline-level: a run wrapped in customXml between two plain anchor runs.
    inline = (
        "<w:p>"
        + t("Inline: before ")
        + '<w:customXml w:element="InlineTag">'
        + t("[inline-customXml]")
        + "</w:customXml>"
        + t(" after")
        + "</w:p>"
    )

    # Table mixing a cell-level wrapper (row 1, 2nd cell) and a row-level
    # wrapper (row 2 in its entirety).
    table = (
        "<w:tbl>"
        f"<w:tblPr><w:tblW w:w=\"0\" w:type=\"auto\"/>{BORDERS}</w:tblPr>"
        '<w:tblGrid><w:gridCol w:w="2400"/><w:gridCol w:w="2400"/></w:tblGrid>'
        "<w:tr>"
        + cell("plain cell")
        + '<w:customXml w:element="CellTag">'
        + cell("[cell-customXml]")
        + "</w:customXml>"
        + "</w:tr>"
        + '<w:customXml w:element="RowTag">'
        + "<w:tr>"
        + cell("[row-customXml A]")
        + cell("[row-customXml B]")
        + "</w:tr>"
        + "</w:customXml>"
        + "</w:tbl>"
    )

    return block + inline + table


def main():
    doc = Document()
    run = doc.add_paragraph().add_run("Custom XML Wrapper (w:customXml)")
    run.bold = True
    run.font.size = Pt(14)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode("utf-8")
        other = {n: zin.read(n) for n in zin.namelist() if n != "word/document.xml"}

    # Final sectPr is a direct child of <w:body>; inject content before it.
    assert "<w:sectPr" in doc_xml, "expected a body-level sectPr to anchor injection"
    doc_xml = doc_xml.replace("<w:sectPr", build_content() + "<w:sectPr", 1)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
