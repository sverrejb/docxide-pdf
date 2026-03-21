#!/usr/bin/env python3
"""case44: Extended WordArt — more warp presets, glow, combined effects, varied sizes.

Complements case43 by testing additional warp presets and effect combinations:
Page 1:
  1. textDeflate — pinched middle
  2. textInflate — bulging middle
  3. textChevron — V-shape
  4. textTriangle — wide top, narrow bottom
  5. textCircle — circular text path

Page 2:
  6. textCascadeDown — descending cascade
  7. textDoubleWave1 — double wave
  8. textFadeRight — perspective fade
  9. textCanDown — barrel bottom
  10. textRingOutside — ring outside

Page 3:
  11. textDeflateBottom — deflate bottom only
  12. textCurveUp — upward curve
  13. Glow effect (textPlain + glow, no warp distortion)
  14. Combined: outline + shadow + textWave2
  15. Small text with textStop (octagon shape)
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("tests/fixtures/cases/case44/input.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
WP14_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
GRAPHIC_DATA_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

EMU_PER_PT = 12700
EMU_PER_INCH = 914400


def wordart_xml(
    text,
    warp_preset,
    y_offset_emu,
    color="4472C4",
    font_size_hps=72,
    outline_color=None,
    outline_width_emu=9525,
    shadow=False,
    glow=False,
    glow_color="FFC000",
    glow_radius_emu=63500,
    adj_values=None,
    width_emu=int(5.5 * EMU_PER_INCH),
    height_emu=int(1.0 * EMU_PER_INCH),
    font_name=None,
):
    """Generate mc:AlternateContent XML for a WordArt textbox."""

    font_size_pt = font_size_hps // 2

    # w14:textOutline
    outline_xml = ""
    if outline_color:
        outline_xml = (
            f'<w14:textOutline w14:w="{outline_width_emu}" w14:cap="flat">'
            f'<w14:solidFill><w14:srgbClr w14:val="{outline_color}"/></w14:solidFill>'
            f"</w14:textOutline>"
        )

    # w14:shadow
    shadow_xml = ""
    if shadow:
        shadow_xml = (
            f'<w14:shadow w14:blurRad="38100" w14:dist="25400" w14:dir="3600000">'
            f'<w14:srgbClr w14:val="808080">'
            f'<w14:alpha w14:val="60000"/>'
            f"</w14:srgbClr>"
            f"</w14:shadow>"
        )

    # w14:glow
    glow_xml = ""
    if glow:
        glow_xml = (
            f'<w14:glow w14:rad="{glow_radius_emu}">'
            f'<w14:srgbClr w14:val="{glow_color}">'
            f'<w14:alpha w14:val="60000"/>'
            f"</w14:srgbClr>"
            f"</w14:glow>"
        )

    # Font override
    font_xml = ""
    if font_name:
        font_xml = f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'

    # Adjustment values
    av_xml = "<a:avLst/>"
    if adj_values:
        gds = "".join(
            f'<a:gd name="{n}" fmla="val {v}"/>' for n, v in adj_values
        )
        av_xml = f"<a:avLst>{gds}</a:avLst>"

    doc_pr_id = abs(hash(text + warp_preset)) % 100000

    return f"""<mc:AlternateContent xmlns:mc="{MC_NS}">
  <mc:Choice Requires="wps">
    <w:drawing xmlns:w="{W_NS}">
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
          simplePos="0" relativeHeight="251658240" behindDoc="0"
          locked="0" layoutInCell="1" allowOverlap="1"
          xmlns:wp="{WP_NS}" xmlns:wp14="{WP14_NS}">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column"><wp:align>center</wp:align></wp:positionH>
        <wp:positionV relativeFrom="paragraph"><wp:posOffset>{y_offset_emu}</wp:posOffset></wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{doc_pr_id}" name="WordArt {text[:10]}"/>
        <a:graphic xmlns:a="{A_NS}">
          <a:graphicData uri="{GRAPHIC_DATA_URI}">
            <wps:wsp xmlns:wps="{WPS_NS}">
              <wps:cNvSpPr txBox="1"/>
              <wps:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                <a:noFill/>
                <a:ln><a:noFill/></a:ln>
              </wps:spPr>
              <wps:txbx>
                <w:txbxContent xmlns:w="{W_NS}">
                  <w:p>
                    <w:pPr><w:jc w:val="center"/></w:pPr>
                    <w:r>
                      <w:rPr>
                        <w:b/>
                        {font_xml}
                        <w:color w:val="{color}"/>
                        <w:sz w:val="{font_size_hps}"/>
                        <w:szCs w:val="{font_size_hps}"/>
                        {outline_xml}
                        {shadow_xml}
                        {glow_xml}
                      </w:rPr>
                      <w:t>{text}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </wps:txbx>
              <wps:bodyPr fromWordArt="1" wrap="none" lIns="0" tIns="0" rIns="0" bIns="0">
                <a:prstTxWarp prst="{warp_preset}">{av_xml}</a:prstTxWarp>
                <a:spAutoFit/>
              </wps:bodyPr>
            </wps:wsp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </mc:Choice>
</mc:AlternateContent>"""


def build_page(wordarts_spec, page_y_base=0):
    """Build WordArt XML for a set of specs on one page."""
    xml_parts = []
    for spec in wordarts_spec:
        text, preset = spec[0], spec[1]
        y_inches = spec[2]
        kwargs = spec[3] if len(spec) > 3 else {}
        y_emu = int(y_inches * EMU_PER_INCH) + page_y_base
        xml_parts.append(wordart_xml(
            text=text,
            warp_preset=preset,
            y_offset_emu=y_emu,
            **kwargs,
        ))
    return "".join(xml_parts)


def main():
    doc = Document()

    # Page 1: title + first batch of warp presets
    title = doc.add_paragraph()
    run = title.add_run("Extended WordArt Tests — Page 1")
    run.bold = True
    run.font.size = Pt(14)

    # Spacers for page 1
    for _ in range(22):
        doc.add_paragraph("")

    # Page break via run break type="page"
    p_break = doc.add_paragraph()
    run_br = p_break.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_br._element.append(br)

    # Page 2: title + more presets
    title2 = doc.add_paragraph()
    run2 = title2.add_run("Extended WordArt Tests — Page 2")
    run2.bold = True
    run2.font.size = Pt(14)

    for _ in range(22):
        doc.add_paragraph("")

    # Page break
    p_break2 = doc.add_paragraph()
    p_break2_run = p_break2.add_run()
    p_break2_run.add_break(docx.oxml.ns.qn("w:br"))

    # Page 3: title + effects combinations
    title3 = doc.add_paragraph()
    run3 = title3.add_run("Extended WordArt Tests — Page 3")
    run3.bold = True
    run3.font.size = Pt(14)

    for _ in range(22):
        doc.add_paragraph("")

    # Save base document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    # Read document XML
    with zipfile.ZipFile(tmp_path, "r") as zin:
        doc_xml = zin.read("word/document.xml").decode("utf-8")
        other_files = {}
        for name in zin.namelist():
            if name != "word/document.xml":
                other_files[name] = zin.read(name)

    # Add w14 namespace
    if f'xmlns:w14="{W14_NS}"' not in doc_xml:
        doc_xml = doc_xml.replace(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xmlns:w14="{W14_NS}"',
        )
    if "mc:Ignorable" in doc_xml and "w14" not in re.search(r'mc:Ignorable="([^"]*)"', doc_xml).group(1):
        doc_xml = re.sub(
            r'mc:Ignorable="([^"]*)"',
            r'mc:Ignorable="\1 w14"',
            doc_xml,
        )

    # Page 1 WordArts: diverse warp presets
    page1 = [
        ("DEFLATE", "textDeflate", 0.7, {"color": "2E75B6", "outline_color": "1F4E79"}),
        ("INFLATE", "textInflate", 2.2, {"color": "ED7D31", "outline_color": "C44F1A"}),
        ("CHEVRON", "textChevron", 3.8, {"color": "70AD47"}),
        ("TRIANGLE", "textTriangle", 5.5, {"color": "C00000", "outline_color": "800000"}),
        ("CIRCLE", "textCircle", 7.2, {
            "color": "7030A0",
            "height_emu": int(1.5 * EMU_PER_INCH),
            "adj_values": [("adj", "10800000")],
        }),
    ]

    # Page 2 WordArts: more warps
    page2 = [
        ("CASCADE", "textCascadeDown", 0.7, {"color": "4472C4", "outline_color": "2E5FA1"}),
        ("DOUBLE WAVE", "textDoubleWave1", 2.2, {"color": "ED7D31"}),
        ("FADE RIGHT", "textFadeRight", 3.8, {"color": "70AD47", "outline_color": "4A7A2E"}),
        ("BARREL", "textCanDown", 5.5, {"color": "C00000"}),
        ("RING", "textRingOutside", 7.2, {
            "color": "7030A0",
            "height_emu": int(1.5 * EMU_PER_INCH),
        }),
    ]

    # Page 3 WordArts: effects and combinations
    page3 = [
        ("DEFLATE BTM", "textDeflateBottom", 0.7, {"color": "2E75B6", "outline_color": "1F4E79"}),
        ("CURVE UP", "textCurveUp", 2.2, {"color": "ED7D31"}),
        ("GLOW TEXT", "textPlain", 3.8, {
            "color": "4472C4",
            "glow": True,
            "glow_color": "FFC000",
            "font_size_hps": 64,
        }),
        ("WAVE COMBO", "textWave2", 5.5, {
            "color": "C00000",
            "outline_color": "800000",
            "shadow": True,
        }),
        ("STOP", "textStop", 7.2, {
            "color": "70AD47",
            "outline_color": "4A7A2E",
            "font_size_hps": 48,
        }),
    ]

    # Find all <w:p> elements and inject into the first one on each page
    # The structure is: title1, spacers..., break, title2, spacers..., break, title3, spacers...
    # We inject all page1 arts into title1's <w:p>, page2 into title2, page3 into title3

    # Count <w:p> to find injection points
    # Page 1 title = 1st <w:p>
    # Page 1 has 22 spacers → positions 2-23
    # Page break paragraph = 24th <w:p>
    # Page 2 title = 25th <w:p>
    # Page 2 has 22 spacers → positions 26-47
    # Page break paragraph = 48th <w:p>
    # Page 3 title = 49th <w:p>

    page1_xml = build_page(page1)
    page2_xml = build_page(page2)
    page3_xml = build_page(page3)

    # Inject by replacing </w:p> at specific positions
    parts = doc_xml.split("</w:p>")

    # Inject page 1 arts after the first paragraph (index 0)
    parts[0] = parts[0] + page1_xml

    # Inject page 2 arts after the 25th paragraph (index 24)
    if len(parts) > 24:
        parts[24] = parts[24] + page2_xml

    # Inject page 3 arts after the 49th paragraph (index 48)
    if len(parts) > 48:
        parts[48] = parts[48] + page3_xml

    doc_xml = "</w:p>".join(parts)

    # Write output
    OUT.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in other_files.items():
            zout.writestr(name, data)
        zout.writestr("word/document.xml", doc_xml.encode("utf-8"))

    os.unlink(tmp_path)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    import docx.oxml.ns
    main()
