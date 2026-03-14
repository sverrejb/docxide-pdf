#!/usr/bin/env python3
"""case35: Adjusted shapes — preset shapes with non-default adjustment values.

Tests that the geometry formula evaluator correctly handles avLst overrides.
Each row shows the same shape twice: left = default adjustments, right = custom adjustments.
This makes visual comparison easy — the shape outline should visibly differ between columns.
"""

import os
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT = Path("tests/fixtures/cases/case35/input.docx")

# XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
GRAPHIC_DATA_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

# Shapes with (preset, color, label, adjustments_dict_or_None)
# None = use defaults, dict = override adjustments
# Each pair: default on left, adjusted on right
SHAPES = [
    # Row 1: roundRect — default corner radius vs very rounded
    ("roundRect", "4472C4", "roundRect\ndefault", None),
    ("roundRect", "2F5597", "roundRect\nadj=40000", {"adj": 40000}),
    # Row 2: star5 — default inner radius vs shallow star
    # star5 has 3 adjustments (adj, hf, vf) — must provide all when overriding
    ("star5", "BF8F00", "star5\ndefault", None),
    ("star5", "806000", "star5\nadj=35000", {"adj": 35000, "hf": 105146, "vf": 110557}),
    # Row 3: rightArrow — default vs narrow head + thin shaft
    ("rightArrow", "264478", "rightArrow\ndefault", None),
    ("rightArrow", "1B3050", "rightArrow\nadj1=25000\nadj2=75000", {"adj1": 25000, "adj2": 75000}),
    # Row 4: hexagon — default vs very tapered
    # hexagon has 2 adjustments (adj, vf) — must provide all when overriding
    ("hexagon", "5B9BD5", "hexagon\ndefault", None),
    ("hexagon", "2E75B6", "hexagon\nadj=10000", {"adj": 10000, "vf": 115470}),
    # Row 5: trapezoid — default vs nearly rectangular
    ("trapezoid", "FF6699", "trapezoid\ndefault", None),
    ("trapezoid", "CC3366", "trapezoid\nadj=10000", {"adj": 10000}),
    # Row 6: plus — default vs thin cross
    ("plus", "7030A0", "plus\ndefault", None),
    ("plus", "501878", "plus\nadj=10000", {"adj": 10000}),
    # Row 7: donut — default vs thin ring
    ("donut", "F4B183", "donut\ndefault", None),
    ("donut", "C48040", "donut\nadj=10000", {"adj": 10000}),
    # Row 8: frame — default vs thick frame
    ("frame", "002060", "frame\ndefault", None),
    ("frame", "001040", "frame\nadj1=30000", {"adj1": 30000}),
    # Row 9: chevron — default vs narrow chevron
    ("chevron", "A5A5A5", "chevron\ndefault", None),
    ("chevron", "757575", "chevron\nadj=25000", {"adj": 25000}),
    # Row 10: triangle — default (centered) vs skewed right
    ("triangle", "70AD47", "triangle\ndefault", None),
    ("triangle", "488030", "triangle\nadj=80000", {"adj": 80000}),
]

# Layout: 2 columns, 10 rows
EMU_PER_INCH = 914400
COLS = 2
ROWS = 10
SHAPE_W_EMU = int(2.5 * EMU_PER_INCH)
SHAPE_H_EMU = int(0.75 * EMU_PER_INCH)

CONTENT_W = 6.5 * EMU_PER_INCH
CONTENT_H = 9.0 * EMU_PER_INCH
COL_PITCH = CONTENT_W / COLS
ROW_PITCH = CONTENT_H / ROWS


def avlst_xml(adjustments):
    """Build avLst XML from adjustment dict or return empty avLst."""
    if not adjustments:
        return "<a:avLst/>"
    gds = "".join(
        f'<a:gd name="{name}" fmla="val {val}"/>'
        for name, val in adjustments.items()
    )
    return f"<a:avLst>{gds}</a:avLst>"


def shape_anchor_xml(idx, preset, color, label, adjustments):
    """Build wp:anchor XML for a single floating shape."""
    row = idx // COLS
    col = idx % COLS
    x_emu = int(col * COL_PITCH + (COL_PITCH - SHAPE_W_EMU) / 2)
    y_emu = int(row * ROW_PITCH + (ROW_PITCH - SHAPE_H_EMU) / 2)

    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    luma = 0.299 * r + 0.587 * g + 0.114 * b
    text_color = "FFFFFF" if luma < 140 else "000000"

    # Use single-line label (replace newlines with spaces) — matches case34 pattern
    flat_label = label.replace("\n", " ")

    return (
        f'<w:r>'
        f'<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f' simplePos="0" relativeHeight="{idx}" behindDoc="0"'
        f' locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="margin">'
        f'<wp:posOffset>{x_emu}</wp:posOffset>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="margin">'
        f'<wp:posOffset>{y_emu}</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{idx + 1}" name="Shape {idx + 1}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="{A_NS}">'
        f'<a:graphicData uri="{GRAPHIC_DATA_URI}">'
        f'<wps:wsp xmlns:wps="{WPS_NS}">'
        f'<wps:cNvSpPr/>'
        f'<wps:spPr>'
        f'<a:xfrm>'
        f'<a:off x="0" y="0"/>'
        f'<a:ext cx="{SHAPE_W_EMU}" cy="{SHAPE_H_EMU}"/>'
        f'</a:xfrm>'
        f'<a:prstGeom prst="{preset}">'
        f'{avlst_xml(adjustments)}'
        f'</a:prstGeom>'
        f'<a:solidFill>'
        f'<a:srgbClr val="{color}"/>'
        f'</a:solidFill>'
        f'<a:ln w="12700">'
        f'<a:solidFill>'
        f'<a:srgbClr val="000000"/>'
        f'</a:solidFill>'
        f'</a:ln>'
        f'</wps:spPr>'
        f'<wps:txbx>'
        f'<w:txbxContent xmlns:w="{W_NS}">'
        f'<w:p>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:sz w:val="16"/>'
        f'<w:color w:val="{text_color}"/>'
        f'</w:rPr>'
        f'<w:t>{flat_label}</w:t>'
        f'</w:r>'
        f'</w:p>'
        f'</w:txbxContent>'
        f'</wps:txbx>'
        f'<wps:bodyPr anchor="ctr" anchorCtr="0" lIns="0" tIns="0" rIns="0" bIns="0"/>'
        f'</wps:wsp>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
        f'</w:r>'
    )


# Step 1: Create base document with python-docx
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p = doc.add_paragraph()
run = p.add_run("SHAPE_PLACEHOLDER")

tmp = tempfile.mktemp(suffix=".docx")
doc.save(tmp)

# Step 2: Post-process ZIP to inject shape drawings
all_shapes_xml = "".join(
    shape_anchor_xml(i, preset, color, label, adjustments)
    for i, (preset, color, label, adjustments) in enumerate(SHAPES)
)

with zipfile.ZipFile(tmp, "r") as zin:
    with zipfile.ZipFile(str(OUT), "w", zipfile.ZIP_DEFLATED) as zout:
        doc_xml = zin.read("word/document.xml").decode()

        ns_decls = (
            f' xmlns:wp="{WP_NS}"'
            f' xmlns:a="{A_NS}"'
            f' xmlns:wps="{WPS_NS}"'
            f' xmlns:r="{R_NS}"'
            f' xmlns:mc="{MC_NS}"'
        )
        if 'xmlns:wp=' not in doc_xml:
            doc_xml = doc_xml.replace(
                '<w:document ',
                f'<w:document {ns_decls} ',
                1,
            )

        placeholder_pattern = r'<w:r>.*?<w:t>SHAPE_PLACEHOLDER</w:t>\s*</w:r>'
        doc_xml = re.sub(
            placeholder_pattern,
            all_shapes_xml,
            doc_xml,
            count=1,
            flags=re.DOTALL,
        )

        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml)
            else:
                zout.writestr(item, zin.read(item.filename))

os.unlink(tmp)
print(f"Generated {OUT}")
