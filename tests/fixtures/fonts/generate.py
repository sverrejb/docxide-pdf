#!/usr/bin/env python3
"""Generate font test fixtures — same content rendered in different fonts.

Each fixture exercises headings, subheadings, body text, and various sizes
to test font metrics, line spacing, and glyph rendering accuracy.

Usage:
    uv run tests/fixtures/fonts/generate.py                    # generate all
    uv run tests/fixtures/fonts/generate.py arial calibri      # generate specific fonts
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path("tests/fixtures/fonts")

FONTS = {
    "arial": "Arial",
    "times_new_roman": "Times New Roman",
    "calibri": "Calibri",
    "cambria": "Cambria",
    "georgia": "Georgia",
}

# Shared content — body paragraphs, headings, size samples
BODY_TEXT = (
    "The quick brown fox jumps over the lazy dog. "
    "Pack my box with five dozen liquor jugs. "
    "How vexingly quick daft zebras jump."
)

BODY_TEXT_2 = (
    "Typography is the art and technique of arranging type to make written language "
    "legible, readable, and appealing when displayed. The arrangement of type involves "
    "selecting typefaces, point sizes, line lengths, line spacing, and letter spacing."
)

BODY_TEXT_3 = (
    "Effective document rendering requires accurate measurement of glyph advance widths, "
    "ascender and descender metrics, and kerning pairs. Differences in font metrics between "
    "typefaces directly impact line breaking, page layout, and overall visual fidelity."
)

SIZE_SAMPLES = [8, 9, 10, 11, 12, 14, 16, 18, 24, 36]


def set_paragraph_font(paragraph, font_name, size_pt=None, bold=False, italic=False):
    """Set font on all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.name = font_name
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic


def add_run(paragraph, text, font_name, size_pt=None, bold=False, italic=False):
    """Add a run with specific font settings."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    return run


def generate_fixture(case_name, font_name):
    """Generate a DOCX fixture for the given font."""
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"Font Test: {font_name}", font_name, size_pt=28, bold=True)

    # --- Heading 1 ---
    p = doc.add_paragraph()
    add_run(p, "Section One: Introduction", font_name, size_pt=20, bold=True)

    p = doc.add_paragraph()
    add_run(p, BODY_TEXT, font_name, size_pt=12)

    p = doc.add_paragraph()
    add_run(p, BODY_TEXT_2, font_name, size_pt=12)

    # --- Heading 2 ---
    p = doc.add_paragraph()
    add_run(p, "1.1 Typography Fundamentals", font_name, size_pt=16, bold=True)

    p = doc.add_paragraph()
    add_run(p, BODY_TEXT_3, font_name, size_pt=12)

    p = doc.add_paragraph()
    add_run(
        p,
        "Italic text is commonly used for emphasis, technical terms, and foreign phrases. "
        "Bold italic combines both weights for strong emphasis within body copy.",
        font_name,
        size_pt=12,
        italic=True,
    )

    # --- Heading 2 ---
    p = doc.add_paragraph()
    add_run(p, "1.2 Mixed Formatting", font_name, size_pt=16, bold=True)

    p = doc.add_paragraph()
    add_run(p, "This paragraph contains ", font_name, size_pt=12)
    add_run(p, "bold", font_name, size_pt=12, bold=True)
    add_run(p, ", ", font_name, size_pt=12)
    add_run(p, "italic", font_name, size_pt=12, italic=True)
    add_run(p, ", and ", font_name, size_pt=12)
    add_run(p, "bold italic", font_name, size_pt=12, bold=True, italic=True)
    add_run(p, " runs within the same line to test run-level font switching.", font_name, size_pt=12)

    # --- Heading 1 ---
    p = doc.add_paragraph()
    add_run(p, "Section Two: Size Comparison", font_name, size_pt=20, bold=True)

    for size in SIZE_SAMPLES:
        p = doc.add_paragraph()
        add_run(p, f"{size}pt — The quick brown fox jumps over the lazy dog.", font_name, size_pt=size)

    # --- Heading 1 ---
    p = doc.add_paragraph()
    add_run(p, "Section Three: Longer Body Text", font_name, size_pt=20, bold=True)

    # --- Heading 2 ---
    p = doc.add_paragraph()
    add_run(p, "3.1 Document Layout", font_name, size_pt=16, bold=True)

    p = doc.add_paragraph()
    add_run(
        p,
        "Word processors must calculate precise line breaks based on the advance widths of "
        "individual glyphs, taking into account kerning pairs and contextual adjustments. "
        "The resulting layout depends heavily on accurate font metrics — even small "
        "discrepancies in glyph width measurements compound across a line, leading to "
        "different break points and misaligned text.",
        font_name,
        size_pt=12,
    )

    p = doc.add_paragraph()
    add_run(
        p,
        "Paragraph spacing, both before and after, follows collapsing rules similar to CSS "
        "margin collapsing: the space between two consecutive paragraphs is the maximum of "
        "the first paragraph's space-after and the second paragraph's space-before, not the "
        "sum. This behavior must be replicated precisely for accurate rendering.",
        font_name,
        size_pt=12,
    )

    # --- Heading 2 ---
    p = doc.add_paragraph()
    add_run(p, "3.2 Character Coverage", font_name, size_pt=16, bold=True)

    p = doc.add_paragraph()
    add_run(
        p,
        'Numbers and punctuation: 0123456789 !@#$%^&*() "quotes" \'apostrophes\' — '
        "em-dash, semi-colon; colon: ellipsis... brackets [square] {curly} (round).",
        font_name,
        size_pt=12,
    )

    p = doc.add_paragraph()
    add_run(
        p,
        "Common ligatures and pairs: fi fl ff ffi ffl AV AW WA To Ta Te "
        "Yo Ya Yd LT Ty Tw Tc. Kerning quality varies significantly across these pairs.",
        font_name,
        size_pt=12,
    )

    # --- Right-aligned closing ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f"— End of {font_name} specimen —", font_name, size_pt=10, italic=True)

    out_path = OUT_DIR / case_name / "input.docx"
    doc.save(str(out_path))
    print(f"Generated {out_path}")


def main():
    requested = sys.argv[1:] if len(sys.argv) > 1 else list(FONTS.keys())

    for case_name in requested:
        if case_name not in FONTS:
            print(f"Unknown font case: {case_name} (available: {', '.join(FONTS.keys())})")
            continue
        case_dir = OUT_DIR / case_name
        case_dir.mkdir(parents=True, exist_ok=True)
        generate_fixture(case_name, FONTS[case_name])


if __name__ == "__main__":
    main()
