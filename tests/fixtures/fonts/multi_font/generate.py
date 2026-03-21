#!/usr/bin/env python3
"""Generate a fixture with many different fonts, one per line.

Each line is regular 12pt body text set in a different font,
with the font name as the text content.

Usage:
    uv run tests/fixtures/fonts/multi_font/generate.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

OUT_DIR = Path(__file__).parent

FONTS = [
    "Arial",
    "Times New Roman",
    "Calibri",
    "Cambria",
    "Georgia",
    "Verdana",
    "Trebuchet MS",
    "Tahoma",
    "Garamond",
    "Palatino Linotype",
    "Book Antiqua",
    "Century Gothic",
    "Lucida Sans Unicode",
    "Courier New",
    "Consolas",
    "Segoe UI",
    "Candara",
    "Constantia",
    "Corbel",
    "Franklin Gothic Medium",
    "Gill Sans MT",
    "Rockwell",
    "Bodoni MT",
    "Copperplate Gothic Light",
    "Bookman Old Style",
    "Aptos",
]

SAMPLE_TEXT = "The quick brown fox jumps over the lazy dog. 0123456789"


def main():
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for font_name in FONTS:
        p = doc.add_paragraph()
        run = p.add_run(f"{font_name}: {SAMPLE_TEXT}")
        run.font.name = font_name
        run.font.size = Pt(12)

    out_path = OUT_DIR / "input.docx"
    doc.save(str(out_path))
    print(f"Generated {out_path}")


if __name__ == "__main__":
    main()
