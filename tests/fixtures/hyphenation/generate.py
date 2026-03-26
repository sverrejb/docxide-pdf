#!/usr/bin/env python3
"""Generate hyphenation test fixtures in various languages.

Each fixture uses real Wikipedia article text (CC BY-SA 4.0) with
autoHyphenation enabled and narrow margins to force frequent line breaks.
Texts are loaded from texts.json (fetched from Wikipedia REST API).

Sources per language:
  english:   Electroencephalography, Psychopharmacology
  german:    Elektroenzephalografie, Bundesverfassungsgericht
  french:    Électroencéphalographie, Psychopharmacologie, Typographie
  spanish:   Electroencefalografía, Otorrinolaringología
  norwegian: Stortinget, Norges Grunnlov
  dutch:     Elektro-encefalografie, Typografie
  italian:   Elettroencefalografia, Psicofarmacologia
  slovak:    Elektroencefalografia, Ústava SR, Slovensko, Bratislava

Usage:
    uv run tests/fixtures/hyphenation/generate.py                # generate all
    uv run tests/fixtures/hyphenation/generate.py english german  # specific ones
"""

import json
import sys
import zipfile
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path("tests/fixtures/hyphenation")


def generate_fixture(name, config):
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.75)
        section.right_margin = Inches(1.75)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(config["title"])
    run.font.name = config["font"]
    run.font.size = Pt(16)
    run.bold = True

    # Body paragraphs (justified)
    for text in config["paragraphs"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = config["font"]
        run.font.size = Pt(12)

    # Save initial docx
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # ZIP post-processing: inject autoHyphenation + w:lang on runs
    case_dir = OUT_DIR / name
    case_dir.mkdir(parents=True, exist_ok=True)
    out_path = case_dir / "input.docx"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(buf.read())
        tmp_path = tmp.name

    lang = config["lang"]

    with zipfile.ZipFile(tmp_path, "r") as zin:
        with zipfile.ZipFile(str(out_path), "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == "word/settings.xml":
                    text_data = data.decode("utf-8")
                    text_data = text_data.replace(
                        "</w:settings>",
                        "<w:autoHyphenation/></w:settings>",
                    )
                    data = text_data.encode("utf-8")

                elif item.filename == "word/document.xml":
                    text_data = data.decode("utf-8")
                    text_data = text_data.replace(
                        "</w:rPr>",
                        f'<w:lang w:val="{lang}"/></w:rPr>',
                    )
                    data = text_data.encode("utf-8")

                zout.writestr(item, data)

    Path(tmp_path).unlink()
    print(f"Generated {out_path}")


def main():
    texts_path = OUT_DIR / "texts.json"
    with open(texts_path) as f:
        fixtures = json.load(f)

    requested = sys.argv[1:] if len(sys.argv) > 1 else list(fixtures.keys())

    for name in requested:
        if name not in fixtures:
            print(f"Unknown fixture: {name} (available: {', '.join(fixtures.keys())})")
            continue
        generate_fixture(name, fixtures[name])


if __name__ == "__main__":
    main()
